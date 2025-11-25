import os, json, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(__file__)
CONTACT_FILE = os.path.join(ROOT, 'contact_info.json')
SHORT_TXT = os.path.join(ROOT, 'tailored_resume_3M_short.txt')  # base content
EXP_JSON = os.path.join(ROOT, 'experience_clean.json')
OUT_STYLED = os.path.join(ROOT, 'tailored_resume_3M_modern.docx')
OUT_ATS = os.path.join(ROOT, 'tailored_resume_3M_modern_ats.docx')

ACCENT = RGBColor(0x12, 0x34, 0x56)
FONT_PRIMARY = 'Calibri'
FONT_ATS = 'Arial'
STYLE_SPEC_PATH = os.path.join(ROOT, 'style_spec.json')

def load_style_spec():
    if not os.path.exists(STYLE_SPEC_PATH):
        return {}
    try:
        with open(STYLE_SPEC_PATH,'r',encoding='utf-8') as f:
            return json.load(f)
    except (OSError, json.JSONDecodeError):
        return {}

ACTION_VERBS = [
    'Accelerated','Achieved','Analyzed','Built','Collaborated','Created','Cut','Delivered','Designed','Drove','Enhanced',
    'Improved','Implemented','Led','Optimized','Reduced','Resolved','Standardized','Streamlined','Strengthened','Supported'
]

def load_contact():
    try:
        with open(CONTACT_FILE,'r',encoding='utf-8') as f:
            return json.load(f)
    except (OSError, json.JSONDecodeError):
        return {}

def read_lines(path):
    with open(path,'r',encoding='utf-8') as f:
        return [l.rstrip() for l in f]

def load_experience():
    """Load pre-cleaned experience entries from JSON."""
    if not os.path.exists(EXP_JSON):
        return []
    try:
        with open(EXP_JSON,'r',encoding='utf-8') as f:
            data = json.load(f)
    except (OSError, json.JSONDecodeError):
        return []
    # Data is already clean and structured
    return data[:3]

def sanitize_bullet(text, spec):
    max_words = spec.get('bullet_rules',{}).get('max_words',28)
    t = re.sub(r'^[-•*]\s*','',text).strip()
    words = t.split()
    if not words:
        return t
    # Preserve existing starting verb if in ACTION_VERBS; otherwise just capitalize first word without forcing verb
    first = words[0]
    # Capitalize first word (avoid forcing new verb insertion)
    words[0] = first.capitalize()
    # Truncate
    if len(words) > max_words:
        t = ' '.join(words[:max_words]) + spec.get('bullet_rules',{}).get('truncate_suffix','…')
    else:
        t = ' '.join(words)
    if spec.get('bullet_rules',{}).get('remove_trailing_punctuation', True):
        t = t.rstrip('.;')
    return t

def add_dans_metadata(doc, contact):
    """Add DANS-compliant document metadata for digital application navigation."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - Professional Resume"
    core_props.subject = "Resume - ATS & DANS Optimized"
    core_props.keywords = "resume, ATS, DANS, professional, application"
    core_props.category = "Resume"
    core_props.comments = "Digital Application Navigation System (DANS) compliant with ATS optimization"

def configure_dans_layout(doc, spec):
    """Configure DANS-compliant page layout with standard dimensions and margins."""
    section = doc.sections[0]
    # DANS: Standard page dimensions
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    # Margins from spec or DANS standard
    m = spec.get('margins_inches', {})
    section.top_margin = Inches(m.get('top', 0.75))  # DANS standard
    section.bottom_margin = Inches(m.get('bottom', 0.75))
    section.left_margin = Inches(m.get('left', 0.75))
    section.right_margin = Inches(m.get('right', 0.75))

def apply_margins(doc, spec):
    """Legacy margin function - now delegates to DANS layout."""
    configure_dans_layout(doc, spec)

def build_header(doc, contact, accent=True, ats=False, spec=None):
    """Build header with DANS-compliant Heading 1 style."""
    fonts = spec.get('fonts', {}) if spec else {}
    primary_font = fonts.get('primary', FONT_PRIMARY)
    ats_font = fonts.get('ats', FONT_ATS)
    name_size = fonts.get('name_size', 20 if not ats else 18)
    # Use style spec accent or fallback to default hex values
    accent_rgb = fonts.get('header_accent_rgb', [0x12, 0x34, 0x56])
    name_p = doc.add_paragraph()
    # DANS: Use Heading 1 for proper document hierarchy
    name_p.style = 'Heading 1'
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(contact.get('name','Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(name_size if not ats else max(name_size-2, 16))
    r.font.name = ats_font if ats else primary_font
    if accent and not ats and isinstance(accent_rgb, list) and len(accent_rgb)==3:
        r.font.color.rgb = RGBColor(*accent_rgb)
    else:
        # DANS: Explicit color for accessibility
        r.font.color.rgb = RGBColor(0, 0, 0)
    line2_parts = [contact.get('address',''), contact.get('location_statement','')]
    line2 = ' • '.join([p for p in line2_parts if p])
    if line2:
        p2 = doc.add_paragraph(line2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in p2.runs:
            rr.font.size = Pt(fonts.get('body_size',9))
            rr.font.name = ats_font if ats else primary_font
    line3_parts = [contact.get('email',''), contact.get('phone',''), contact.get('linkedin',''), contact.get('github','')]
    line3 = ' • '.join([p for p in line3_parts if p])
    p3 = doc.add_paragraph(line3)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rr in p3.runs:
        rr.font.size = Pt(fonts.get('body_size',9))
        rr.font.name = ats_font if ats else primary_font

def add_section_title(doc, title, ats=False):
    """Add DANS-compliant section heading with Heading 2 style."""
    p = doc.add_paragraph()
    # DANS: Use Heading 2 for proper document hierarchy
    p.style = 'Heading 2'
    # DANS: Uppercase for clarity and consistency
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(11 if ats else 12)
    r.font.name = FONT_ATS if ats else FONT_PRIMARY
    if not ats:
        r.font.color.rgb = ACCENT
    else:
        # DANS: Explicit color for accessibility
        r.font.color.rgb = RGBColor(0, 0, 0)
    # Add spacing after title
    p.paragraph_format.space_after = Pt(4)

def extract_summary(lines):
    out=[]; capture=False
    for l in lines:
        if l.strip().upper().startswith('SUMMARY'): capture=True; continue
        if capture:
            if l.isupper() and len(l.split())<6: break
            if not l.strip(): continue
            out.append(l)
    # Condense to 3 sentences max
    text = ' '.join(out)
    parts = re.split(r'(?<=[.!?])\s+', text)
    return ' '.join(parts[:3])

def extract_core_skills(lines):
    skills=[]; capture=False
    for l in lines:
        if l.strip().upper().startswith('CORE SKILLS'): capture=True; continue
        if capture:
            if l.isupper() and len(l.split())<6: break
            if l.lower().startswith('emerging/'): continue
            skills.extend([p.strip() for p in re.split(r'\s*•\s*|,\s*', l) if p.strip()])
    # Deduplicate and cap
    uniq=[]
    for s in skills:
        if s not in uniq: uniq.append(s)
    return uniq[:12]

def write_skills(doc, skills, ats=False, spec=None):
    add_section_title(doc,'Core Skills', ats=ats)
    line = ' • '.join(skills)
    p = doc.add_paragraph(line)
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.name = FONT_ATS if ats else FONT_PRIMARY
    p.paragraph_format.space_after = Pt(spec.get('spacing',{}).get('skills_after_pt',6) if spec else 6)

def write_experience(doc, entries, ats=False, spec=None):
    add_section_title(doc,'Professional Experience', ats=ats)
    spacing = spec.get('spacing',{}) if spec else {}
    for e in entries:
        # Use 2-column table for alignment (dates right-aligned)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell, right_cell = table.rows[0].cells
        # Format: Company | Title
        company = e.get('company','').strip()
        title = e.get('title','').strip()
        location = e.get('location','').strip()
        dates = e.get('dates','').strip()
        
        header_text_left = f"{company} | {title}"
        header_text_right = f"{dates}"
        if location:
            header_text_right = f"{location} | {dates}"
        
        left_p = left_cell.paragraphs[0]
        lrun = left_p.add_run(header_text_left)
        lrun.bold = True
        lrun.font.size = Pt(10)
        lrun.font.name = FONT_ATS if ats else FONT_PRIMARY
        
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rrun = right_p.add_run(header_text_right)
        rrun.bold = True
        rrun.font.size = Pt(10)
        rrun.font.name = FONT_ATS if ats else FONT_PRIMARY
        
        # Bullets
        bullets = e.get('bullets',[])[:6]
        for b in bullets:
            sb = sanitize_bullet(b, spec)
            p = doc.add_paragraph(style='List Bullet' if not ats else None)
            p.add_run(sb)
            for rr in p.runs:
                rr.font.size = Pt(10)
                rr.font.name = FONT_ATS if ats else FONT_PRIMARY
            p.paragraph_format.space_after = Pt(spacing.get('bullet_after_pt',2))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(spacing.get('section_after_pt',4))



def write_education(doc, lines, ats=False, spec=None):
    # Extract from original file
    add_section_title(doc,'Education', ats=ats)
    capture=False; out=[]
    for l in lines:
        if l.strip().upper().startswith('EDUCATION'): capture=True; continue
        if capture:
            if l.isupper() and len(l.split())<6 and not l.strip().startswith('Bachelor'): break
            if not l.strip(): continue
            out.append(l)
    for line in out[:3]:
        p=doc.add_paragraph(line)
        for r in p.runs:
            r.font.size=Pt(10)
            r.font.name = FONT_ATS if ats else FONT_PRIMARY
        p.paragraph_format.space_after = Pt(spec.get('spacing',{}).get('bullet_after_pt',2) if spec else 2)

def build_styled(contact, base_lines, exp_entries, spec):
    doc = Document()
    add_dans_metadata(doc, contact)  # DANS compliance
    apply_margins(doc, spec)
    build_header(doc, contact, accent=True, ats=False, spec=spec)
    add_section_title(doc,'Summary')
    summary = extract_summary(base_lines)
    sp = doc.add_paragraph(summary)
    for r in sp.runs:
        r.font.size = Pt(10); r.font.name = FONT_PRIMARY
    skills = extract_core_skills(base_lines)
    write_skills(doc, skills, ats=False, spec=spec)
    write_experience(doc, exp_entries, ats=False, spec=spec)
    write_education(doc, base_lines, ats=False, spec=spec)
    doc.save(OUT_STYLED)
    return OUT_STYLED

def build_ats(contact, base_lines, exp_entries, spec):
    doc = Document()
    add_dans_metadata(doc, contact)  # DANS compliance
    apply_margins(doc, spec)
    build_header(doc, contact, accent=False, ats=True, spec=spec)
    add_section_title(doc,'Summary', ats=True)
    summary = extract_summary(base_lines)
    sp = doc.add_paragraph(summary)
    for r in sp.runs:
        r.font.size = Pt(10); r.font.name = FONT_ATS
    skills = extract_core_skills(base_lines)
    write_skills(doc, skills, ats=True, spec=spec)
    write_experience(doc, exp_entries, ats=True, spec=spec)
    write_education(doc, base_lines, ats=True, spec=spec)
    out = OUT_ATS
    try:
        doc.save(out)
    except PermissionError:
        alt = out.replace('.docx','_updated.docx')
        doc.save(alt)
        out = alt
    return out

def main():
    contact = load_contact()
    lines = read_lines(SHORT_TXT)
    exp_entries = load_experience()
    spec = load_style_spec()
    styled = build_styled(contact, lines, exp_entries, spec)
    ats = build_ats(contact, lines, exp_entries, spec)
    print('Modern resumes generated:', styled, ats)

if __name__ == '__main__':
    main()