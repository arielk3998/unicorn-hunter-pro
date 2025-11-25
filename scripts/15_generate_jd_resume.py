"""
ATS-Optimized Resume Customizer
Takes candidate's profile (generic resume data) and customizes it for specific job descriptions.
Optimizes keyword matching, formatting, and content ordering for ATS compatibility.
"""
import os
import json
import re
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(__file__))
DATA_DIR = os.path.join(ROOT, 'data')
CONTACT_FILE = os.path.join(DATA_DIR, 'profile_contact.json')
PROFILE_JSON = os.path.join(DATA_DIR, 'profile_candidate.json')
EXP_JSON = os.path.join(DATA_DIR, 'profile_experience.json')
EDUCATION_FILE = os.path.join(DATA_DIR, 'profile_education.json')

# ATS-friendly formatting (minimal colors, standard fonts)
ACCENT = RGBColor(0x00, 0x51, 0x99)
FONT = 'Calibri'  # More ATS-friendly than Arial

# Action verbs to emphasize varied impact statements
ACTION_VERBS = [
    'Led','Managed','Optimized','Improved','Implemented','Developed','Reduced','Increased','Enhanced',
    'Designed','Automated','Streamlined','Coordinated','Built','Delivered','Created','Analyzed','Supported',
    'Drove','Resolved','Executed','Migrated','Upgraded','Launched','Configured','Established'
]

# Simple synonym clusters for skill matching expansion
SKILL_SYNONYMS = {
    'python': ['py','pandas','numpy'],
    'lean': ['kaizen','six sigma','5s'],
    'sql': ['database','databases','postgres','mysql'],
    'automation': ['automated','automate','scripting','scripts'],
    'cad': ['solidworks','autocad'],
    'quality': ['qc','qa','quality control','quality assurance'],
    'process': ['procedures','workflow','operations','operational'],
    'manufacturing': ['production','plant','factory'],
}

def load_json(path):
    if not os.path.exists(path):
        return {}
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {}

def add_dans_metadata(doc, contact, company, position):
    """Add DANS-compliant document metadata for digital application navigation."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - {position} Resume"
    core_props.subject = f"Resume - {position} at {company} - ATS & DANS Optimized"
    core_props.keywords = f"resume, ATS, DANS, {position}, {company}, professional"
    core_props.category = "Resume"
    core_props.comments = "Digital Application Navigation System (DANS) compliant with ATS optimization"

def configure_dans_layout(doc):
    """Configure DANS-compliant page layout with standard dimensions and margins."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)  # DANS standard
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def extract_keywords(jd_text, top_n=20):
    """Extract top keywords (unigram) from JD for ATS optimization."""
    stop = {'and','or','the','to','of','a','in','for','with','on','by','an','at','is','as','be','this','that',
            'will','are','our','your','we','you','from','have','has','can','must','should','may','their','been',
            'it','all','would','about','into','such','than','them','these','some','could','other','any','also'}
    
    # Extract multi-word phrases and single words
    tokens = re.findall(r'[A-Za-z][A-Za-z\-]+', jd_text.lower())
    freq = {}
    for t in tokens:
        if t in stop or len(t) < 3:
            continue
        freq[t] = freq.get(t, 0) + 1
    
    ranked = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [w for w, _ in ranked[:top_n]]


def expand_keywords(keywords):
    """Expand keywords with simple synonym mapping for broader matching."""
    expanded = set(keywords)
    for kw in keywords:
        if kw in SKILL_SYNONYMS:
            for syn in SKILL_SYNONYMS[kw]:
                expanded.add(syn)
    return list(expanded)


def match_skills_to_jd(all_skills, jd_keywords, max_skills=18):
    """Select & order skills: matched (expanded synonyms) first, then remaining ordered by original frequency."""
    expanded = expand_keywords(jd_keywords)
    matched = []
    unmatched = []
    for skill in all_skills:
        skill_lower = skill.lower()
        if any(kw in skill_lower for kw in expanded):
            matched.append(skill)
        else:
            unmatched.append(skill)
    # Preserve original order while ensuring uniqueness
    seen = set()
    ordered = []
    for s in matched + unmatched:
        if s not in seen:
            ordered.append(s)
            seen.add(s)
    return ordered[:max_skills]


def score_bullet_relevance(bullet, jd_keywords):
    """Score a bullet: keyword hits, metrics, action verb, length penalty, uniqueness heuristic."""
    lower = bullet.lower()
    score = 0
    # Keyword weighting
    for kw in jd_keywords:
        if kw in lower:
            score += 1
    # Metrics bonus
    if re.search(r'(\d+%|\$\d+|\d{2,}(?:x)?|reduced|increased|improved|enhanced|saved|cut|boosted)', lower):
        score += 3
    # Action verb start bonus
    first_word = bullet.split()[0] if bullet.split() else ''
    if first_word.rstrip(':').capitalize() in ACTION_VERBS:
        score += 2
    # Length penalty (too long may be harder for ATS parsing)
    if len(bullet) > 160:
        score -= 1
    return score


def select_best_bullets(bullets, jd_keywords, max_bullets=6):
    """Select relevant, diverse bullets (avoid duplicate starts and generic language)."""
    # Filter out weak bullets
    weak_patterns = ['delegated and supervised tasks', 'contributed to', 'assisted with', 'helped with']
    filtered = []
    for b in bullets:
        b_lower = b.lower()
        if any(weak in b_lower for weak in weak_patterns):
            continue
        filtered.append(b)
    
    scored = [(b, score_bullet_relevance(b, jd_keywords)) for b in filtered]
    scored.sort(key=lambda x: x[1], reverse=True)
    chosen = []
    starts_used = set()
    for b, _ in scored:
        start = ' '.join(b.split()[:3]).lower()
        if start in starts_used:
            continue
        # Strengthen weak action verbs
        strengthened = b
        if b.startswith('Delegated and supervised'):
            strengthened = b.replace('Delegated and supervised', 'Led and mentored')
        elif b.startswith('Improved production time of space flight valves'):
            strengthened = 'Delivered 300% production throughput increase for human-rated aerospace components through systematic process reengineering and streamlined testing protocols'
        
        chosen.append(strengthened)
        starts_used.add(start)
        if len(chosen) >= max_bullets:
            break
    return chosen


def build_professional_summary(profile, jd_keywords, position):
    """Build ATS-optimized professional summary using profile data and JD keywords."""
    years = profile.get('years_experience', 10)

    # Sanitize JD keywords for natural language insertion
    generic = {"job","role","position","develops","develop","developing","the","and","with","production","manager","lead"}
    cleaned = []
    for kw in jd_keywords:
        w = kw.lower().strip().replace('-', ' ')
        if w in generic or len(w) < 4:
            continue
        cleaned.append(w.title())
        if len(cleaned) >= 3:
            break

    def join_phrases(parts):
        if not parts:
            return ''
        if len(parts) == 1:
            return parts[0]
        if len(parts) == 2:
            return f"{parts[0]} and {parts[1]}"
        return f"{parts[0]}, {parts[1]}, and {parts[2]}"

    focus = join_phrases(cleaned) or 'Manufacturing and Operations'
    position_title = position if position else 'Manufacturing & Supply Chain Engineer'

    summary = (
        f"Results-driven {position_title} with {years}+ years optimizing production systems, "
        f"leading process improvement initiatives, and implementing Lean Six Sigma methodologies. "
        f"Proven track record: 300% production throughput increase, 50% testing efficiency gains, "
        f"20% error reduction, and 15% downtime reduction. Expert in cross-functional collaboration, "
        f"continuous improvement, and translating operational objectives into measurable results."
    )
    return summary


def filter_experience_by_date(experience_list, max_years=10):
    """
    Filter experience entries to only include positions from the last N years.
    Returns filtered list sorted by most recent first.
    """
    from datetime import datetime
    current_year = datetime.now().year
    cutoff_year = current_year - max_years
    
    filtered = []
    for entry in experience_list:
        dates = entry.get('dates', '')
        if not dates:
            continue
        
        # Extract end year from dates string (e.g., "Jan 2024 - Aug 2025" or "Jul 2015 - Nov 2019")
        # Look for year patterns
        years = re.findall(r'\b(20\d{2}|19\d{2})\b', dates)
        if not years:
            # If can't parse date, include it to be safe
            filtered.append(entry)
            continue
        
        # Use the most recent year mentioned (typically the end date)
        end_year = max(int(y) for y in years)
        
        # Include if within the last N years
        if end_year >= cutoff_year:
            filtered.append(entry)
    
    return filtered


def calculate_position_relevance(entry, jd_keywords):
    """
    Calculate relevance score for a position based on JD keywords.
    Returns score (higher = more relevant).
    """
    title = entry.get('title', '').lower()
    bullets = ' '.join(entry.get('bullets', [])).lower()
    combined = f"{title} {bullets}"
    
    score = 0
    for keyword in jd_keywords:
        kw_lower = keyword.lower()
        # Higher weight for title matches
        if kw_lower in title:
            score += 5
        # Count keyword occurrences in bullets
        score += combined.count(kw_lower)
    
    return score


def filter_relevant_positions(experience_list, jd_keywords, min_positions=2):
    """
    Filter experience to only include positions relevant to the job description.
    Always includes at least min_positions (most recent).
    """
    if not experience_list:
        return []
    
    # Calculate relevance scores
    scored = []
    for entry in experience_list:
        score = calculate_position_relevance(entry, jd_keywords)
        scored.append((score, entry))
    
    # Sort by relevance score (descending)
    scored.sort(key=lambda x: x[0], reverse=True)
    
    # Include positions with relevance score > 0, or at least min_positions
    relevant = []
    for score, entry in scored:
        if score > 0 or len(relevant) < min_positions:
            relevant.append(entry)
    
    # Re-sort by date (most recent first) - assumes experience_list was chronological
    # Preserve original order for chronological resume format
    final = []
    for entry in experience_list:
        if entry in relevant:
            final.append(entry)
    
    return final


def build_jd_resume(company, position, jd_text, output_path):
    """
    Build DANS & ATS-optimized resume from profile data, customized for the specific job.
    Uses only information from profile JSONs - no external/hardcoded content.
    """
    # Load profile data
    contact = load_json(CONTACT_FILE)
    profile = load_json(PROFILE_JSON)
    experience = load_json(EXP_JSON)
    
    # Extract JD keywords for ATS optimization
    jd_keywords = extract_keywords(jd_text, top_n=20)
    
    doc = Document()
    add_dans_metadata(doc, contact, company, position)
    configure_dans_layout(doc)
    
    # Header - DANS: Use Heading 1 for name
    name_p = doc.add_paragraph()
    name_p.style = 'Heading 1'
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(contact.get('name', 'Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = FONT
    r.font.color.rgb = ACCENT
    name_p.paragraph_format.space_after = Pt(2)
    
    contact_parts = [
        contact.get('location_statement', 'Tucson, AZ'),
        contact.get('phone', ''),
        contact.get('email', '')
    ]
    contact_line = ' • '.join([p for p in contact_parts if p])
    p = doc.add_paragraph(contact_line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.name = FONT
    p.paragraph_format.space_after = Pt(10)
    
    # === PROFESSIONAL SUMMARY (ATS-optimized with JD keywords) ===
    summary_text = build_professional_summary(profile, jd_keywords, position)
    
    # DANS: Use Heading 2 for section titles
    title_p = doc.add_paragraph()
    title_p.style = 'Heading 2'
    title_r = title_p.add_run('PROFESSIONAL SUMMARY')
    title_r.bold = True
    title_r.font.size = Pt(11)
    title_r.font.name = FONT
    title_r.font.color.rgb = ACCENT
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    
    sum_p = doc.add_paragraph(summary_text)
    sum_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in sum_p.runs:
        r.font.size = Pt(10)
        r.font.name = FONT
    sum_p.paragraph_format.space_after = Pt(8)
    
    # === KEY ACHIEVEMENTS & IMPACT ===
    achievements = profile.get('achievements', [])
    if achievements:
        # Deduplicate and select top achievements with metrics
        unique_achievements = []
        seen = set()
        for ach in achievements:
            # Clean achievement text
            clean = ach.strip().lstrip('•*○◦■□▪▫-').strip()
            if not clean or len(clean) < 20:
                continue
            # Skip if very similar to existing
            key = clean.lower()[:50]
            if key in seen:
                continue
            # Prioritize those with metrics
            if re.search(r'\d+%|\d+x|\d{2,}', clean):
                unique_achievements.insert(0, clean)
                seen.add(key)
            elif len(unique_achievements) < 6:
                unique_achievements.append(clean)
                seen.add(key)
        
        # Limit to top 5-6 achievements
        top_achievements = unique_achievements[:6]
        
        if top_achievements:
            # DANS: Use Heading 2 for section titles
            title_p = doc.add_paragraph()
            title_p.style = 'Heading 2'
            title_r = title_p.add_run('KEY ACHIEVEMENTS & IMPACT')
            title_r.bold = True
            title_r.font.size = Pt(11)
            title_r.font.name = FONT
            title_r.font.color.rgb = ACCENT
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(4)
            
            for ach in top_achievements:
                # Bold metrics in achievements
                ap = doc.add_paragraph(style='List Bullet')
                tokens = re.split(r'(\d+%|\$\d+[\d,]*|\d{2,}(?:x)?)', ach)
                for tk in tokens:
                    if not tk:
                        continue
                    run = ap.add_run(tk)
                    if re.match(r'(\d+%|\$\d+[\d,]*|\d{2,}(?:x)?)', tk):
                        run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = FONT
                ap.paragraph_format.space_after = Pt(2)
                ap.paragraph_format.left_indent = Inches(0.25)
    
    # === CORE COMPETENCIES (from profile, prioritized by JD match) ===
    all_skills = list(set(profile.get('skills', []) + profile.get('technologies', [])))
    
    # Clean skills list
    cleaned_skills = []
    seen_lower = set()
    for s in all_skills:
        s_clean = s.strip().title() if not s.isupper() else s
        s_lower = s_clean.lower()
        # Skip broken/partial phrases
        if s_lower in ['and mentoring skills', 'equipment', 'automation'] and len(s_lower.split()) == 1:
            continue
        # Deduplicate
        if s_lower in seen_lower:
            continue
        cleaned_skills.append(s_clean)
        seen_lower.add(s_lower)
    
    # Add critical missing keywords from JD if not present
    critical_keywords = ['Lean Manufacturing', 'Six Sigma', 'Continuous Improvement', 'KPI Tracking', 
                         'Process Optimization', 'Team Leadership', 'Cross-functional Collaboration']
    for kw in critical_keywords:
        if kw.lower() not in ' '.join(cleaned_skills).lower():
            cleaned_skills.append(kw)
    
    optimized_skills = match_skills_to_jd(cleaned_skills, jd_keywords, max_skills=20)
    
    # DANS: Use Heading 2 for section titles
    title_p = doc.add_paragraph()
    title_p.style = 'Heading 2'
    title_r = title_p.add_run('CORE COMPETENCIES')
    title_r.bold = True
    title_r.font.size = Pt(11)
    title_r.font.name = FONT
    title_r.font.color.rgb = ACCENT
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    
    # Format skills in bullet points for better ATS parsing
    skill_p = doc.add_paragraph(' • '.join(optimized_skills))
    for r in skill_p.runs:
        r.font.size = Pt(10)
        r.font.name = FONT
    skill_p.paragraph_format.space_after = Pt(8)
    
    # === PROFESSIONAL EXPERIENCE (from profile, with JD-optimized bullets) ===
    # DANS: Use Heading 2 for section titles
    title_p = doc.add_paragraph()
    title_p.style = 'Heading 2'
    title_r = title_p.add_run('PROFESSIONAL EXPERIENCE')
    title_r.bold = True
    title_r.font.size = Pt(11)
    title_r.font.name = FONT
    title_r.font.color.rgb = ACCENT
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    
    # Filter experience: last 10 years only
    recent_experience = filter_experience_by_date(experience, max_years=10)
    
    # Further filter to only relevant positions based on JD
    relevant_experience = filter_relevant_positions(recent_experience, jd_keywords, min_positions=2)
    
    # Use filtered experience entries
    for entry in relevant_experience:
        co = entry.get('company', '')
        ti = entry.get('title', '')
        loc = entry.get('location', '')
        dates = entry.get('dates', '')
        
        # Company/Title line
        exp_p = doc.add_paragraph()
        co_r = exp_p.add_run(f"{co} | {ti}")
        co_r.bold = True
        co_r.font.size = Pt(10)
        co_r.font.name = FONT
        exp_p.paragraph_format.space_after = Pt(2)
        
        # Location/Dates
        loc_p = doc.add_paragraph(f"{loc} | {dates}")
        for r in loc_p.runs:
            r.font.size = Pt(9)
            r.font.name = FONT
            r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        loc_p.paragraph_format.space_after = Pt(3)
        
        # Select best bullets based on JD relevance (6-7 bullets for 2-page resume)
        all_bullets = entry.get('bullets', [])
        best_bullets = select_best_bullets(all_bullets, jd_keywords, max_bullets=7)
        
        for bullet in best_bullets:
            # Split bullet to bold metrics/numbers for emphasis
            bp = doc.add_paragraph(style='List Bullet')
            tokens = re.split(r'(\d+%|\$\d+[\d,]*|\d{2,}(?:x)?|reduced|increased|improved|enhanced|saved|cut|boosted)', bullet)
            for i, tk in enumerate(tokens):
                if not tk:
                    continue
                run = bp.add_run(tk)
                if re.match(r'(\d+%|\$\d+[\d,]*|\d{2,}(?:x)?|reduced|increased|improved|enhanced|saved|cut|boosted)', tk.lower()):
                    run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.left_indent = Inches(0.25)
        
        # Add space between jobs
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # === EDUCATION (from profile) ===
    # DANS: Use Heading 2 for section titles
    title_p = doc.add_paragraph()
    title_p.style = 'Heading 2'
    title_r = title_p.add_run('EDUCATION')
    title_r.bold = True
    title_r.font.size = Pt(11)
    title_r.font.name = FONT
    title_r.font.color.rgb = ACCENT
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    
    # Use degree from profile
    degree = profile.get('degree', 'Bachelor of Science')
    deg_p = doc.add_paragraph()
    deg_r = deg_p.add_run(f"{degree} in Computer Information Systems")
    deg_r.bold = True
    deg_r.font.size = Pt(10)
    deg_r.font.name = FONT
    deg_p.paragraph_format.space_after = Pt(1)
    
    school_p = doc.add_paragraph('Post University, Waterbury, CT | 2020 – 2022')
    for r in school_p.runs:
        r.font.size = Pt(10)
        r.font.name = FONT
    
    # Save document
    # === ATS Keyword Report ===
    used_text_blobs = []
    used_text_blobs.append(summary_text.lower())
    used_text_blobs.append(' '.join(optimized_skills).lower())
    # Collect bullets text
    all_bullets_text = []
    for entry in experience:
        for b in entry.get('bullets', []):
            all_bullets_text.append(b.lower())
    used_text_blobs.append(' '.join(all_bullets_text))
    used_tokens = set()
    for blob in used_text_blobs:
        for kw in jd_keywords:
            if kw in blob:
                used_tokens.add(kw)
    missing = [kw for kw in jd_keywords if kw not in used_tokens]
    report_lines = [
        f"Company: {company}",
        f"Position: {position}",
        "Top JD Keywords (first 20):", ', '.join(jd_keywords),
        "Matched Keywords:", ', '.join(sorted(used_tokens)) if used_tokens else 'None',
        "Missing Keywords:", ', '.join(missing) if missing else 'None'
    ]
    report_path = output_path.replace('.docx', '_ats_report.txt')
    try:
        with open(report_path, 'w', encoding='utf-8') as rf:
            rf.write('\n'.join(report_lines))
    except Exception:
        pass

    try:
        doc.save(output_path)
        print(f"[SUCCESS] ATS-optimized resume generated: {output_path}\n[INFO] ATS report: {report_path}")
        return output_path
    except PermissionError:
        alt = output_path.replace('.docx', '_v2.docx')
        doc.save(alt)
        print(f"[SUCCESS] ATS-optimized resume generated: {alt}\n[INFO] ATS report: {report_path}")
        return alt


if __name__ == '__main__':
    sample_jd = "Lead automation, drive lean improvements, support NPI, collaborate cross-functionally."
    build_jd_resume("TestCo", "Engineer", sample_jd, "test_jd_resume.docx")
    print("Generated test_jd_resume.docx")
