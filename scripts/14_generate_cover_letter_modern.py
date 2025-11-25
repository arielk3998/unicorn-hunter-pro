import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(__file__)
CONTACT_FILE = os.path.join(ROOT, 'contact_info.json')
COVER_SRC = os.path.join(ROOT, 'cover_letter_3M_short.txt')
OUT_STYLED = os.path.join(ROOT, 'cover_letter_3M_modern.docx')
OUT_ATS = os.path.join(ROOT, 'cover_letter_3M_modern_ats.docx')
STYLE_SPEC_PATH = os.path.join(ROOT, 'style_spec.json')
ACCENT = RGBColor(0x12, 0x34, 0x56)
FONT_PRIMARY = 'Calibri'
FONT_ATS = 'Arial'

def load_contact():
    try:
        with open(CONTACT_FILE,'r',encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}

def load_style_spec():
    if os.path.exists(STYLE_SPEC_PATH):
        try:
            with open(STYLE_SPEC_PATH,'r',encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def read_lines(path):
    with open(path,'r',encoding='utf-8') as f:
        return [l.rstrip() for l in f]

def add_dans_metadata(doc, contact):
    """Add DANS-compliant document metadata for digital application navigation."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - Cover Letter"
    core_props.subject = "Cover Letter - ATS & DANS Optimized"
    core_props.keywords = "cover letter, ATS, DANS, professional, application"
    core_props.category = "Cover Letter"
    core_props.comments = "Digital Application Navigation System (DANS) compliant with ATS optimization"

def configure_dans_layout(doc, spec):
    """Configure DANS-compliant page layout with standard dimensions and margins."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    m = spec.get('margins_inches', {})
    section.top_margin = Inches(m.get('top', 0.75))  # DANS standard
    section.bottom_margin = Inches(m.get('bottom', 0.75))
    section.left_margin = Inches(m.get('left', 0.75))
    section.right_margin = Inches(m.get('right', 0.75))

def apply_margins(doc, spec):
    """Legacy margin function - now delegates to DANS layout."""
    configure_dans_layout(doc, spec)

def build_header(doc, contact, ats=False, spec=None):
    """Build header with DANS-compliant Heading 1 style."""
    p = doc.add_paragraph()
    p.style = 'Heading 1'  # DANS: Use Heading 1 for name
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(contact.get('name','Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(18 if ats else 20)
    r.font.name = FONT_ATS if ats else FONT_PRIMARY
    if not ats:
        r.font.color.rgb = ACCENT
    else:
        # DANS: Explicit color for accessibility
        r.font.color.rgb = RGBColor(0, 0, 0)
    info_parts = [contact.get('address',''), contact.get('location_statement','')]
    info = ' • '.join([i for i in info_parts if i])
    if info:
        p2 = doc.add_paragraph(info)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in p2.runs:
            rr.font.size = Pt(9)
            rr.font.name = FONT_ATS if ats else FONT_PRIMARY
    line3_parts = [contact.get('email',''), contact.get('phone',''), contact.get('linkedin',''), contact.get('github','')]
    line3 = ' • '.join([p for p in line3_parts if p])
    p3 = doc.add_paragraph(line3)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rr in p3.runs:
        rr.font.size = Pt(9)
        rr.font.name = FONT_ATS if ats else FONT_PRIMARY

def add_section_gap(doc, pt=6):
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(pt)

def build_cover(contact, lines, ats=False, spec=None):
    doc = Document()
    add_dans_metadata(doc, contact)  # DANS compliance
    apply_margins(doc, spec or {})
    build_header(doc, contact, ats=ats, spec=spec)
    add_section_gap(doc, pt=spec.get('spacing',{}).get('section_after_pt',6) if spec else 6)
    body_font = FONT_ATS if ats else FONT_PRIMARY
    for raw in lines:
        if not raw.strip():
            continue
        p = doc.add_paragraph(raw)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = body_font
        p.paragraph_format.space_after = Pt(spec.get('spacing',{}).get('bullet_after_pt',2) if spec else 2)
    out_path = OUT_ATS if ats else OUT_STYLED
    try:
        doc.save(out_path)
    except PermissionError:
        alt = out_path.replace('.docx','_updated.docx')
        doc.save(alt)
        out_path = alt
    return out_path

def main():
    contact = load_contact()
    spec = load_style_spec()
    lines = read_lines(COVER_SRC)
    styled = build_cover(contact, lines, ats=False, spec=spec)
    ats = build_cover(contact, lines, ats=True, spec=spec)
    print('Modern cover letters generated:', styled, ats)

if __name__ == '__main__':
    main()
