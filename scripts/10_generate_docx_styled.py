import os, json, re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(__file__)
CONTACT_FILE = os.path.join(ROOT, 'contact_info.json')
RESUME_TXT = os.path.join(ROOT, 'tailored_resume_3M.txt')
COVER_LETTER_TXT = os.path.join(ROOT, 'cover_letter_3M.txt')
RESUME_DOCX = os.path.join(ROOT, 'tailored_resume_3M.docx')
COVER_LETTER_DOCX = os.path.join(ROOT, 'cover_letter_3M.docx')
EXPERIENCE_JSON = os.path.join(ROOT, 'experience_entries.json')

SECTION_PATTERN = re.compile(r'^([A-Z][A-Z &/]+)$')

ACCENT_COLOR = RGBColor(0x18, 0x3C, 0x5A)  # deep blue
BODY_FONT = 'Calibri'

def load_contact():
    if os.path.exists(CONTACT_FILE):
        with open(CONTACT_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def read_lines(path):
    with open(path, 'r', encoding='utf-8') as f:
        return [l.rstrip() for l in f.readlines()]

def is_bullet(line: str) -> bool:
    return line.strip().startswith(('•','*','-'))

def clean_bullet(text: str) -> str:
    return text.lstrip('•*- ').strip()

def split_core_skills(skills_line: str):
    # Split by bullet separators • or comma
    parts = re.split(r'\s*•\s*|,\s*', skills_line)
    cleaned = [p.strip() for p in parts if p.strip()]
    return cleaned

def add_dans_metadata(doc, contact):
    """Add DANS-compliant document metadata."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - Professional Resume"
    core_props.subject = "Resume - ATS & DANS Optimized"
    core_props.keywords = "resume, ATS, DANS, professional"
    core_props.category = "Resume"
    core_props.comments = "Digital Application Navigation System (DANS) compliant"

def configure_dans_layout(doc):
    """Configure DANS-compliant page layout."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def build_resume(contact, lines):
    doc = Document()
    add_dans_metadata(doc, contact)
    configure_dans_layout(doc)
    # Header - DANS: Use Heading 1
    header = doc.add_paragraph()
    header.style = 'Heading 1'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(contact.get('name','Ariel Karagodskiy'))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = BODY_FONT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"{contact.get('address','')} • {contact.get('location_statement','')}".strip())
    sub_run.font.size = Pt(10)
    sub_run.font.name = BODY_FONT
    contact_line_parts = [contact.get('email',''), contact.get('phone','')]
    if contact.get('linkedin'): contact_line_parts.append(contact['linkedin'])
    if contact.get('github'): contact_line_parts.append(contact['github'])
    contact_line = ' • '.join([p for p in contact_line_parts if p])
    contact_para = doc.add_paragraph(contact_line)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in contact_para.runs:
        r.font.size = Pt(9)
        r.font.name = BODY_FONT

    current_section = None
    buffer = []

    # Load structured experience entries if available
    experience_entries = []
    if os.path.exists(EXPERIENCE_JSON):
        try:
            with open(EXPERIENCE_JSON, 'r', encoding='utf-8') as f:
                experience_entries = json.load(f)
        except Exception:
            experience_entries = []

    def write_experience(entries):
        if not entries:
            return False
        title = doc.add_paragraph()
        run = title.add_run('Professional Experience')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = ACCENT_COLOR
        run.font.name = BODY_FONT
        for e in entries:
            header_p = doc.add_paragraph()
            header_run = header_p.add_run(f"{e.get('title','')} | {e.get('company','')} | {e.get('dates','')}")
            header_run.bold = True
            header_run.font.size = Pt(10)
            header_run.font.name = BODY_FONT
            for b in e.get('bullets', []):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(clean_bullet(b))
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = BODY_FONT
        return True

    def flush_section(sec, buf):
        if not sec or not buf:
            return
        title = doc.add_paragraph()
        run = title.add_run(sec)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = ACCENT_COLOR
        run.font.name = BODY_FONT

        # Detect CORE SKILLS special handling: create two-column table
        if sec.lower().startswith('core skills'):
            # Merge all buffer lines into one skill string (excluding emerging areas line)
            skill_text_lines = [l for l in buf if not l.lower().startswith('emerging/')]
            skill_flat = ' '.join(skill_text_lines)
            skills = split_core_skills(skill_flat)
            if skills:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = True
                left_cell, right_cell = table.rows[0].cells
                mid = (len(skills) + 1) // 2
                left_list = skills[:mid]
                right_list = skills[mid:]
                for lst, cell in ((left_list, left_cell), (right_list, right_cell)):
                    for s in lst:
                        p = cell.add_paragraph(style='List Bullet')
                        r = p.add_run(s)
                        r.font.name = BODY_FONT
                        r.font.size = Pt(10)
                # Add emerging/development line after table
                emerging_lines = [l for l in buf if l.lower().startswith('emerging/')]
                for el in emerging_lines:
                    p = doc.add_paragraph(el)
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = BODY_FONT
                return

        for b in buf:
            if is_bullet(b):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(clean_bullet(b))
            else:
                p = doc.add_paragraph(b)
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = BODY_FONT

    skip_prof_exp = False
    for idx, line in enumerate(lines):
        if not line.strip():
            continue
        if line.strip().upper().startswith('NOTES FOR 3M'):
            break
        # Detect section headings
        if line.isupper() and len(line.split()) < 10 and not is_bullet(line) and not line.startswith('Ariel'):
            # When encountering PROFESSIONAL EXPERIENCE in text, replace with structured entries
            if line.strip().upper().startswith('PROFESSIONAL EXPERIENCE') and experience_entries:
                flush_section(current_section, buffer)
                write_experience(experience_entries)
                current_section = None
                buffer = []
                skip_prof_exp = True
                continue
            else:
                flush_section(current_section, buffer)
                current_section = line.title()
                buffer = []
        else:
            if skip_prof_exp:
                # Skip original unstructured experience lines until next section heading
                if line.isupper() and len(line.split()) < 10 and not is_bullet(line):
                    skip_prof_exp = False
                    current_section = line.title()
                else:
                    continue
            if not skip_prof_exp:
                buffer.append(line)
    flush_section(current_section, buffer)

    doc.save(RESUME_DOCX)
    return RESUME_DOCX

def build_cover_letter(contact, lines):
    doc = Document()
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(contact.get('name','Ariel Karagodskiy'))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = BODY_FONT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"{contact.get('address','')} • {contact.get('email','')} • {contact.get('phone','')}")
    sub_run.font.size = Pt(9)
    sub_run.font.name = BODY_FONT
    if contact.get('linkedin') or contact.get('github'):
        social = doc.add_paragraph(' '.join([contact.get('linkedin',''), contact.get('github','')]).strip())
        social.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in social.runs:
            r.font.size = Pt(9)
            r.font.name = BODY_FONT

    today = date.today().strftime('%B %d, %Y')
    date_p = doc.add_paragraph(today)
    for r in date_p.runs:
        r.font.size = Pt(10)
        r.font.name = BODY_FONT

    # Body (remove header lines duplicated)
    for line in lines:
        if not line.strip():
            continue
        if line.startswith(contact.get('name','Ariel')):
            continue
        if 'Ariel.pk@outlook.com' in line or contact.get('email','') in line:
            continue
        if line.startswith('Re:'):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.color.rgb = ACCENT_COLOR
            r.font.name = BODY_FONT
            continue
        if is_bullet(line):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(clean_bullet(line))
        else:
            p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = BODY_FONT

    # Safe save (handle file locked by Word)
    target = COVER_LETTER_DOCX
    try:
        doc.save(target)
    except PermissionError:
        alt = target.replace('.docx','_updated.docx')
        try:
            doc.save(alt)
            target = alt
        except Exception:
            raise
    return target

def main():
    contact = load_contact()
    resume_lines = read_lines(RESUME_TXT)
    cover_lines = read_lines(COVER_LETTER_TXT)
    r_path = build_resume(contact, resume_lines)
    c_path = build_cover_letter(contact, cover_lines)
    print('Generated styled DOCX:', r_path, 'and', c_path)

if __name__ == '__main__':
    main()