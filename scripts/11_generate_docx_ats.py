import os, json
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(__file__)
CONTACT_FILE = os.path.join(ROOT, 'contact_info.json')
RESUME_TXT = os.path.join(ROOT, 'tailored_resume_3M.txt')
COVER_TXT = os.path.join(ROOT, 'cover_letter_3M.txt')
RESUME_DOCX = os.path.join(ROOT, 'tailored_resume_3M.docx')  # overwrite
COVER_DOCX = os.path.join(ROOT, 'cover_letter_3M.docx')      # overwrite
EXPERIENCE_JSON = os.path.join(ROOT, 'experience_entries.json')

BODY_FONT = 'Arial'

def load_contact():
    if os.path.exists(CONTACT_FILE):
        with open(CONTACT_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def read_lines(path):
    with open(path, 'r', encoding='utf-8') as f:
        return [l.rstrip() for l in f.readlines()]

def is_heading(line: str) -> bool:
    # Basic heading heuristic: uppercase line with <= 8 words
    words = line.split()
    return line.isupper() and 0 < len(words) <= 8

def normalize_line(line: str) -> str:
    return line.replace('•', '-').replace('–', '-').strip()

def add_header(doc: Document, contact: dict, title_size=18):
    """Add DANS-compliant header with structured contact information."""
    # DANS: Name as document title (Heading 1 for structure)
    p = doc.add_paragraph()
    p.style = 'Heading 1'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(contact.get('name', 'Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(title_size)
    r.font.name = BODY_FONT
    r.font.color.rgb = RGBColor(0, 0, 0)  # DANS: explicit color

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    loc = contact.get('address', '')
    stmt = contact.get('location_statement', '')
    r2 = p2.add_run(' | '.join([x for x in [loc, stmt] if x]))
    r2.font.size = Pt(9)
    r2.font.name = BODY_FONT

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    items = [contact.get('email',''), contact.get('phone','')]
    if contact.get('linkedin'): items.append(contact['linkedin'])
    if contact.get('github'): items.append(contact['github'])
    r3 = p3.add_run(' | '.join([i for i in items if i]))
    r3.font.size = Pt(9)
    r3.font.name = BODY_FONT

def write_section_heading(doc: Document, text: str):
    """Write DANS-compliant section heading with consistent styling and structure."""
    p = doc.add_paragraph()
    # DANS: Use heading style for proper document structure
    p.style = 'Heading 2'
    r = p.add_run(text.upper())  # DANS: uppercase for clear section delineation
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = BODY_FONT
    r.font.color.rgb = RGBColor(0, 0, 0)  # DANS: explicit black for accessibility
    # Add spacing for DANS readability
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

def add_dans_metadata(doc: Document, contact: dict):
    """Add DANS-compliant document metadata for digital application systems."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - Professional Resume"
    core_props.subject = "Resume - ATS & DANS Optimized"
    core_props.keywords = "resume, ATS, DANS, professional, application"
    core_props.category = "Resume"
    core_props.comments = "Digital Application Navigation System (DANS) compliant resume with ATS optimization"

def configure_dans_layout(doc: Document):
    """Configure DANS-compliant page layout: standard margins, consistent spacing."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)

def build_resume(contact, lines):
    doc = Document()
    # DANS compliance: metadata and layout
    add_dans_metadata(doc, contact)
    configure_dans_layout(doc)
    add_header(doc, contact, title_size=18)
    current_heading = None
    buffer = []

    experience_entries = []
    if os.path.exists(EXPERIENCE_JSON):
        try:
            with open(EXPERIENCE_JSON, 'r', encoding='utf-8') as f:
                experience_entries = json.load(f)
        except Exception:
            experience_entries = []

    def write_experience(entries):
        if not entries:
            return False
        write_section_heading(doc, 'Professional Experience')
        for e in entries:
            header = doc.add_paragraph()
            run = header.add_run(f"{e.get('title','')} | {e.get('company','')} | {e.get('dates','')}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = BODY_FONT
            for b in e.get('bullets', []):
                p = doc.add_paragraph(normalize_line(b))
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = BODY_FONT
        return True

    def flush():
        nonlocal buffer, current_heading
        if current_heading and buffer:
            write_section_heading(doc, current_heading)
            for raw in buffer:
                if not raw.strip():
                    continue
                line = normalize_line(raw)
                if line.startswith('- '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:].strip())
                elif line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:].strip())
                elif line.startswith('• '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:].strip())
                else:
                    p = doc.add_paragraph(line)
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = BODY_FONT
            buffer = []

    skip_prof_exp = False
    for line in lines:
        if line.strip().upper().startswith('ADDITIONAL INFO'):
            # treat as normal heading
            pass
        if is_heading(line) and not line.startswith('Ariel'):
            # Inject structured experience when hitting heading
            if line.strip().upper().startswith('PROFESSIONAL EXPERIENCE') and experience_entries:
                flush()
                write_experience(experience_entries)
                current_heading = None
                buffer = []
                skip_prof_exp = True
                continue
            else:
                flush()
                current_heading = line.strip()
        else:
            # Skip internal planning or emerging lines not needed for ATS if empty
            if line.strip().lower().startswith('growth & extension targets'): continue
            if skip_prof_exp:
                if is_heading(line) and not line.startswith('Ariel'):
                    skip_prof_exp = False
                    current_heading = line.strip()
                else:
                    continue
            if not skip_prof_exp:
                buffer.append(line)
    flush()
    doc.save(RESUME_DOCX)

def build_cover_letter(contact, lines):
    doc = Document()
    add_header(doc, contact, title_size=16)
    # Date
    d = doc.add_paragraph(date.today().strftime('%B %d, %Y'))
    for r in d.runs:
        r.font.size = Pt(10)
        r.font.name = BODY_FONT
    for line in lines:
        if not line.strip():
            continue
        if line.startswith(contact.get('name','Ariel')):
            continue
        if contact.get('email','') in line or contact.get('phone','') in line:
            continue
        txt = normalize_line(line)
        p = doc.add_paragraph(txt)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = BODY_FONT
    target = COVER_DOCX
    try:
        doc.save(target)
    except PermissionError:
        alt = target.replace('.docx','_updated.docx')
        try:
            doc.save(alt)
            target = alt
        except Exception:
            raise

def main():
    contact = load_contact()
    resume_lines = read_lines(RESUME_TXT)
    cover_lines = read_lines(COVER_TXT)
    build_resume(contact, resume_lines)
    build_cover_letter(contact, cover_lines)
    print('ATS-friendly DOCX files written.')

if __name__ == '__main__':
    main()