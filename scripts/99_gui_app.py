"""
GUI Application for Job Application Automation
Provides user-friendly interface for applying to jobs with auto-generated resumes and cover letters.
"""
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from pathlib import Path
import sys
import threading
from datetime import datetime
import webbrowser

# Add parent directory to path for imports
ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT / 'scripts'))

# Import the automation module
import importlib.util
import re
spec = importlib.util.spec_from_file_location("apply_to_job", ROOT / 'scripts' / '00_apply_to_job.py')
apply_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(apply_module)
JobApplicationAutomation = apply_module.JobApplicationAutomation

# Cover letter generator (extracted logic)
try:
    from cover_letter_generator import generate_cover_letter
except Exception:
    generate_cover_letter = None

# Import AI assistant
spec_ai = importlib.util.spec_from_file_location("ai_assistant", ROOT / 'scripts' / 'ai_assistant.py')
ai_module = importlib.util.module_from_spec(spec_ai)
spec_ai.loader.exec_module(ai_module)
get_ai_assistant = ai_module.get_ai_assistant

# ATS analyzer (extracted)
try:
    from ats_analyzer import analyze as ats_analyze
except Exception:
    ats_analyze = None

# Import Framework Validator
spec_fv = importlib.util.spec_from_file_location("framework_validator", ROOT / 'scripts' / 'framework_validator.py')
fv_module = importlib.util.module_from_spec(spec_fv)
spec_fv.loader.exec_module(fv_module)
FrameworkValidator = fv_module.FrameworkValidator

# Import Session Manager (for session-level data capture & finalization)
spec_sm = importlib.util.spec_from_file_location("session_manager", ROOT / 'scripts' / 'session_manager.py')
sm_module = importlib.util.module_from_spec(spec_sm)
sys.modules[spec_sm.name] = sm_module
spec_sm.loader.exec_module(sm_module)
SessionManager = sm_module.SessionManager
spec_ca = importlib.util.spec_from_file_location("compensation_advisor", ROOT / 'scripts' / 'compensation_advisor.py')
ca_module = importlib.util.module_from_spec(spec_ca)
sys.modules[spec_ca.name] = ca_module
spec_ca.loader.exec_module(ca_module)
CompensationAdvisor = ca_module.CompensationAdvisor
# Import Offline AI Core for bullet optimization & tailoring
spec_ai_core = importlib.util.spec_from_file_location("offline_ai_core", ROOT / 'scripts' / 'offline_ai_core.py')
ai_core_module = importlib.util.module_from_spec(spec_ai_core)
sys.modules[spec_ai_core.name] = ai_core_module
spec_ai_core.loader.exec_module(ai_core_module)
OfflineAICore = ai_core_module.OfflineAICore

# Accessibility Manager import
spec_acc = importlib.util.spec_from_file_location("accessibility_manager", ROOT / 'scripts' / 'accessibility_manager.py')
acc_module = importlib.util.module_from_spec(spec_acc)
sys.modules[spec_acc.name] = acc_module
spec_acc.loader.exec_module(acc_module)
AccessibilityManager = acc_module.AccessibilityManager
ACCESSIBLE_PALETTES = acc_module.ACCESSIBLE_PALETTES

# Template Engine import (registered to sys.modules to satisfy dataclasses reliably)
spec_te = importlib.util.spec_from_file_location("template_engine", ROOT / 'scripts' / 'template_engine.py')
te_module = importlib.util.module_from_spec(spec_te)
sys.modules[spec_te.name] = te_module
spec_te.loader.exec_module(te_module)
TemplateEngine = te_module.TemplateEngine

# Resume Version Manager import
spec_rvm = importlib.util.spec_from_file_location("resume_version_manager", ROOT / 'scripts' / 'resume_version_manager.py')
rvm_module = importlib.util.module_from_spec(spec_rvm)
sys.modules[spec_rvm.name] = rvm_module
spec_rvm.loader.exec_module(rvm_module)
ResumeVersionManager = rvm_module.ResumeVersionManager

# Skills dashboard import
skills_spec = importlib.util.spec_from_file_location("skills_dashboard", ROOT / 'scripts' / 'skills_dashboard.py')
skills_module = importlib.util.module_from_spec(skills_spec)
sys.modules[skills_spec.name] = skills_module
skills_spec.loader.exec_module(skills_module)
extract_skills = skills_module.extract_skills
top_skills = skills_module.top_skills


class JobApplicationGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Resume Toolkit - Job Application Automation System")
        self.root.geometry("1000x700")  # Wider default to better fit table
        self.root.minsize(900, 600)
        
        # Initialize Accessibility Manager FIRST for WCAG 2.1 AA compliance
        self.accessibility = AccessibilityManager(self.root)
        
        # Icon placeholder (temporary modern glyph)
        try:
            icon_data = (
                "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAMAAAAoLQ9TAAAAM1BMVEUAAAD///////////////////////////////////////////////8AAABn4nQeAAAADHRSTlMAEBAgIDg4QEBgcHCAcRBcAAAAP0lEQVQY02WPRw6AIAxEB0lZ0P7/39tJgWJpQxzGJNsH7gCwH2g0gijvLwDE2QhHSAcJxwCkIVy4i3zjqEHZIrYxbCkDBkJ5nKyeL6mv1Z3yB6Q9w+LFU9RP9YQFsxh3ZKAAAAAElFTkSuQmCC"
            )
            self._icon_img = tk.PhotoImage(data=icon_data)
            self.root.iconphoto(True, self._icon_img)
        except Exception:
            pass
        
        # Configure style
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Use WCAG 2.1 AA compliant color palettes
        self._palettes = ACCESSIBLE_PALETTES
        
        self.current_theme = 'light'
        self._apply_palette()
        self.root.configure(bg=self.bg_color)

        # Global font definitions - modern and clean
        self.font_regular = ("Segoe UI", 10)
        self.font_bold = ("Segoe UI", 10, "bold")
        self.font_heading = ("Segoe UI", 22, "bold")
        self.font_subheading = ("Segoe UI", 14, "bold")
        self.font_small = ("Segoe UI", 9)

        # Enhanced ttk styles - modern and sleek
        self.style.configure('TFrame', background=self.bg_color)
        self.style.configure('Card.TFrame', background=self.subtle_bg, relief='flat')
        self.style.configure('TNotebook', background=self.bg_color, borderwidth=0)
        self.style.configure('TNotebook.Tab', 
                           font=self.font_bold, 
                           padding=(20, 12),
                           borderwidth=0)
        self.style.map('TNotebook.Tab', 
                      background=[('selected', self.subtle_bg), ('!selected', self.bg_color)],
                      foreground=[('selected', self.accent_color), ('!selected', self.text_muted)])
        self.style.configure('TLabel', background=self.bg_color, foreground=self.text_color, font=self.font_regular)
        self.style.configure('Heading.TLabel', font=self.font_heading, foreground=self.text_color)
        self.style.configure('Subheading.TLabel', font=self.font_subheading, foreground=self.text_color)
        self.style.configure('Muted.TLabel', foreground=self.text_muted, font=self.font_small)
        self.style.configure('TButton', 
                           font=self.font_regular, 
                           padding=(12, 8),
                           borderwidth=1,
                           relief='flat')
        self.style.map('TButton', 
                      background=[('active', self.border_color), ('!active', self.subtle_bg)],
                      bordercolor=[('focus', self.accent_color)])
        self.style.configure('Accent.TButton', 
                           background=self.accent_color, 
                           foreground='white', 
                           font=self.font_bold, 
                           padding=(16, 10),
                           borderwidth=0,
                           relief='flat')
        self.style.map('Accent.TButton', 
                      background=[('active', self.accent_dark), ('!active', self.accent_color)])
        self.style.configure('Success.TButton', 
                           background=self.success_color, 
                           foreground='white', 
                           font=self.font_bold, 
                           padding=(12, 8),
                           borderwidth=0)
        self.style.map('Success.TButton', background=[('active', '#059669')])
        self.style.configure('Treeview', 
                           font=self.font_regular, 
                           rowheight=32, 
                           bordercolor=self.border_color, 
                           borderwidth=1,
                           background=self.subtle_bg,
                           fieldbackground=self.subtle_bg)
        self.style.configure('Treeview.Heading', 
                           font=self.font_bold, 
                           background=self.bg_color,
                           foreground=self.text_color,
                           borderwidth=1,
                           relief='flat')
        self.style.map('Treeview', 
                      background=[('selected', self.accent_light)], 
                      foreground=[('selected', self.text_color)])
        self.style.map('Treeview.Heading',
                      background=[('active', self.border_color)])
        
        self.automation = JobApplicationAutomation()
        self.ai_assistant = get_ai_assistant()
        self.offline_core = OfflineAICore()  # Offline AI for tailoring & optimization
        self.template_engine = TemplateEngine() if 'TemplateEngine' in globals() else None
        self.version_manager = ResumeVersionManager() if 'ResumeVersionManager' in globals() else None
        # Root directory reference for profile data
        self.root_dir = ROOT
        self.default_output_dir = str(ROOT / 'outputs')
        self.reference_resumes = []  # Store paths to reference resumes
        self.tracker_path = self._find_tracker()  # Auto-locate tracker
        # Feedback config paths & load
        self.feedback_log_path = ROOT / 'data' / 'feedback_log.json'
        self.feedback_config_path = ROOT / 'data' / 'feedback_config.json'
        self.feedback_config = self._load_feedback_config()
        # Simple mode preference (must be set before setup_ui)
        self.simple_mode = False
        self._hidden_tabs = []
        self._preferences_path = ROOT / 'data' / 'config_preferences.json'
        self._load_preferences()
        self.setup_ui()
    
    def _find_tracker(self) -> Path:
        """Automatically locate the job application tracker Excel file.
        Searches in multiple locations and returns the first match found.
        """
        # Search locations in priority order
        search_locations = [
            ROOT / 'job_application_master_tracker.xlsx',  # resume-toolkit folder
            ROOT.parent / 'job_application_master_tracker.xlsx',  # Ariels-Resumes folder
            Path.home() / 'Documents' / 'job_application_master_tracker.xlsx',
            Path.home() / 'Desktop' / 'job_application_master_tracker.xlsx',
        ]
        
        for path in search_locations:
            if path.exists():
                return path
        
        # If not found, default to resume-toolkit location (will be created if needed)
        return ROOT / 'job_application_master_tracker.xlsx'
    
    def setup_ui(self):
        """Create the main UI layout"""
        # Modern header with gradient
        header_height = 80
        header_canvas = tk.Canvas(self.root, height=header_height, highlightthickness=0, bd=0)
        header_canvas.pack(fill=tk.X)
        # Smooth gradient from dark to light
        for i in range(header_height):
            ratio = i / header_height
            # Interpolate between accent_dark and accent_color
            def _interp(c1, c2):
                c1 = (int(c1[1:3],16), int(c1[3:5],16), int(c1[5:7],16))
                c2 = (int(c2[1:3],16), int(c2[3:5],16), int(c2[5:7],16))
                ch = tuple(int(c1[j] + (c2[j]-c1[j])*ratio) for j in range(3))
                return f"#{ch[0]:02x}{ch[1]:02x}{ch[2]:02x}"
            color_line = _interp(self.accent_dark, self.accent_color)
            header_canvas.create_line(0, i, 2000, i, fill=color_line)
        
        # App icon and title
        header_canvas.create_text(
            24,
            header_height//2,
            anchor='w',
            text="🚀",
            font=("Segoe UI", 32),
            fill='white'
        )
        header_canvas.create_text(
            75,
            header_height//2 - 8,
            anchor='w',
            text="Job Application Automation",
            font=self.font_heading,
            fill='white'
        )
        header_canvas.create_text(
            75,
            header_height//2 + 16,
            anchor='w',
            text="Smart resume generation powered by AI",
            font=("Segoe UI", 10),
            fill='#bfdbfe'
        )
        # Version badge
        badge_x, badge_y = 920, header_height//2
        badge_width, badge_height = 140, 28
        header_canvas.create_rectangle(
            badge_x, badge_y - badge_height//2,
            badge_x + badge_width, badge_y + badge_height//2,
            fill='#1e40af', outline='', width=0
        )
        header_canvas.create_text(
            badge_x + badge_width//2,
            badge_y,
            text=datetime.now().strftime("v1.7.0 • %b %d, %Y"),
            font=("Segoe UI", 9, "bold"),
            fill='white'
        )
        # Simple Mode toggle button (overlay right of header)
        self.simple_btn = tk.Button(header_canvas, text='Simple Mode: OFF', font=("Segoe UI",9,'bold'), bg='#1e3a8a', fg='white', relief='flat', command=self.toggle_simple_mode, cursor='hand2')
        header_canvas.create_window(badge_x - 160, badge_y, window=self.simple_btn)
        # Theme cycle button
        self.theme_btn = tk.Button(header_canvas, text='Theme: Light', font=("Segoe UI",9,'bold'), bg='#1e3a8a', fg='white', relief='flat', command=self.cycle_theme, cursor='hand2')
        header_canvas.create_window(badge_x - 300, badge_y, window=self.theme_btn)
        
        # Main content area with notebook (tabs)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Tab 1: Apply to Job
        self.apply_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.apply_tab, text="Apply to Job")
        self.create_apply_tab()
        
        # Tab 2: View Applications
        self.view_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.view_tab, text="View Applications")
        self.create_view_tab()
        
        # Tab 3: Framework Validator
        self.validator_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.validator_tab, text="Framework Validator")
        self.create_validator_tab()
        
        # Tab 4: Profile & Settings (combined)
        self.profile_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.profile_tab, text="Profile & Settings")
        self.create_profile_tab()

        # Tab 5: Compensation Advisor
        self.comp_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.comp_tab, text="Compensation Advisor")
        self.create_comp_tab()

        # Tab 6: Optimize & Tailor (Offline AI)
        self.tailor_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.tailor_tab, text="Optimize & Tailor")
        self.create_tailor_tab()

        # Tab 7: Cover Letter Generator
        self.cover_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.cover_tab, text="Cover Letter")
        self.create_cover_letter_tab()

        # Tab 8: ATS Scanner (Beta stub)
        self.ats_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.ats_tab, text="ATS Scanner (Beta)")
        self.create_ats_tab()

        # Tab 9: Templates
        self.templates_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.templates_tab, text="Templates")
        self.create_templates_tab()

        # Tab 10: Versions
        self.versions_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.versions_tab, text="Versions")
        self.create_versions_tab()

        # Tab 11: Skills Dashboard
        self.skills_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.skills_tab, text="Skills")
        self.create_skills_tab()
        # Tab 12: Distribution (Stub)
        self.distribution_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.distribution_tab, text='Distribution')
        self.create_distribution_tab()
        
        # Status bar (enhanced: includes session tracking + finalize)
        self.status_bar_frame = tk.Frame(self.root, bg="#e0e0e0")
        self.status_bar_frame.pack(side=tk.BOTTOM, fill=tk.X)
        self.status_bar = tk.Label(self.status_bar_frame, text="Ready", anchor=tk.W, bg="#e0e0e0")
        self.status_bar.pack(side=tk.LEFT, padx=6)
        self.session_info_label = tk.Label(self.status_bar_frame, text="Session active", anchor=tk.W, bg="#e0e0e0")
        self.session_info_label.pack(side=tk.LEFT, padx=12)
        self.finalize_btn = tk.Button(self.status_bar_frame, text="Finalize Session", command=self._finalize_session, bg="#0284c7", fg="white")
        self.finalize_btn.pack(side=tk.RIGHT, padx=6, pady=2)
        
        # Load reference resumes on startup
        self._load_reference_resumes()
        # Schedule initial width adjustment after layout settles
        self.root.after(350, self._initial_resize)
        # Start periodic adaptation checks
        self.root.after(2500, self._maybe_adapt_thresholds)
        # Initialize session manager & wrap job fetch
        self.session_manager = SessionManager(ROOT)
        self._wrap_fetch_job()
        
        # Enable keyboard navigation for WCAG 2.1 AA compliance (ADA/EAA)
        self.accessibility.enable_keyboard_navigation()
        # On window close, finalize session
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        # Apply simple mode if preference saved
        if self.simple_mode:
            self._apply_simple_mode()

    def _initial_resize(self):
        """Ensure window width accommodates the applications table."""
        try:
            if hasattr(self, 'app_tree'):
                self.root.update_idletasks()
                tree_width = self.app_tree.winfo_reqwidth()
                padding = 140  # account for notebook margins
                target_width = max(self.root.winfo_width(), tree_width + padding)
                if target_width > self.root.winfo_width():
                    self.root.geometry(f"{target_width}x{self.root.winfo_height()}")
                    # Lock in a wider minsize so future refreshes don't shrink below content
                    self.root.minsize(target_width, 600)
        except Exception:
            pass

    def _wrap_fetch_job(self):
        """Monkey-patch AI assistant fetch to auto-record successful job descriptions."""
        if not hasattr(self, 'ai_assistant'):
            return
        original = self.ai_assistant.fetch_job_from_url

        def wrapped(url: str):
            result = original(url)
            try:
                if result.get('success') and result.get('job_description'):
                    self.session_manager.add_job(result)
                    self.session_info_label.config(text=f"Session jobs: {len(self.session_manager.jobs)}")
            except Exception:
                pass
            return result

        self.ai_assistant.fetch_job_from_url = wrapped

    def _finalize_session(self):
        """Finalize session manually (writes folder with collected data)."""
        path = self.session_manager.finalize_session()
        if path:
            self.status_bar.config(text=f"Session finalized: {path.name}")
            self.finalize_btn.config(state='disabled')
        else:
            self.status_bar.config(text="Session already finalized or no jobs collected")

    def _on_close(self):
        """Finalize session then close GUI."""
        try:
            self.session_manager.finalize_session()
        except Exception:
            pass
        self.root.destroy()

    # === COMPENSATION ADVISOR TAB ===
    def create_comp_tab(self):
        advisor = CompensationAdvisor()
        container = ttk.Frame(self.comp_tab, padding="14")
        container.pack(fill=tk.BOTH, expand=True)

        header = ttk.Label(container, text="Compensation & Salary Guidance", style='Subheading.TLabel')
        header.pack(anchor=tk.W)
        ttk.Label(container, text="Estimate realistic salary ranges by role & zipcode; adjust for experience & skills.", style='Muted.TLabel').pack(anchor=tk.W, pady=(2,12))

        form = ttk.Frame(container)
        form.pack(fill=tk.X)

        ttk.Label(form, text="Job Title:").grid(row=0, column=0, sticky='w', padx=(0,6), pady=4)
        self.comp_title_var = tk.StringVar()
        ttk.Entry(form, textvariable=self.comp_title_var, width=32).grid(row=0, column=1, sticky='w')

        ttk.Label(form, text="Zipcode:").grid(row=0, column=2, sticky='w', padx=(18,6))
        self.comp_zip_var = tk.StringVar()
        ttk.Entry(form, textvariable=self.comp_zip_var, width=10).grid(row=0, column=3, sticky='w')

        ttk.Label(form, text="Years Experience:").grid(row=1, column=0, sticky='w', padx=(0,6), pady=4)
        self.comp_years_var = tk.StringVar()
        ttk.Entry(form, textvariable=self.comp_years_var, width=10).grid(row=1, column=1, sticky='w')

        ttk.Label(form, text="Skills Count:").grid(row=1, column=2, sticky='w', padx=(18,6))
        self.comp_skills_var = tk.StringVar()
        ttk.Entry(form, textvariable=self.comp_skills_var, width=10).grid(row=1, column=3, sticky='w')

        self.comp_result_text = scrolledtext.ScrolledText(container, height=14, wrap=tk.WORD, font=self.font_regular)
        self.comp_result_text.pack(fill=tk.BOTH, expand=True, pady=(12,8))

        def run_estimate():
            title = self.comp_title_var.get().strip()
            zipcode = self.comp_zip_var.get().strip()
            years = self.comp_years_var.get().strip()
            skills = self.comp_skills_var.get().strip()
            years_i = int(years) if years.isdigit() else None
            skills_i = int(skills) if skills.isdigit() else None
            est = advisor.estimate_salary(title, zipcode, years_i, skills_i)
            self.comp_result_text.delete('1.0', tk.END)
            if not est:
                self.comp_result_text.insert(tk.END, "No local data match. Try another title or zipcode.\n")
                return
            self.comp_result_text.insert(tk.END, f"ESTIMATED RANGE (Adjusted): ${est['adjusted_p25']:,.0f} - ${est['adjusted_p75']:,.0f}\n")
            self.comp_result_text.insert(tk.END, f"Mean: ${est['adjusted_mean']:,.0f} | National Equivalent: ${est['national_equivalent_mean']:,.0f}\n")
            self.comp_result_text.insert(tk.END, f"Data Points: {est['data_points']} | Cost-of-Living Index: {est['cost_of_living_index']}\n")
            self.comp_result_text.insert(tk.END, f"Experience Adj: {est['experience_adjustment_pct']}% | Skills Adj: {est['skills_adjustment_pct']}%\n\n")
            tips = advisor.negotiation_tips(est)
            self.comp_result_text.insert(tk.END, "NEGOTIATION TIPS:\n")
            for t in tips:
                self.comp_result_text.insert(tk.END, f"• {t}\n")
            self.comp_result_text.insert(tk.END, "\nUse these figures to anchor offers; refine with market research.")

        ttk.Button(form, text="Estimate", style='Accent.TButton', command=run_estimate).grid(row=0, column=4, rowspan=2, sticky='nsw', padx=(24,0))

    # === OPTIMIZE & TAILOR TAB ===
    def create_tailor_tab(self):
        container = ttk.Frame(self.tailor_tab, padding="14")
        container.pack(fill=tk.BOTH, expand=True)

        header = ttk.Label(container, text="Bullet Optimization & Tailoring", style='Subheading.TLabel')
        header.pack(anchor=tk.W)
        ttk.Label(container, text="Enhance resume bullets locally: scoring, tailoring to a job description, example retrieval.", style='Muted.TLabel').pack(anchor=tk.W, pady=(2,12))

        # Layout frames
        left = ttk.Frame(container)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0,8))
        right = ttk.Frame(container)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(8,0))

        # Job Description input
        jd_card = ttk.LabelFrame(left, text="Target Job Description", padding="8")
        jd_card.pack(fill=tk.BOTH, expand=True)
        self.tailor_jd_text = scrolledtext.ScrolledText(jd_card, height=12, wrap=tk.WORD, font=("Consolas",9))
        self.tailor_jd_text.pack(fill=tk.BOTH, expand=True)

        # Bullets input
        bullets_card = ttk.LabelFrame(left, text="Resume Bullets (one per line)", padding="8")
        bullets_card.pack(fill=tk.BOTH, expand=True, pady=(12,0))
        self.tailor_bullets_text = scrolledtext.ScrolledText(bullets_card, height=10, wrap=tk.WORD, font=("Consolas",9))
        self.tailor_bullets_text.pack(fill=tk.BOTH, expand=True)

        action_frame = ttk.Frame(left)
        action_frame.pack(fill=tk.X, pady=(10,0))
        ttk.Button(action_frame, text="🔍 Enhance Selected", style='Accent.TButton', command=self._enhance_single_bullet).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(action_frame, text="⚙️ Enhance All", command=self._enhance_all_bullets).pack(side=tk.LEFT)
        ttk.Button(action_frame, text="Clear", command=lambda: self.tailor_bullets_text.delete('1.0', tk.END)).pack(side=tk.LEFT, padx=(8,0))

        # Results area
        results_card = ttk.LabelFrame(right, text="Results", padding="8")
        results_card.pack(fill=tk.BOTH, expand=True)

        self.tailor_result_text = scrolledtext.ScrolledText(results_card, height=28, wrap=tk.WORD, font=("Consolas",9), state='normal')
        self.tailor_result_text.pack(fill=tk.BOTH, expand=True)

        # Helper labels
        ttk.Label(right, text="Shows tailored version, component scores, suggestions, similar examples & extracted keywords.", style='Muted.TLabel').pack(anchor=tk.W, pady=(8,0))
        # Real-Time ATS coverage widgets
        ats_frame = ttk.Frame(right)
        ats_frame.pack(fill=tk.X, pady=(8,4))
        ttk.Label(ats_frame, text='Keyword Coverage:', style='Muted.TLabel').pack(side=tk.LEFT)
        self.ats_progress = ttk.Progressbar(ats_frame, length=180, maximum=100, value=0)
        self.ats_progress.pack(side=tk.LEFT, padx=(6,4))
        self.ats_cov_label = ttk.Label(ats_frame, text='0%', style='Muted.TLabel')
        self.ats_cov_label.pack(side=tk.LEFT)
        self.ats_warn_label = ttk.Label(ats_frame, text='', style='Muted.TLabel', foreground=self.warning_color)
        self.ats_warn_label.pack(side=tk.LEFT, padx=(10,0))
        # Bind updates
        self.tailor_jd_text.bind('<KeyRelease>', self._update_ats_score)
        self.tailor_bullets_text.bind('<KeyRelease>', self._update_ats_score)
        self.root.after(1000, self._update_ats_score)
        # Micro-Lessons panel
        lessons_frame = ttk.LabelFrame(right, text='Micro-Lessons', padding=6)
        lessons_frame.pack(fill=tk.X, pady=(4,0))
        self.micro_lessons_text = tk.Text(lessons_frame, height=6, font=('Consolas',8), wrap=tk.WORD)
        self.micro_lessons_text.pack(fill=tk.BOTH, expand=True)
        self.micro_lessons_text.insert('1.0', 'Tips will appear here when coverage or metrics are low.')
        self.micro_lessons_text.config(state='disabled')

    def _enhance_single_bullet(self):
        bullet = self._get_selected_bullet()
        if not bullet:
            messagebox.showinfo("No Bullet Selected", "Place cursor on a line containing a bullet.")
            return
        jd = self.tailor_jd_text.get('1.0', tk.END)
        result = self.offline_core.enhance(bullet, jd)
        self._render_tailor_result([result])

    def _enhance_all_bullets(self):
        text = self.tailor_bullets_text.get('1.0', tk.END).strip()
        if not text:
            messagebox.showwarning("No Bullets", "Enter one or more bullets (one per line).")
            return
        jd = self.tailor_jd_text.get('1.0', tk.END)
        bullets = [b.strip() for b in text.split('\n') if b.strip()]
        results = [self.offline_core.enhance(b, jd) for b in bullets]
        self._render_tailor_result(results)

    def _get_selected_bullet(self):
        try:
            index = self.tailor_bullets_text.index(tk.INSERT)
            line = index.split('.')[0]
            line_start = f"{line}.0"
            line_end = f"{line}.end"
            content = self.tailor_bullets_text.get(line_start, line_end).strip()
            return content
        except Exception:
            return None

    def _render_tailor_result(self, results):
        self.tailor_result_text.delete('1.0', tk.END)
        for r in results:
            comp = r['components']
            self.tailor_result_text.insert(tk.END, f"ORIGINAL: {r['original']}\n")
            if r['tailored'] != r['original']:
                self.tailor_result_text.insert(tk.END, f"TAILORED: {r['tailored']}\n")
            self.tailor_result_text.insert(tk.END, f"SCORE: {r['score']} | ACTION:{comp['action']} METRIC:{comp['metric']} CONTEXT:{comp['context']} RESULT:{comp['result']} KW:{comp['keywords']} LEN:{comp['length']}\n")
            self.tailor_result_text.insert(tk.END, "SUGGESTIONS:\n")
            for s in r['suggestions']:
                self.tailor_result_text.insert(tk.END, f"  • {s}\n")
            if r['similar_examples']:
                self.tailor_result_text.insert(tk.END, "EXAMPLES:\n")
                for ex in r['similar_examples'][:5]:
                    self.tailor_result_text.insert(tk.END, f"  ↪ {ex}\n")
            if r['keywords_used']:
                self.tailor_result_text.insert(tk.END, f"KEYWORDS USED: {', '.join(r['keywords_used'])}\n")
            if r.get('metric_ideas'):
                self.tailor_result_text.insert(tk.END, "METRIC IDEAS:\n")
                for idea in r['metric_ideas']:
                    self.tailor_result_text.insert(tk.END, f"  • {idea}\n")
            self.tailor_result_text.insert(tk.END, "\n" + ("-"*70) + "\n\n")

    # Passive voice / real-time formatting suggestions (basic)
        # Add tag configuration for passive voice highlight
        try:
            self.tailor_bullets_text.tag_configure('passive', background='#ffe4e6')
            self.tailor_bullets_text.bind('<KeyRelease>', self._scan_passive_voice)
        except Exception:
            pass

    def _scan_passive_voice(self, event=None):
        text = self.tailor_bullets_text.get('1.0', tk.END)
        self.tailor_bullets_text.tag_remove('passive', '1.0', tk.END)
        passive_pattern = re.compile(r'\b(was|were|is|are|been|being|be)\s+[A-Za-z]+ed\b')
        lines = text.split('\n')
        for i, line in enumerate(lines, start=1):
            if passive_pattern.search(line):
                start = f"{i}.0"
                end = f"{i}.end"
                self.tailor_bullets_text.tag_add('passive', start, end)
        # Also update ATS score on passive scan events
        self._update_ats_score()

    def _update_ats_score(self, event=None):
        """Compute simple keyword coverage: JD unique meaningful tokens vs presence in bullets.
        Also flag basic formatting hazards (tables markers, images, excessive bullets)."""
        import re
        jd = self.tailor_jd_text.get('1.0', tk.END).lower()
        bullets_raw = self.tailor_bullets_text.get('1.0', tk.END).lower()
        # Tokenize: alphanum words length >=4 (avoid stop words roughly)
        tokens = set(re.findall(r'[a-zA-Z]{4,}', jd))
        # Strip very common words
        stop = {'with','from','this','that','have','will','your','such','into','over','under','used','been','also','more','must','need'}
        core_tokens = {t for t in tokens if t not in stop and len(t) < 20}
        matched = sum(1 for t in core_tokens if t in bullets_raw)
        coverage = int((matched / len(core_tokens))*100) if core_tokens else 0
        self.ats_progress.configure(value=coverage)
        self.ats_cov_label.config(text=f"{coverage}%")
        # Hazards: table indicators ("|"), image words, excessive bullet symbols
        hazards = []
        if '|' in bullets_raw:
            hazards.append('Table-like layout detected (|)')
        if re.search(r'image|graphic|screenshot', bullets_raw):
            hazards.append('Image references may confuse parsers')
        if bullets_raw.count('\n') > 40:
            hazards.append('Too many bullet lines (>40)')
        if hazards:
            self.ats_warn_label.config(text='; '.join(hazards[:2]))
        else:
            self.ats_warn_label.config(text='')
        # Color feedback (green/yellow/red)
        if coverage >= 60:
            clr = self.success_color
        elif coverage >= 35:
            clr = self.warning_color
        else:
            clr = '#ef4444'
        try:
            self.ats_cov_label.config(foreground=clr)
        except Exception:
            pass
        # Micro-lessons logic
        try:
            self.micro_lessons_text.config(state='normal')
            self.micro_lessons_text.delete('1.0', tk.END)
            lessons = []
            if coverage < 50:
                lessons.append('Integrate missing high-value job keywords into bullet context naturally (avoid stuffing).')
            if coverage < 30:
                lessons.append('Add a summary line referencing 2–3 core role competencies explicitly.')
            # Metric heuristic: count numbers in bullets_raw
            import re as _re
            numbers = _re.findall(r'\d+', bullets_raw)
            if len(numbers) < 5:
                lessons.append('Increase quantified impact: add before/after metrics (saved $, reduced %, improved time).')
            if _re.search(r'\b(responsible for|duties included)\b', bullets_raw):
                lessons.append('Rewrite passive phrasing ("Responsible for") into direct action ("Led", "Implemented").')
            if not lessons:
                lessons.append('Great progress. Refine verbs for variety (Optimized, Accelerated, Streamlined).')
            for tip in lessons[:5]:
                self.micro_lessons_text.insert(tk.END, f"• {tip}\n")
            self.micro_lessons_text.config(state='disabled')
        except Exception:
            pass

    # === COVER LETTER TAB ===
    def create_cover_letter_tab(self):
        container = ttk.Frame(self.cover_tab, padding='14')
        container.pack(fill=tk.BOTH, expand=True)
        ttk.Label(container, text='Instant Cover Letter Generator', style='Subheading.TLabel').pack(anchor=tk.W)
        ttk.Label(container, text='Paste a Job Description or use fetched one. Generates structured letter using your profile data.', style='Muted.TLabel').pack(anchor=tk.W, pady=(2,10))
        # JD input
        jd_frame = ttk.LabelFrame(container, text='Job Description', padding=8)
        jd_frame.pack(fill=tk.BOTH, expand=True)
        self.cover_jd_text = scrolledtext.ScrolledText(jd_frame, height=14, wrap=tk.WORD, font=('Consolas',9))
        self.cover_jd_text.pack(fill=tk.BOTH, expand=True)
        # Output
        out_frame = ttk.LabelFrame(container, text='Generated Cover Letter', padding=8)
        out_frame.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        self.cover_output_text = scrolledtext.ScrolledText(out_frame, height=16, wrap=tk.WORD, font=('Consolas',9))
        self.cover_output_text.pack(fill=tk.BOTH, expand=True)
        # Actions
        act_frame = ttk.Frame(container)
        act_frame.pack(fill=tk.X, pady=8)
        ttk.Button(act_frame, text='Generate Letter', style='Accent.TButton', command=self._generate_cover_letter).pack(side=tk.LEFT)
        ttk.Button(act_frame, text='Clear', command=lambda: self.cover_output_text.delete('1.0', tk.END)).pack(side=tk.LEFT, padx=6)
        ttk.Button(act_frame, text='Save DOCX', command=self._save_cover_letter_docx).pack(side=tk.LEFT, padx=6)

    def _generate_cover_letter(self):
        import json
        from pathlib import Path
        data_dir = Path(self.root_dir, 'data')
        jd = self.cover_jd_text.get('1.0', tk.END).strip()
        if not jd and getattr(self, 'jd_text', None):  # fallback to apply tab
            jd = self.jd_text.get('1.0', tk.END).strip()

        def _load(name, default):
            p = data_dir / name
            if p.exists():
                try:
                    return json.loads(p.read_text(encoding='utf-8'))
                except Exception:
                    return default
            return default

        contact = _load('profile_contact.json', {})
        candidate = _load('profile_candidate.json', {})
        experience = _load('profile_experience.json', [])

        position = self.position_entry.get().strip() if getattr(self, 'position_entry', None) else 'the role'
        company = self.company_entry.get().strip() if getattr(self, 'company_entry', None) else 'your organization'

        if not generate_cover_letter:
            messagebox.showerror('Module Error', 'Cover letter generator module not available.')
            return

        letter = generate_cover_letter(contact, candidate, experience, position, company, jd)
        self.cover_output_text.delete('1.0', tk.END)
        self.cover_output_text.insert('1.0', letter)
        self.update_status('Cover letter generated (module)')

    def _save_cover_letter_docx(self):
        from pathlib import Path
        from docx import Document
        text = self.cover_output_text.get('1.0', tk.END).strip()
        if not text:
            messagebox.showinfo('No Letter', 'Generate a cover letter first.')
            return
        out_dir = Path(self.default_output_dir)
        out_dir.mkdir(exist_ok=True)
        file_name = f"cover_letter_{datetime.now():%Y%m%d_%H%M%S}.docx"
        out_path = out_dir / file_name
        try:
            doc = Document()
            for para in text.split('\n'):
                if para.strip() == '':
                    doc.add_paragraph('')
                elif para.startswith('•'):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(para.lstrip('• ').strip())
                else:
                    doc.add_paragraph(para)
            doc.save(str(out_path))
            messagebox.showinfo('Saved', f'Cover letter saved to {out_path}')
        except Exception as e:
            messagebox.showerror('Save Error', str(e))

    # === DISTRIBUTION STUB TAB ===
    def create_distribution_tab(self):
        container = ttk.Frame(self.distribution_tab, padding='14')
        container.pack(fill=tk.BOTH, expand=True)
        ttk.Label(container, text='Recruiter Distribution (Roadmap Stub)', style='Subheading.TLabel').pack(anchor=tk.W)
        ttk.Label(container, text='Future capability: curated export packages & tracked outreach.', style='Muted.TLabel').pack(anchor=tk.W, pady=(2,10))
        info = (
            'Planned Features:\n'
            '• Bulk export tailored resume + cover letter sets per target list.\n'
            '• Smart batching with variant rotation to avoid generic impressions.\n'
            '• Contact enrichment (email pattern inference, LinkedIn mapping).\n'
            '• Outreach sequencing (initial, follow-up, value-add message).\n'
            '• Response tracking dashboard & KPI alerts.\n'
            '\nCurrent Placeholder: Use generated files manually for outreach. '
            'Add recruiter list CSV support and email template engine in future release.'
        )
        txt = scrolledtext.ScrolledText(container, height=22, wrap=tk.WORD, font=('Consolas',9))
        txt.pack(fill=tk.BOTH, expand=True)
        txt.insert('1.0', info)
        txt.config(state='disabled')

    # === Templates Tab ===
    def create_templates_tab(self):
        container = ttk.Frame(self.templates_tab, padding='14')
        container.pack(fill=tk.BOTH, expand=True)
        ttk.Label(container, text="Resume Template Profiles", style='Subheading.TLabel').pack(anchor=tk.W)
        ttk.Label(container, text="Select a template to render a preview ordering of sections.", style='Muted.TLabel').pack(anchor=tk.W, pady=(2,10))
        self.template_choice = tk.StringVar(value='Traditional')
        if self.template_engine:
            ttk.Combobox(container, textvariable=self.template_choice, values=self.template_engine.list_templates(), state='readonly', width=22).pack(anchor=tk.W)
        self.template_preview = scrolledtext.ScrolledText(container, height=24, wrap=tk.WORD, font=('Consolas',9))
        self.template_preview.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        ttk.Button(container, text='Render Preview', style='Accent.TButton', command=self._render_template_preview).pack(anchor=tk.W, pady=8)

    def _render_template_preview(self):
        if not self.template_engine:
            return
        tp = self.template_engine.get_template(self.template_choice.get())
        # Mock sections for preview
        sections = {
            'Summary': ["Operations leader improving throughput and cost."],
            'Skills': ["Lean", "SQL", "Process Optimization", "Analytics"],
            'Experience': ["Led plant digitization initiative resulting in 8% cost reduction."],
            'Education': ["B.S. Industrial Engineering"],
            'Projects': ["Implemented Python automation pipeline for reporting."],
        }
        lines = self.template_engine.render_sections(tp, sections)
        self.template_preview.delete('1.0', tk.END)
        self.template_preview.insert('1.0', '\n'.join(lines))

    # === Versions Tab ===
    def create_versions_tab(self):
        container = ttk.Frame(self.versions_tab, padding='14')
        container.pack(fill=tk.BOTH, expand=True)
        ttk.Label(container, text="Resume Versions", style='Subheading.TLabel').pack(anchor=tk.W)
        ttk.Label(container, text="Store and diff multiple resume variants.", style='Muted.TLabel').pack(anchor=tk.W, pady=(2,6))
        self.version_listbox = tk.Listbox(container, height=6)
        self.version_listbox.pack(fill=tk.X)
        btn_frame = ttk.Frame(container)
        btn_frame.pack(fill=tk.X, pady=6)
        ttk.Button(btn_frame, text='Refresh', command=self._refresh_versions).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text='Diff Selected (2)', command=self._diff_versions).pack(side=tk.LEFT, padx=4)
        self.version_diff_text = scrolledtext.ScrolledText(container, height=22, wrap=tk.WORD, font=('Consolas',9))
        self.version_diff_text.pack(fill=tk.BOTH, expand=True)
        self._refresh_versions()

    def _refresh_versions(self):
        if not self.version_manager:
            return
        self.version_listbox.delete(0, tk.END)
        for name in self.version_manager.list_names():
            self.version_listbox.insert(tk.END, name)

    def _diff_versions(self):
        if not self.version_manager:
            return
        selection = self.version_listbox.curselection()
        if len(selection) != 2:
            messagebox.showinfo('Select Two Versions', 'Please select exactly two versions to diff.')
            return
        a = self.version_listbox.get(selection[0])
        b = self.version_listbox.get(selection[1])
        diff = self.version_manager.diff(a, b)
        self.version_diff_text.delete('1.0', tk.END)
        if 'error' in diff:
            self.version_diff_text.insert('1.0', diff['error'])
            return
        self.version_diff_text.insert(tk.END, f"Diff: {a} -> {b}\n\n")
        for k,v in diff.items():
            if k == 'summary':
                self.version_diff_text.insert(tk.END, f"Summary Variants:\n")
                for line in v:
                    self.version_diff_text.insert(tk.END, f"  {line}\n")
                self.version_diff_text.insert(tk.END, '\n')
            else:
                self.version_diff_text.insert(tk.END, f"{k}:\n")
                for item in v:
                    self.version_diff_text.insert(tk.END, f"  • {item}\n")
                self.version_diff_text.insert(tk.END, '\n')

    # === Skills Dashboard Tab ===
    def create_skills_tab(self):
        container = ttk.Frame(self.skills_tab, padding='14')
        container.pack(fill=tk.BOTH, expand=True)
        ttk.Label(container, text="Skills Strength Dashboard", style='Subheading.TLabel').pack(anchor=tk.W)
        ttk.Label(container, text="Aggregated frequency from current experience bullets.", style='Muted.TLabel').pack(anchor=tk.W, pady=(2,6))
        ttk.Button(container, text='Recompute', style='Accent.TButton', command=self._compute_skills_dashboard).pack(anchor=tk.W, pady=(0,8))
        self.skills_tree = ttk.Treeview(container, columns=('skill','count'), show='headings', height=22)
        self.skills_tree.heading('skill', text='Skill / Token')
        self.skills_tree.heading('count', text='Count')
        self.skills_tree.column('skill', width=240)
        self.skills_tree.column('count', width=70, anchor=tk.CENTER)
        self.skills_tree.pack(fill=tk.BOTH, expand=True)
        self._compute_skills_dashboard()

    def _compute_skills_dashboard(self):
        # Load experience bullets from profile_experience.json
        import json
        exp_path = ROOT / 'data' / 'profile_experience.json'
        bullets: list[str] = []
        try:
            if exp_path.exists():
                data = json.loads(exp_path.read_text(encoding='utf-8'))
                for role in data:
                    for b in role.get('bullets', []):
                        bullets.append(b)
        except Exception:
            pass
        freq = extract_skills(bullets)
        top = top_skills(freq, 50)
        for item in self.skills_tree.get_children():
            self.skills_tree.delete(item)
        for skill,count in top:
            self.skills_tree.insert('', tk.END, values=(skill, count))

    # === ATS SCANNER (Stub) ===
    def create_ats_tab(self):
        container = ttk.Frame(self.ats_tab, padding="14")
        container.pack(fill=tk.BOTH, expand=True)
        ttk.Label(container, text="ATS Deep Scanner", style='Subheading.TLabel').pack(anchor=tk.W)
        ttk.Label(container, text="Analyze resume vs job description: keyword gaps, passive voice density, formatting hazards.", style='Muted.TLabel').pack(anchor=tk.W, pady=(2,10))
        top_frame = ttk.Frame(container)
        top_frame.pack(fill=tk.X)
        # Buttons
        ttk.Button(top_frame, text='Load Resume File', style='Accent.TButton', command=self._ats_load_resume).pack(side=tk.LEFT)
        ttk.Button(top_frame, text='Use Current JD', command=self._ats_use_current_jd).pack(side=tk.LEFT, padx=6)
        ttk.Button(top_frame, text='Clear', command=self._ats_clear).pack(side=tk.LEFT, padx=6)
        # Text areas
        split = ttk.Frame(container)
        split.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        resume_frame = ttk.LabelFrame(split, text='Resume Text', padding=6)
        resume_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.ats_resume_text = scrolledtext.ScrolledText(resume_frame, height=20, wrap=tk.WORD, font=('Consolas',9))
        self.ats_resume_text.pack(fill=tk.BOTH, expand=True)
        jd_frame = ttk.LabelFrame(split, text='Job Description', padding=6)
        jd_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(8,0))
        self.ats_jd_text = scrolledtext.ScrolledText(jd_frame, height=20, wrap=tk.WORD, font=('Consolas',9))
        self.ats_jd_text.pack(fill=tk.BOTH, expand=True)
        # Results
        results_frame = ttk.LabelFrame(container, text='Analysis Results', padding=6)
        results_frame.pack(fill=tk.BOTH, expand=True, pady=(10,0))
        self.ats_results_text = scrolledtext.ScrolledText(results_frame, height=14, wrap=tk.WORD, font=('Consolas',9))
        self.ats_results_text.pack(fill=tk.BOTH, expand=True)
        ttk.Button(container, text='Run Analysis', command=self._run_ats_analysis).pack(anchor=tk.W, pady=(8,0))

    def _ats_clear(self):
        self.ats_resume_text.delete('1.0', tk.END)
        self.ats_jd_text.delete('1.0', tk.END)
        self.ats_results_text.delete('1.0', tk.END)

    def _ats_use_current_jd(self):
        try:
            jd = self.jd_text.get('1.0', tk.END)
            self.ats_jd_text.delete('1.0', tk.END)
            self.ats_jd_text.insert('1.0', jd)
        except Exception:
            pass

    def _ats_load_resume(self):
        path = filedialog.askopenfilename(title='Select Resume', filetypes=[('Documents','*.docx *.txt *.pdf')])
        if not path:
            return
        text = ''
        try:
            import os
            ext = os.path.splitext(path)[1].lower()
            if ext == '.docx':
                from docx import Document
                d = Document(path)
                text = '\n'.join(p.text for p in d.paragraphs if p.text.strip())
            elif ext == '.txt':
                with open(path,'r',encoding='utf-8',errors='ignore') as f:
                    text = f.read()
            elif ext == '.pdf':
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(path)
                    parts = []
                    for page in reader.pages:
                        try: parts.append(page.extract_text() or '')
                        except Exception: pass
                    text = '\n'.join(parts)
                except Exception as e:
                    messagebox.showwarning('PDF Parse', f'PDF parsing failed: {e}')
                    return
        except Exception as e:
            messagebox.showerror('Load Error', str(e)); return
        self.ats_resume_text.delete('1.0', tk.END)
        self.ats_resume_text.insert('1.0', text)
        self.update_status('Resume loaded for ATS analysis')

    def _run_ats_analysis(self):
        resume = self.ats_resume_text.get('1.0', tk.END)
        jd = self.ats_jd_text.get('1.0', tk.END)
        if not resume or not jd:
            messagebox.showinfo('Missing Data','Load resume and job description first.')
            return
        if not ats_analyze:
            messagebox.showerror('Module Error', 'ATS analyzer module not available.')
            return
        result = ats_analyze(resume, jd)
        self.ats_results_text.delete('1.0', tk.END)
        self.ats_results_text.insert(tk.END, f"KEYWORD COVERAGE: {result['keyword_coverage']}%\n")
        self.ats_results_text.insert(tk.END, f"MISSING KEYWORDS (top 25): {', '.join(result['missing_keywords'][:25]) or 'None'}\n")
        self.ats_results_text.insert(tk.END, f"PASSIVE VOICE DENSITY: {result['passive_density_pct_lines']:.1f}% of lines\n")
        self.ats_results_text.insert(tk.END, f"LINES WITH METRICS: {result['metric_lines_ratio_pct']}%\n")
        hazards_line = ', '.join(result['hazards']) if result['hazards'] else 'None detected'
        self.ats_results_text.insert(tk.END, f"FORMAT HAZARDS: {hazards_line}\n\n")
        self.ats_results_text.insert(tk.END, 'RECOMMENDATIONS:\n')
        for r in result['recommendations']:
            self.ats_results_text.insert(tk.END, f"• {r}\n")
        self.update_status('ATS analysis complete')
    
    def create_apply_tab(self):
        """Create the apply to job tab with full-window scrolling"""
        canvas = tk.Canvas(self.apply_tab)
        v_scroll = ttk.Scrollbar(self.apply_tab, orient="vertical", command=canvas.yview)
        inner = ttk.Frame(canvas, padding="20")
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0,0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=v_scroll.set)
        canvas.pack(side="left", fill="both", expand=True)
        v_scroll.pack(side="right", fill="y")
        def _wheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _wheel)
        container = inner
        
        # Modern card-style input sections
        # Company Name Card
        company_card = tk.Frame(container, bg=self.subtle_bg, relief='flat', bd=0)
        company_card.grid(row=0, column=0, sticky=tk.EW, pady=(0, 16), padx=2)
        tk.Label(company_card, text="Company Name", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color, anchor='w').pack(fill=tk.X, padx=16, pady=(12, 4))
        self.company_entry = tk.Entry(company_card, font=self.font_regular, relief='flat', bd=0, bg='#f1f5f9', fg=self.text_color)
        self.company_entry.pack(fill=tk.X, padx=16, pady=(0, 12), ipady=8)
        
        # Position Title Card
        position_card = tk.Frame(container, bg=self.subtle_bg, relief='flat', bd=0)
        position_card.grid(row=1, column=0, sticky=tk.EW, pady=(0, 16), padx=2)
        tk.Label(position_card, text="Position Title", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color, anchor='w').pack(fill=tk.X, padx=16, pady=(12, 4))
        self.position_entry = tk.Entry(position_card, font=self.font_regular, relief='flat', bd=0, bg='#f1f5f9', fg=self.text_color)
        self.position_entry.pack(fill=tk.X, padx=16, pady=(0, 12), ipady=8)
        
        # Job URL Card
        url_card = tk.Frame(container, bg=self.subtle_bg, relief='flat', bd=0)
        url_card.grid(row=2, column=0, sticky=tk.EW, pady=(0, 16), padx=2)
        tk.Label(url_card, text="Job URL (optional - auto-fetch)", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color, anchor='w').pack(fill=tk.X, padx=16, pady=(12, 4))
        
        url_inner = tk.Frame(url_card, bg=self.subtle_bg)
        url_inner.pack(fill=tk.X, padx=16, pady=(0, 12))
        
        self.url_entry = tk.Entry(url_inner, font=self.font_regular, relief='flat', bd=0, bg='#f1f5f9', fg=self.text_color)
        self.url_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8)
        
        fetch_btn = tk.Button(
            url_inner,
            text="🌐 Fetch",
            command=self.fetch_from_url,
            font=self.font_bold,
            bg=self.accent_color,
            fg='white',
            relief='flat',
            bd=0,
            padx=16,
            pady=8,
            cursor='hand2'
        )
        fetch_btn.pack(side=tk.LEFT, padx=(8, 0))
        
        # Job Description Card
        jd_card = tk.Frame(container, bg=self.subtle_bg, relief='flat', bd=0)
        jd_card.grid(row=3, column=0, sticky=tk.NSEW, pady=(0, 16), padx=2)
        tk.Label(jd_card, text="Job Description", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color, anchor='w').pack(fill=tk.X, padx=16, pady=(12, 8))
        
        self.jd_text = scrolledtext.ScrolledText(
            jd_card,
            width=60,
            height=15,
            font=("Consolas", 9),
            wrap=tk.WORD,
            relief='flat',
            bd=0,
            bg='#f1f5f9',
            fg=self.text_color
        )
        self.jd_text.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 12))
        # Automatic field inference when user pastes a raw JD
        self._jd_auto_timer = None
        self.jd_text.bind('<KeyRelease>', self._on_jd_key_activity)
        self.jd_text.bind('<<Paste>>', self._on_jd_key_activity)
        self.jd_text.bind('<FocusOut>', self._on_jd_focus_out)
        
        # Action Buttons
        button_frame = tk.Frame(container, bg=self.bg_color)
        button_frame.grid(row=4, column=0, pady=(8, 0), sticky=tk.EW)
        
        # Load JD from file button
        load_btn = tk.Button(
            button_frame,
            text="📁 Load from File",
            command=self.load_jd_file,
            font=self.font_regular,
            bg=self.subtle_bg,
            fg=self.text_color,
            relief='flat',
            bd=0,
            padx=16,
            pady=10,
            cursor='hand2'
        )
        load_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Clear button
        clear_btn = tk.Button(
            button_frame,
            text="🗑 Clear",
            command=self.clear_form,
            font=self.font_regular,
            bg=self.subtle_bg,
            fg=self.text_color,
            relief='flat',
            bd=0,
            padx=16,
            pady=10,
            cursor='hand2'
        )
        clear_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Submit button
        submit_btn = tk.Button(
            button_frame,
            text="✅ Generate Application",
            command=self.submit_application,
            font=self.font_bold,
            bg=self.success_color,
            fg='white',
            relief='flat',
            bd=0,
            padx=24,
            pady=12,
            cursor='hand2'
        )
        submit_btn.pack(side=tk.LEFT)
        
        # Configure grid weights
        container.columnconfigure(0, weight=1)
        container.rowconfigure(3, weight=1)
    
    def fetch_from_url(self):
        """Fetch job description from URL using AI web scraper"""
        url = self.url_entry.get().strip()
        
        if not url:
            messagebox.showinfo("URL Required", "Please enter a job posting URL first.")
            return
        
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
            self.url_entry.delete(0, tk.END)
            self.url_entry.insert(0, url)
        
        # Domain-aware status
        if 'myworkdayjobs.com' in url.lower():
            self.update_status("🔍 Workday job detected – enhanced parser active…")
        else:
            self.update_status("🌐 Fetching job description from URL…")
        
        # Monday.com auto-detection: if user just pastes a Monday item URL, fetch directly
        if 'monday.com' in url:
            # Try enhanced Monday URL handling (no manual token entry, attempt env/config, else public scrape)
            thread = threading.Thread(target=self._fetch_monday_from_url_worker, args=(url,))
            thread.daemon = True
            thread.start()
            return

        # Generic fetch
        thread = threading.Thread(target=self._fetch_url_worker, args=(url,))
        thread.daemon = True
        thread.start()

    def _extract_monday_item_id(self, url: str):
        """Extract Monday item (pulse) ID from a Monday.com URL.
        Supports URLs containing /pulses/<id>, /items/<id>, or query param pulseId=<id>.
        Returns int or None.
        """
        import re
        patterns = [r'/pulses/(\d+)', r'/items/(\d+)', r'[?&]pulseId=(\d+)']
        for pat in patterns:
            m = re.search(pat, url)
            if m:
                try:
                    return int(m.group(1))
                except ValueError:
                    pass
        return None

    def _fetch_monday_from_url_worker(self, url: str):
        """Worker to fetch Monday job description given full item URL.
        Tries GraphQL if token present (env MONDAY_API_TOKEN or config), else attempts public HTML scrape.
        """
        try:
            import os, json, requests, re
            from bs4 import BeautifulSoup
            item_id = self._extract_monday_item_id(url)
            if not item_id:
                self.update_status("❌ Could not parse Monday item ID")
                messagebox.showerror("Monday.com", "Could not extract item ID from URL. Please use the item page URL.")
                return

            cfg_path = ROOT / 'data' / 'monday_config.json'
            cfg = {}
            if cfg_path.exists():
                try:
                    with open(cfg_path, 'r', encoding='utf-8') as f:
                        cfg = json.load(f)
                except Exception:
                    cfg = {}
            api_token = os.getenv('MONDAY_API_TOKEN') or cfg.get('api_token')
            jd_column_id = cfg.get('jd_column_id')
            company_column_id = cfg.get('company_column_id')

            position_title = ''
            company_name = ''
            jd_text = ''

            if api_token:
                # GraphQL path
                self.update_status("🔗 Monday.com (API)…")
                headers = {'Authorization': api_token, 'Content-Type': 'application/json'}
                query = (
                    'query ($ids: [Int]) {\n'
                    '  items (ids: $ids) {\n'
                    '    name\n'
                    '    column_values { id title text value }\n'
                    '  }\n'
                    '}'
                )
                payload = {"query": query, "variables": {"ids": [item_id]}}
                resp = requests.post('https://api.monday.com/v2', headers=headers, json=payload, timeout=30)
                if resp.status_code == 200:
                    data = resp.json()
                    items = (data or {}).get('data', {}).get('items', [])
                    if items:
                        item = items[0]
                        position_title = item.get('name') or ''
                        for col in item.get('column_values', []):
                            cid = col.get('id') or ''
                            title = (col.get('title') or '').lower()
                            text = col.get('text') or ''
                            if not jd_text and (cid == jd_column_id or 'description' in title or 'job description' in title):
                                jd_text = text
                            if not company_name and (cid == company_column_id or 'company' in title or 'employer' in title):
                                company_name = text
                else:
                    # Fall back to scrape if token invalid
                    api_token = None

            if not api_token:
                # Public HTML scrape (best effort; may fail if board is private)
                self.update_status("🌐 Monday.com (public scrape)…")
                try:
                    html_resp = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
                    if html_resp.status_code == 200:
                        soup = BeautifulSoup(html_resp.text, 'html.parser')
                        # Heuristics: look for large text blocks that might be description
                        text_blobs = []
                        for tag in soup.find_all(['p','div','span']):
                            t = tag.get_text(strip=True)
                            if t and len(t) > 60:
                                text_blobs.append(t)
                        # Deduplicate
                        dedup = []
                            
                        for t in text_blobs:
                            if t not in dedup:
                                dedup.append(t)
                        # Choose the longest as description candidate
                        if not jd_text and dedup:
                            jd_text = max(dedup, key=len)
                        # Position guess: first heading-like text
                        headings = [h.get_text(strip=True) for h in soup.find_all(['h1','h2','h3']) if h.get_text(strip=True)]
                        if headings and not position_title:
                            position_title = headings[0][:120]
                        # Company guess: domain segment
                        if not company_name:
                            m = re.search(r'https?://([^/]+)/', url + '/')
                            if m:
                                company_name = m.group(1).split('.')[0].title()
                except Exception:
                    pass

            # Populate UI
            if position_title:
                self.position_entry.delete(0, tk.END)
                self.position_entry.insert(0, position_title)
            if company_name:
                self.company_entry.delete(0, tk.END)
                self.company_entry.insert(0, company_name)
            if jd_text:
                self.jd_text.delete('1.0', tk.END)
                self.jd_text.insert('1.0', jd_text)

            if jd_text or position_title:
                self.update_status("✅ Monday.com URL loaded")
                messagebox.showinfo("Monday.com", "Loaded job info from URL.")
            else:
                self.update_status("⚠️ Monday.com URL processed")
                messagebox.showwarning("Monday.com", "Could not extract job description automatically. You may need to use the manual Monday button or copy/paste.")

        except Exception as e:
            self.update_status("❌ Monday.com fetch failed")
            messagebox.showerror("Monday.com", f"Failed to fetch from URL:\n\n{e}")

    def load_jd_from_monday(self):
        """Fetch a job description from Monday.com given an item ID.
        Looks for config in data/monday_config.json; prompts for missing values.
        """
        try:
            import json
            import os
            import requests
            from tkinter import simpledialog

            cfg_path = ROOT / 'data' / 'monday_config.json'
            cfg = {}
            if cfg_path.exists():
                try:
                    with open(cfg_path, 'r', encoding='utf-8') as f:
                        cfg = json.load(f)
                except Exception:
                    cfg = {}

            # Get token (env wins, then config, then prompt)
            api_token = os.getenv('MONDAY_API_TOKEN') or cfg.get('api_token')
            if not api_token:
                api_token = simpledialog.askstring("Monday.com", "Enter Monday API token:")
                if not api_token:
                    return

            # Ask for Item ID
            item_id_str = simpledialog.askstring("Monday.com", "Enter Monday Item ID (from the item URL):")
            if not item_id_str:
                return
            try:
                item_id = int(item_id_str)
            except ValueError:
                messagebox.showerror("Invalid Item ID", "Item ID must be a number.")
                return

            # Optional: JD column id from config
            jd_column_id = cfg.get('jd_column_id')
            company_column_id = cfg.get('company_column_id')

            # GraphQL query to fetch item + all column values
            url = 'https://api.monday.com/v2'
            headers = {
                'Authorization': api_token,
                'Content-Type': 'application/json'
            }
            query = (
                'query ($ids: [Int]) {\n'
                '  items (ids: $ids) {\n'
                '    name\n'
                '    column_values { id title text value }\n'
                '  }\n'
                '}'
            )
            payload = {"query": query, "variables": {"ids": [item_id]}}

            self.update_status("🔗 Fetching from Monday.com…")
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            if resp.status_code != 200:
                messagebox.showerror("Monday.com Error", f"HTTP {resp.status_code}: {resp.text[:300]}")
                return
            data = resp.json()
            items = (data or {}).get('data', {}).get('items', [])
            if not items:
                messagebox.showinfo("Not Found", "No item found with that ID.")
                return

            item = items[0]
            position_title = item.get('name') or ''
            jd_text = ''
            company_name = ''

            # Try direct IDs first, then fallback by title matching
            for col in item.get('column_values', []):
                cid = col.get('id') or ''
                title = (col.get('title') or '').lower()
                text = col.get('text') or ''
                if not jd_text and (cid == jd_column_id or 'description' in title or 'job description' in title):
                    jd_text = text
                if not company_name and (cid == company_column_id or 'company' in title or 'employer' in title):
                    company_name = text

            # Populate UI
            if position_title:
                self.position_entry.delete(0, tk.END)
                self.position_entry.insert(0, position_title)
            if company_name:
                self.company_entry.delete(0, tk.END)
                self.company_entry.insert(0, company_name)
            if jd_text:
                self.jd_text.delete('1.0', tk.END)
                self.jd_text.insert('1.0', jd_text)

            if jd_text or position_title or company_name:
                messagebox.showinfo("Monday.com", "✅ Loaded job details from Monday.com")
                self.update_status("✅ Monday.com job loaded")
            else:
                messagebox.showinfo("Monday.com", "Item loaded but no recognizable fields were found. You may need to configure column IDs in data/monday_config.json.")
                self.update_status("ℹ️ Monday.com item loaded with no JD")

        except Exception as e:
            messagebox.showerror("Monday.com", f"Failed to load from Monday.com:\n\n{e}")
            self.update_status("❌ Monday.com load failed")
    
    def _fetch_url_worker(self, url):
        """Background worker to fetch URL"""
        try:
            result = self.ai_assistant.fetch_job_from_url(url)
            self.root.after(0, lambda: self._handle_url_fetch_result(result))
        except Exception as e:
            self.root.after(0, lambda: self._handle_url_fetch_error(str(e)))
    
    def _handle_url_fetch_result(self, result):
        """Handle URL fetch result in main thread"""
        if result.get('success'):
            # Fill in the JD text
            jd_text = result.get('job_description', '')
            if jd_text:
                self.jd_text.delete('1.0', tk.END)
                self.jd_text.insert('1.0', jd_text)
            
            # Auto-fill company and position if found
            if result.get('company'):
                self.company_entry.delete(0, tk.END)
                self.company_entry.insert(0, result['company'])
            
            if result.get('position'):
                self.position_entry.delete(0, tk.END)
                self.position_entry.insert(0, result['position'])
            
            info_parts = []
            if result.get('company'):
                info_parts.append(f"Company: {result['company']}")
            if result.get('position'):
                info_parts.append(f"Position: {result['position']}")
            if result.get('location'):
                info_parts.append(f"Location: {result['location']}")
            
            message = "✅ Successfully fetched job description!\n\n"
            if info_parts:
                message += "Extracted:\n" + "\n".join(info_parts)
            
            messagebox.showinfo("URL Fetch Complete", message)
            self.update_status("✅ Job description loaded from URL")
        else:
            error = result.get('error', 'Unknown error')
            messagebox.showwarning("URL Fetch Failed", error)
            self.update_status("❌ Failed to fetch from URL")
    
    def _handle_url_fetch_error(self, error_msg):
        """Handle URL fetch error"""
        messagebox.showerror("Error", f"Failed to fetch URL:\n{error_msg}")
        self.update_status("❌ URL fetch error")
    
    def create_view_tab(self):
        """Create the view applications tab with full-window scrolling"""
        canvas = tk.Canvas(self.view_tab)
        v_scroll = ttk.Scrollbar(self.view_tab, orient="vertical", command=canvas.yview)
        inner = ttk.Frame(canvas, padding="20")
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0,0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=v_scroll.set)
        canvas.pack(side="left", fill="both", expand=True)
        v_scroll.pack(side="right", fill="y")
        def _wheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _wheel)
        container = inner
        
        # Title
        ttk.Label(
            container,
            text="Recent Applications",
            font=("Segoe UI", 14, "bold")
        ).pack(pady=(0, 10))
        
        # Buttons
        btn_frame = ttk.Frame(container)
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(
            btn_frame,
            text="📤 Upload Tracker",
            command=self.upload_tracker
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="📊 Open Tracker",
            command=self.open_tracker
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="📂 Open Output Folder",
            command=self.open_output_folder
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame,
            text="🔄 Refresh",
            command=self.refresh_applications
        ).pack(side=tk.LEFT, padx=5)
        
        # Treeview for applications
        tree_frame = ttk.Frame(container)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(tree_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Treeview
        self.app_tree = ttk.Treeview(
            tree_frame,
            columns=("Company", "Position", "Date", "Source", "Status", "Match", "Fit"),
            show="headings",
            yscrollcommand=scrollbar.set
        )
        scrollbar.config(command=self.app_tree.yview)
        
        # Configure columns
        self.app_tree.heading("Company", text="Company")
        self.app_tree.heading("Position", text="Position")
        self.app_tree.heading("Date", text="Date Applied")
        self.app_tree.heading("Source", text="Source")
        self.app_tree.heading("Status", text="Status")
        self.app_tree.heading("Match", text="Skills Match %")
        self.app_tree.heading("Fit", text="Job Fit %")
        
        self.app_tree.column("Company", width=160)
        self.app_tree.column("Position", width=200)
        self.app_tree.column("Date", width=90)
        self.app_tree.column("Source", width=80)
        self.app_tree.column("Status", width=90)
        self.app_tree.column("Match", width=90)
        self.app_tree.column("Fit", width=80)
        
        self.app_tree.pack(fill=tk.BOTH, expand=True)
        
        # Configure tags for color-coded match scores - modern colors
        self.app_tree.tag_configure('high_match', background='#d1fae5', foreground='#065f46')
        self.app_tree.tag_configure('medium_match', background='#fef3c7', foreground='#92400e')
        self.app_tree.tag_configure('low_match', background='#fee2e2', foreground='#991b1b')
        # Zebra striping (applied second; color tags override when present)
        self.app_tree.tag_configure('zebra_even', background='#ffffff')
        self.app_tree.tag_configure('zebra_odd', background='#f9fafb')

        # Bind selection event for editing
        self.app_tree.bind("<<TreeviewSelect>>", self.on_application_select)

        # Modern info card for disclaimer
        info_card = tk.Frame(container, bg='#eff6ff', relief='flat', bd=0)
        info_card.pack(fill=tk.X, pady=(16, 0), padx=2)
        
        # Blue accent bar on left
        tk.Frame(info_card, bg=self.accent_color, width=4).pack(side=tk.LEFT, fill=tk.Y)
        
        info_content = tk.Frame(info_card, bg='#eff6ff')
        info_content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=16, pady=12)
        
        disclaimer_text = (
            "💡 Skills Match % = How well YOU match the JOB (your skills/experience vs their requirements). "
            "Job Fit % = How well the JOB matches YOU (role characteristics vs your preferences like location, salary, work style). "
            "Both are automated estimates only. Never let low scores discourage you from applying to roles you're interested in!"
        )
        
        disclaimer_label = tk.Label(
            info_content,
            text=disclaimer_text,
            wraplength=900,
            justify=tk.LEFT,
            bg='#eff6ff',
            fg='#1e40af',
            font=("Segoe UI", 9),
            anchor='w'
        )
        disclaimer_label.pack(fill=tk.X)

        # Enhanced Details / Edit pane with modern card styling
        details_card = tk.Frame(container, bg=self.subtle_bg, relief='flat', bd=0)
        details_card.pack(fill=tk.BOTH, expand=True, pady=(16, 0), padx=2)
        
        tk.Label(details_card, text="Selected Application Details", font=self.font_subheading, bg=self.subtle_bg, fg=self.text_color, anchor='w').pack(fill=tk.X, padx=16, pady=(16, 12))
        
        details = tk.Frame(details_card, bg=self.subtle_bg)
        details.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 16))

        # Row 0: Company and Source
        tk.Label(details, text="Company:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=0, column=0, sticky=tk.W, pady=4, padx=(0, 8))
        self.view_company_var = tk.StringVar()
        company_entry = tk.Entry(details, textvariable=self.view_company_var, font=self.font_regular, relief='flat', bd=1, bg='#f1f5f9')
        company_entry.grid(row=0, column=1, sticky=tk.EW, pady=4, ipady=6)
        
        tk.Label(details, text="Source:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=0, column=2, sticky=tk.W, pady=4, padx=(16, 8))
        self.view_source_var = tk.StringVar()
        source_label = tk.Label(details, textvariable=self.view_source_var, font=self.font_regular, bg='#f1f5f9', fg=self.text_muted, anchor='w', padx=8, pady=6)
        source_label.grid(row=0, column=3, sticky=tk.EW, pady=4)

        # Row 1: Position
        tk.Label(details, text="Position:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=1, column=0, sticky=tk.W, pady=4, padx=(0, 8))
        self.view_position_var = tk.StringVar()
        position_entry = tk.Entry(details, textvariable=self.view_position_var, font=self.font_regular, relief='flat', bd=1, bg='#f1f5f9')
        position_entry.grid(row=1, column=1, columnspan=3, sticky=tk.EW, pady=4, ipady=6)

        # Row 2: Date, Status, Match Score, Job Fit
        tk.Label(details, text="Date Applied:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=2, column=0, sticky=tk.W, pady=4, padx=(0, 8))
        self.view_date_var = tk.StringVar()
        date_entry = tk.Entry(details, textvariable=self.view_date_var, font=self.font_regular, relief='flat', bd=1, bg='#f1f5f9', width=15)
        date_entry.grid(row=2, column=1, sticky=tk.W, pady=4, ipady=6)

        tk.Label(details, text="Status:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=2, column=2, sticky=tk.W, pady=4, padx=(16, 8))
        self.view_status_var = tk.StringVar()
        status_options = ['Not Applied','Applied','Phone Screen','Interview 1','Interview 2','Interview 3+','Offer','Accepted','Rejected','Withdrawn']
        self.view_status_combo = ttk.Combobox(details, textvariable=self.view_status_var, values=status_options, width=18, state='readonly')
        self.view_status_combo.grid(row=2, column=3, sticky=tk.W, pady=4)
        self.view_status_combo.bind('<<ComboboxSelected>>', self._on_status_changed)

        # Row 3: Skills Match and Job Fit scores (read-only display)
        tk.Label(details, text="Skills Match:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=3, column=0, sticky=tk.W, pady=4, padx=(0, 8))
        self.view_match_var = tk.StringVar(value="N/A")
        match_display = tk.Label(details, textvariable=self.view_match_var, font=("Segoe UI", 10, "bold"), bg='#f1f5f9', fg=self.accent_color, anchor='w', padx=8, pady=6)
        match_display.grid(row=3, column=1, sticky=tk.W, pady=4)

        tk.Label(details, text="Job Fit:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=3, column=2, sticky=tk.W, pady=4, padx=(16, 8))
        self.view_fit_var = tk.StringVar(value="N/A")
        fit_display = tk.Label(details, textvariable=self.view_fit_var, font=("Segoe UI", 10, "bold"), bg='#f1f5f9', fg=self.accent_color, anchor='w', padx=8, pady=6)
        fit_display.grid(row=3, column=3, sticky=tk.W, pady=4)

        # Row 4: URL (if available)
        tk.Label(details, text="Job URL:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=4, column=0, sticky=tk.NW, pady=4, padx=(0, 8))
        self.view_url_var = tk.StringVar()
        url_text = tk.Text(details, height=2, font=("Consolas", 8), relief='flat', bd=1, bg='#f1f5f9', fg=self.text_muted, wrap=tk.WORD)
        url_text.grid(row=4, column=1, columnspan=3, sticky=tk.EW, pady=4)
        self.view_url_text = url_text

        # Row 5: Notes
        tk.Label(details, text="Notes:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=5, column=0, sticky=tk.NW, pady=4, padx=(0, 8))
        self.view_notes_text = tk.Text(details, width=60, height=5, font=("Consolas",9), relief='flat', bd=1, bg='#f1f5f9', wrap=tk.WORD)
        self.view_notes_text.grid(row=5, column=1, columnspan=3, sticky=tk.EW, pady=4)

        # Row 6: Lessons Learned
        tk.Label(details, text="Lessons Learned:", font=self.font_bold, bg=self.subtle_bg, fg=self.text_color).grid(row=6, column=0, sticky=tk.NW, pady=4, padx=(0, 8))
        self.view_lessons_text = tk.Text(details, width=60, height=4, font=("Consolas",9), relief='flat', bd=1, bg='#f1f5f9', wrap=tk.WORD)
        self.view_lessons_text.grid(row=6, column=1, columnspan=3, sticky=tk.EW, pady=4)

        # Configure column weights for proper resizing
        details.columnconfigure(1, weight=1)
        details.columnconfigure(3, weight=1)

        # Action buttons with modern styling
        action_frame = tk.Frame(details_card, bg=self.subtle_bg)
        action_frame.pack(fill=tk.X, padx=16, pady=(0, 16))
        
        save_btn = tk.Button(action_frame, text="💾 Save Changes", command=self.update_selected_application,
                            font=self.font_bold, bg=self.accent_color, fg='white', relief='flat', bd=0,
                            padx=16, pady=10, cursor='hand2')
        save_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        delete_btn = tk.Button(action_frame, text="🗑 Delete", command=self.delete_selected_application,
                               font=self.font_regular, bg='#ef4444', fg='white', relief='flat', bd=0,
                               padx=16, pady=10, cursor='hand2')
        delete_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        export_btn = tk.Button(action_frame, text="📤 Export", command=self.export_selected_application,
                               font=self.font_regular, bg=self.subtle_bg, fg=self.text_color, relief='flat', bd=1,
                               padx=16, pady=10, cursor='hand2')
        export_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        clear_btn = tk.Button(action_frame, text="↻ Clear", command=self.clear_view_selection,
                              font=self.font_regular, bg=self.subtle_bg, fg=self.text_color, relief='flat', bd=1,
                              padx=16, pady=10, cursor='hand2')
        clear_btn.pack(side=tk.LEFT)
        
        # Load initial data
        self.refresh_applications()

    
    # === FRAMEWORK VALIDATOR TAB ===
    def create_validator_tab(self):
        """Create framework validator tab for bullet validation and improvement."""
        # Main container with padding
        container = ttk.Frame(self.validator_tab, padding="16")
        container.pack(fill=tk.BOTH, expand=True)
        
        # Header section
        header_frame = ttk.Frame(container)
        header_frame.pack(fill=tk.X, pady=(0, 16))
        
        ttk.Label(
            header_frame,
            text="Resume Bullet Validator",
            style='Subheading.TLabel'
        ).pack(side=tk.LEFT)
        
        ttk.Label(
            header_frame,
            text="Validate your resume bullets against industry-standard frameworks",
            style='Muted.TLabel'
        ).pack(side=tk.LEFT, padx=(12, 0))
        
        # Two-column layout
        left_panel = ttk.Frame(container)
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 8))
        
        right_panel = ttk.Frame(container)
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(8, 0))
        
        # === LEFT PANEL: Input ===
        input_card = ttk.LabelFrame(left_panel, text="Resume Bullets", padding="12")
        input_card.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(
            input_card,
            text="Paste your resume bullets below (one per line):",
            style='Muted.TLabel'
        ).pack(anchor=tk.W, pady=(0, 8))
        
        # Text area for bullets
        self.validator_input_text = scrolledtext.ScrolledText(
            input_card,
            height=15,
            font=self.font_regular,
            wrap=tk.WORD,
            bg=self.subtle_bg,
            fg=self.text_color,
            relief='solid',
            borderwidth=1
        )
        self.validator_input_text.pack(fill=tk.BOTH, expand=True, pady=(0, 12))
        
        # Framework selection
        framework_frame = ttk.Frame(input_card)
        framework_frame.pack(fill=tk.X, pady=(0, 12))
        
        ttk.Label(framework_frame, text="Framework:").pack(side=tk.LEFT, padx=(0, 8))
        
        self.validator_framework = tk.StringVar(value="STAR")
        framework_combo = ttk.Combobox(
            framework_frame,
            textvariable=self.validator_framework,
            values=["STAR", "CAR", "PAR", "WHO", "LPS"],
            state="readonly",
            width=15
        )
        framework_combo.pack(side=tk.LEFT)
        
        # Info button
        info_btn = ttk.Button(
            framework_frame,
            text="ℹ️ Framework Info",
            command=self.show_framework_info
        )
        info_btn.pack(side=tk.LEFT, padx=(8, 0))
        
        # Action buttons
        btn_frame = ttk.Frame(input_card)
        btn_frame.pack(fill=tk.X)
        
        ttk.Button(
            btn_frame,
            text="🔍 Validate Bullets",
            style='Accent.TButton',
            command=self.validate_bullets
        ).pack(side=tk.LEFT, padx=(0, 8))
        
        ttk.Button(
            btn_frame,
            text="Clear",
            command=lambda: self.validator_input_text.delete('1.0', tk.END)
        ).pack(side=tk.LEFT)
        
        # === RIGHT PANEL: Results ===
        results_card = ttk.LabelFrame(right_panel, text="Validation Results", padding="12")
        results_card.pack(fill=tk.BOTH, expand=True)
        
        # Results summary label
        self.validator_summary_label = ttk.Label(
            results_card,
            text="No validation results yet",
            style='Muted.TLabel'
        )
        self.validator_summary_label.pack(anchor=tk.W, pady=(0, 8))
        
        # Results tree
        tree_frame = ttk.Frame(results_card)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        tree_scroll = ttk.Scrollbar(tree_frame)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.validator_tree = ttk.Treeview(
            tree_frame,
            columns=("bullet", "score", "grade", "components"),
            show="headings",
            yscrollcommand=tree_scroll.set,
            height=12
        )
        tree_scroll.config(command=self.validator_tree.yview)
        
        self.validator_tree.heading("bullet", text="Bullet Point")
        self.validator_tree.heading("score", text="Score")
        self.validator_tree.heading("grade", text="Grade")
        self.validator_tree.heading("components", text="Components")
        
        self.validator_tree.column("bullet", width=350)
        self.validator_tree.column("score", width=60, anchor=tk.CENTER)
        self.validator_tree.column("grade", width=60, anchor=tk.CENTER)
        self.validator_tree.column("components", width=120, anchor=tk.CENTER)
        
        self.validator_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Bind selection to show suggestions
        self.validator_tree.bind('<<TreeviewSelect>>', self.on_bullet_selected)
        
        # Suggestions panel
        suggestions_card = ttk.LabelFrame(right_panel, text="Improvement Suggestions", padding="12")
        suggestions_card.pack(fill=tk.BOTH, expand=True, pady=(12, 0))
        
        self.validator_suggestions_text = scrolledtext.ScrolledText(
            suggestions_card,
            height=8,
            font=self.font_small,
            wrap=tk.WORD,
            bg=self.subtle_bg,
            fg=self.text_color,
            relief='solid',
            borderwidth=1,
            state='disabled'
        )
        self.validator_suggestions_text.pack(fill=tk.BOTH, expand=True)
        
        # Store validation results
        self.validation_results = []
    
    def show_framework_info(self):
        """Display information about available frameworks."""
        info_text = """
FRAMEWORK GUIDE:

📌 STAR (Situation, Task, Action, Result)
   Best for: All roles, behavioral interviews
   Structure: Context → Challenge → Your action → Measurable outcome

📌 CAR (Challenge, Action, Result)
   Best for: Problem-solving roles
   Structure: Problem faced → Your solution → Impact achieved

📌 PAR (Problem, Action, Result)
   Best for: Operations, process improvement
   Structure: Issue identified → Steps taken → Improvement made

📌 WHO (What, How, Outcome)
   Best for: Entry-level, straightforward accomplishments
   Structure: What you did → How you did it → What happened

📌 LPS (Location, Problem, Solution)
   Best for: Geographic/contextual roles
   Structure: Where → What was wrong → How you fixed it

SCORING COMPONENTS:
✓ Action Verb (30 pts): Strong, specific verb
✓ Metric (30 pts): Numbers, %, $, measurable data
✓ Context (20 pts): Situation or scope
✓ Result (20 pts): Outcome or impact

GRADING SCALE:
A+ (90-100): Exceptional - ready to submit
A  (80-89):  Strong - minor tweaks possible
B  (70-79):  Good - needs some improvement
C  (60-69):  Acceptable - needs significant work
D  (50-59):  Weak - major revision needed
F  (<50):    Poor - complete rewrite required
"""
        messagebox.showinfo("Framework Information", info_text)
    
    def validate_bullets(self):
        """Validate the entered resume bullets."""
        # Get input text
        input_text = self.validator_input_text.get('1.0', tk.END).strip()
        
        if not input_text:
            messagebox.showwarning("No Input", "Please enter at least one resume bullet to validate.")
            return
        
        # Split into individual bullets (by newline)
        bullets = [b.strip() for b in input_text.split('\n') if b.strip()]
        
        if not bullets:
            messagebox.showwarning("No Bullets", "Please enter valid resume bullets (one per line).")
            return
        
        # Get selected framework
        framework = self.validator_framework.get()
        
        # Create validator
        validator = FrameworkValidator()
        
        # Clear previous results
        for item in self.validator_tree.get_children():
            self.validator_tree.delete(item)
        self.validation_results = []
        
        # Validate each bullet
        try:
            results = validator.bulk_validate(bullets, framework)
            self.validation_results = results
            
            # Calculate summary stats
            total = len(results)
            avg_score = sum(r.score for r in results) / total if total > 0 else 0
            strong = sum(1 for r in results if r.score >= 80)
            weak = sum(1 for r in results if r.score < 60)
            
            # Update summary label
            grade = self._score_to_grade(avg_score)
            self.validator_summary_label.config(
                text=f"Validated {total} bullets | Average: {avg_score:.1f}/100 (Grade {grade}) | Strong: {strong} | Weak: {weak}"
            )
            
            # Populate tree with color coding
            for result in results:
                # Truncate bullet text for display
                bullet_display = result.bullet[:60] + "..." if len(result.bullet) > 60 else result.bullet
                
                # Component indicators
                components = ""
                components += "✓" if result.has_action else "✗"
                components += "✓" if result.has_metric else "✗"
                components += "✓" if result.has_context else "✗"
                components += "✓" if result.has_result else "✗"
                
                # Insert into tree
                item_id = self.validator_tree.insert(
                    "",
                    tk.END,
                    values=(bullet_display, result.score, result.grade, components)
                )
                
                # Color code by score
                if result.score >= 80:
                    self.validator_tree.item(item_id, tags=('excellent',))
                elif result.score >= 60:
                    self.validator_tree.item(item_id, tags=('good',))
                else:
                    self.validator_tree.item(item_id, tags=('poor',))
            
            # Configure tag colors
            self.validator_tree.tag_configure('excellent', background='#d1fae5')  # Light green
            self.validator_tree.tag_configure('good', background='#fef3c7')      # Light yellow
            self.validator_tree.tag_configure('poor', background='#fee2e2')      # Light red
            
            # Auto-select first item
            if self.validator_tree.get_children():
                first_item = self.validator_tree.get_children()[0]
                self.validator_tree.selection_set(first_item)
                self.validator_tree.focus(first_item)
                self.on_bullet_selected(None)
                
        except Exception as e:
            messagebox.showerror("Validation Error", f"Error validating bullets:\n{str(e)}")
    
    def on_bullet_selected(self, event):
        """Show suggestions when a bullet is selected in the tree."""
        selection = self.validator_tree.selection()
        if not selection:
            return
        
        # Get index of selected item
        item_id = selection[0]
        item_index = self.validator_tree.index(item_id)
        
        if item_index >= len(self.validation_results):
            return
        
        result = self.validation_results[item_index]
        
        # Build suggestions text
        suggestions_text = f"BULLET: {result.bullet}\n\n"
        suggestions_text += f"FRAMEWORK: {result.framework}\n"
        suggestions_text += f"SCORE: {result.score}/100 (Grade {result.grade})\n\n"
        
        suggestions_text += "COMPONENTS:\n"
        suggestions_text += f"  {'✓' if result.has_action else '✗'} Action Verb: {'Found' if result.has_action else 'MISSING - Use strong action verb (Led, Achieved, Improved)'}\n"
        suggestions_text += f"  {'✓' if result.has_metric else '✗'} Metric: {'Found' if result.has_metric else 'MISSING - Add numbers, %, $, or measurable data'}\n"
        suggestions_text += f"  {'✓' if result.has_context else '✗'} Context: {'Found' if result.has_context else 'MISSING - Add situation or scope'}\n"
        suggestions_text += f"  {'✓' if result.has_result else '✗'} Result: {'Found' if result.has_result else 'MISSING - Add outcome (resulting in, achieving, leading to)'}\n\n"
        
        if result.suggestions:
            suggestions_text += "SUGGESTIONS:\n"
            for i, suggestion in enumerate(result.suggestions, 1):
                suggestions_text += f"  {i}. {suggestion}\n"
        else:
            suggestions_text += "✨ Excellent! This bullet meets all criteria.\n"
        
        # Update suggestions display
        self.validator_suggestions_text.config(state='normal')
        self.validator_suggestions_text.delete('1.0', tk.END)
        self.validator_suggestions_text.insert('1.0', suggestions_text)
        self.validator_suggestions_text.config(state='disabled')
    
    def _score_to_grade(self, score):
        """Convert numeric score to letter grade."""
        if score >= 90:
            return "A+"
        elif score >= 80:
            return "A"
        elif score >= 70:
            return "B"
        elif score >= 60:
            return "C"
        elif score >= 50:
            return "D"
        else:
            return "F"
    
    # === PROFILE & SETTINGS TAB ===
    def create_profile_tab(self):
        """Create unified profile editor, settings, and baseline resume generator."""
        # Create canvas with scrollbar
        canvas = tk.Canvas(self.profile_tab)
        scrollbar = ttk.Scrollbar(self.profile_tab, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Enable mousewheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        container = ttk.Frame(scrollable_frame, padding="16")
        container.pack(fill=tk.BOTH, expand=True)

        # Contact/Candidate form section
        form_frame = ttk.LabelFrame(container, text="Contact & Candidate Info", padding="10")
        form_frame.grid(row=0, column=0, columnspan=2, sticky=tk.EW, pady=(0,10))

        labels = [
            ("Name", 'name'),
            ("Email", 'email'),
            ("Phone", 'phone'),
            ("Address", 'address'),
        ]
        self.profile_entries = {}
        for i,(label,key) in enumerate(labels):
            ttk.Label(form_frame, text=label+":").grid(row=i, column=0, sticky=tk.W, pady=2)
            ent = ttk.Entry(form_frame, width=50)
            ent.grid(row=i, column=1, sticky=tk.W, pady=2)
            self.profile_entries[key] = ent

        # Summary multiline
        ttk.Label(form_frame, text="Summary:").grid(row=len(labels), column=0, sticky=tk.NW, pady=2)
        self.summary_text = tk.Text(form_frame, width=50, height=4, font=("Consolas",9))
        self.summary_text.grid(row=len(labels), column=1, sticky=tk.W, pady=2)

        # Skills / Technologies / Achievements
        ttk.Label(form_frame, text="Skills (comma separated):").grid(row=len(labels)+1, column=0, sticky=tk.NW, pady=2)
        skills_frame = ttk.Frame(form_frame)
        skills_frame.grid(row=len(labels)+1, column=1, sticky=tk.W, pady=2)
        self.skills_entry = tk.Text(skills_frame, width=50, height=3, font=("Consolas",9), wrap=tk.WORD)
        self.skills_entry.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        skills_scroll = ttk.Scrollbar(skills_frame, orient=tk.VERTICAL, command=self.skills_entry.yview)
        skills_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.skills_entry.configure(yscrollcommand=skills_scroll.set)

        ttk.Label(form_frame, text="Technologies (comma separated):").grid(row=len(labels)+2, column=0, sticky=tk.NW, pady=2)
        tech_frame = ttk.Frame(form_frame)
        tech_frame.grid(row=len(labels)+2, column=1, sticky=tk.W, pady=2)
        self.tech_entry = tk.Text(tech_frame, width=50, height=3, font=("Consolas",9), wrap=tk.WORD)
        self.tech_entry.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tech_scroll = ttk.Scrollbar(tech_frame, orient=tk.VERTICAL, command=self.tech_entry.yview)
        tech_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tech_entry.configure(yscrollcommand=tech_scroll.set)

        ttk.Label(form_frame, text="Achievements (one per line):").grid(row=len(labels)+3, column=0, sticky=tk.NW, pady=2)
        self.achievements_text = tk.Text(form_frame, width=50, height=4, font=("Consolas",9))
        self.achievements_text.grid(row=len(labels)+3, column=1, sticky=tk.W, pady=2)

        # JSON advanced editor (optional)
        adv_frame = ttk.LabelFrame(container, text="Advanced JSON (optional)", padding="10")
        adv_frame.grid(row=1, column=0, columnspan=2, sticky=tk.NSEW, pady=(0,10))

        self.profile_text = tk.Text(adv_frame, width=90, height=16, font=("Consolas", 9), wrap=tk.NONE)
        self.profile_text.grid(row=0, column=0, sticky=tk.NSEW)
        y_scroll = ttk.Scrollbar(adv_frame, orient=tk.VERTICAL, command=self.profile_text.yview)
        y_scroll.grid(row=0, column=1, sticky=tk.NS)
        self.profile_text.configure(yscrollcommand=y_scroll.set)
        adv_frame.columnconfigure(0, weight=1)
        adv_frame.rowconfigure(0, weight=1)

        # Buttons
        btn_frame = ttk.Frame(container)
        btn_frame.grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(0,10))
        ttk.Button(btn_frame, text="Load Profile", command=self.load_profile_into_editor).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Auto-Fill", command=self.auto_fill_from_files).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Onboarding Wizard", command=self.open_onboarding_wizard).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Save Profile", command=self.save_profile_from_editor, style="Accent.TButton").pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Generate Standard Resume", command=self.generate_standard_resume).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Import Resume to Profile", command=self.import_resume_into_profile).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Add Reference Resumes", command=self.add_reference_resumes).pack(side=tk.LEFT, padx=4)

        # Reference Resumes Section
        ref_frame = ttk.LabelFrame(container, text="Reference Resume Documents", padding="10")
        ref_frame.grid(row=3, column=0, columnspan=2, sticky=tk.EW, pady=(0,10))
        
        ref_info = ttk.Label(ref_frame, text="Reference resumes are used as templates and examples when generating new resumes.", 
                            font=("Segoe UI", 9), foreground="#666")
        ref_info.pack(anchor=tk.W, pady=(0,5))
        
        self.ref_listbox = tk.Listbox(ref_frame, height=4, font=("Segoe UI", 9))
        self.ref_listbox.pack(fill=tk.X, pady=5)
        
        ref_btn_frame = ttk.Frame(ref_frame)
        ref_btn_frame.pack(fill=tk.X)
        ttk.Button(ref_btn_frame, text="Add References", command=self.add_reference_resumes).pack(side=tk.LEFT, padx=2)
        ttk.Button(ref_btn_frame, text="Remove Selected", command=self.remove_reference_resume).pack(side=tk.LEFT, padx=2)
        ttk.Button(ref_btn_frame, text="Clear All", command=self.clear_reference_resumes).pack(side=tk.LEFT, padx=2)
        
        # Settings Section
        settings_frame = ttk.LabelFrame(container, text="Application Settings", padding="10")
        settings_frame.grid(row=4, column=0, columnspan=2, sticky=tk.EW, pady=(0,10))
        
        # Profile & Output configuration in grid
        ttk.Label(settings_frame, text="Profile File:").grid(row=0, column=0, sticky=tk.W, pady=5, padx=(0,10))
        self.profile_path_var = tk.StringVar(value=str(ROOT / 'profile.yaml'))
        profile_entry = ttk.Entry(settings_frame, textvariable=self.profile_path_var, width=40)
        profile_entry.grid(row=0, column=1, pady=5)
        ttk.Button(settings_frame, text="Browse", command=self.browse_profile).grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Label(settings_frame, text="Output Folder:").grid(row=1, column=0, sticky=tk.W, pady=5, padx=(0,10))
        self.output_path_var = tk.StringVar(value=str(ROOT / 'outputs'))
        output_entry = ttk.Entry(settings_frame, textvariable=self.output_path_var, width=40)
        output_entry.grid(row=1, column=1, pady=5)
        ttk.Button(settings_frame, text="Browse", command=self.browse_output).grid(row=1, column=2, padx=5, pady=5)

        # About Section
        about_frame = ttk.LabelFrame(container, text="About", padding="10")
        about_frame.grid(row=5, column=0, columnspan=2, sticky=tk.EW)
        
        about_text = """Job Application Automation System v1.0

This tool generates customized resumes and cover letters based on job descriptions.
It extracts keywords from the JD and creates tailored application materials.

Features:
• JD-driven resume generation
• Keyword-optimized cover letters
• Organized output by company/position
• Automated application tracking
"""
        about_label = ttk.Label(about_frame, text=about_text, justify=tk.LEFT)
        about_label.pack()

        container.columnconfigure(0, weight=1)
        container.rowconfigure(1, weight=1)
        # Attempt silent auto-fill shortly after UI builds (does not overwrite existing values)
        self.root.after(800, self._auto_fill_on_start)

    def _auto_fill_on_start(self):
        """Populate profile form fields from existing data files on initial load if they are empty.
        Does NOT overwrite any already filled field (user or previous session)."""
        try:
            # Only run if name/email empty (heuristic for first load)
            if self.profile_entries['name'].get().strip() or self.profile_entries['email'].get().strip():
                return
            self.auto_fill_from_files(silent=True)
        except Exception:
            pass

    def auto_fill_from_files(self, silent: bool = False):
        """Auto-fill form fields using data in /data/profile_*.json without touching advanced JSON editor.
        Only fills fields that are currently blank. Reuses parsing logic similar to import, but without conflict prompts.
        Parameters:
            silent: when True, suppresses user-facing message boxes (used for startup)."""
        import json
        from pathlib import Path
        data_dir = Path(self.root_dir, 'data')
        files = {
            'contact': data_dir / 'profile_contact.json',
            'candidate': data_dir / 'profile_candidate.json',
            'experience': data_dir / 'profile_experience.json',
            'education': data_dir / 'profile_education.json'
        }
        loaded = {}
        for key, fp in files.items():
            try:
                if fp.exists():
                    loaded[key] = json.loads(fp.read_text(encoding='utf-8'))
                else:
                    loaded[key] = {} if key != 'experience' else []
            except Exception:
                loaded[key] = {} if key != 'experience' else []
        contact = loaded.get('contact', {})
        candidate = loaded.get('candidate', {})
        # Simple entries
        mapping = {
            'name': contact.get('name') or candidate.get('name'),
            'email': contact.get('email'),
            'phone': contact.get('phone'),
            'address': contact.get('address') or contact.get('location_statement')
        }
        filled_any = False
        for field, value in mapping.items():
            if value and not self.profile_entries[field].get().strip():
                self.profile_entries[field].delete(0, tk.END)
                self.profile_entries[field].insert(0, value)
                filled_any = True
        # Summary
        summary_value = candidate.get('summary')
        if summary_value and not self.summary_text.get('1.0', tk.END).strip():
            self.summary_text.delete('1.0', tk.END)
            self.summary_text.insert('1.0', summary_value)
            filled_any = True
        # Skills / technologies / achievements
        def _populate_text_box(widget, items):
            nonlocal filled_any
            if items and not widget.get('1.0', tk.END).strip():
                widget.delete('1.0', tk.END)
                widget.insert('1.0', ', '.join(items) if widget is not self.achievements_text else '\n'.join(items))
                filled_any = True
        skills_list = candidate.get('skills', [])
        tech_list = candidate.get('technologies', [])
        ach_list = candidate.get('achievements', [])
        _populate_text_box(self.skills_entry, skills_list)
        _populate_text_box(self.tech_entry, tech_list)
        _populate_text_box(self.achievements_text, ach_list)
        if filled_any and not silent:
            messagebox.showinfo('Auto-Fill', 'Profile fields populated from existing data files.')
        self.update_status('Auto-Fill complete' if filled_any else 'Auto-Fill found no blank fields to populate')

    # === SIMPLE MODE ===
    def _load_preferences(self):
        import json
        try:
            if self._preferences_path.exists():
                data = json.loads(self._preferences_path.read_text(encoding='utf-8'))
                self.simple_mode = bool(data.get('simple_mode', False))
                self.current_theme = data.get('theme', 'light') if data.get('theme') in self._palettes else 'light'
                self._apply_palette()  # palette might change before UI finished
        except Exception:
            self.simple_mode = False

    def _save_preferences(self):
        import json
        try:
            self._preferences_path.parent.mkdir(exist_ok=True)
            self._preferences_path.write_text(json.dumps({'simple_mode': self.simple_mode, 'theme': self.current_theme}, indent=2), encoding='utf-8')
        except Exception:
            pass

    def _apply_palette(self):
        p = self._palettes[self.current_theme]
        self.bg_color = p['bg']; self.accent_color = p['accent']; self.accent_dark = p['accent_dark']; self.accent_light = p['accent_light']
        self.success_color = p['success']; self.warning_color = p['warning']; self.text_color = p['text']; self.text_muted = p['muted']
        self.subtle_bg = p['subtle']; self.border_color = p['border']; self.highlight_color = p['highlight']
        # If style already initialized update key elements (called during runtime theme change)
        try:
            self.root.configure(bg=self.bg_color)
            self.style.configure('TFrame', background=self.bg_color)
            self.style.configure('Card.TFrame', background=self.subtle_bg)
            self.style.configure('TLabel', background=self.bg_color, foreground=self.text_color)
            self.style.configure('Heading.TLabel', foreground=self.text_color)
            self.style.configure('Subheading.TLabel', foreground=self.text_color)
            self.style.configure('Muted.TLabel', foreground=self.text_muted)
            self.style.configure('Accent.TButton', background=self.accent_color, foreground='white')
        except Exception:
            pass

    def cycle_theme(self):
        order = ['light','dark','contrast']
        idx = order.index(self.current_theme)
        self.current_theme = order[(idx+1) % len(order)]
        self._apply_palette()
        self.theme_btn.config(text=f'Theme: {self.current_theme.title()}')
        self._save_preferences()
        self.update_status(f"Theme switched to {self.current_theme}")

    def toggle_simple_mode(self):
        self.simple_mode = not self.simple_mode
        self._save_preferences()
        if self.simple_mode:
            self._apply_simple_mode()
            self.update_status('Simple Mode enabled')
        else:
            self._restore_advanced_tabs()
            self.update_status('Simple Mode disabled')
        self.simple_btn.config(text=f'Simple Mode: {"ON" if self.simple_mode else "OFF"}')

    def _apply_simple_mode(self):
        # Tabs to hide for simplicity
        to_hide = [self.tailor_tab, self.ats_tab, self.templates_tab, self.versions_tab, self.skills_tab]
        # Ensure list is clean
        self._hidden_tabs = []
        for tab in to_hide:
            try:
                idx = self.notebook.index(tab)
                self.notebook.forget(idx)
                self._hidden_tabs.append(tab)
            except Exception:
                pass
        # Also hide advanced JSON editor frame in profile tab if present
        # (We detect by label text 'Advanced JSON (optional)')
        try:
            for child in self.profile_tab.winfo_children():
                # Deep search label frames in scrollable canvas
                for sub in child.winfo_children():
                    if isinstance(sub, ttk.LabelFrame) and sub.cget('text').startswith('Advanced JSON'):
                        sub.grid_remove()
        except Exception:
            pass

    def _restore_advanced_tabs(self):
        # Re-add hidden tabs at end preserving order names
        for tab in self._hidden_tabs:
            try:
                if tab == self.tailor_tab:
                    self.notebook.add(tab, text='Optimize & Tailor')
                elif tab == self.ats_tab:
                    self.notebook.add(tab, text='ATS Scanner (Beta)')
                elif tab == self.templates_tab:
                    self.notebook.add(tab, text='Templates')
                elif tab == self.versions_tab:
                    self.notebook.add(tab, text='Versions')
                elif tab == self.skills_tab:
                    self.notebook.add(tab, text='Skills')
            except Exception:
                pass
        self._hidden_tabs = []
        # Restore advanced JSON editor
        try:
            for child in self.profile_tab.winfo_children():
                for sub in child.winfo_children():
                    if isinstance(sub, ttk.LabelFrame) and sub.cget('text').startswith('Advanced JSON'):
                        sub.grid()  # show again
        except Exception:
            pass

    # === ONBOARDING WIZARD ===
    def open_onboarding_wizard(self):
        """Launch a multi-step onboarding wizard to collect core profile data in a guided flow.
        Steps: 0 Contact | 1 Summary | 2 Experience Bullets | 3 Skills & Technologies | 4 Review & Save"""
        import json
        from pathlib import Path
        wiz = tk.Toplevel(self.root)
        wiz.title("Onboarding Wizard")
        wiz.geometry("640x520")
        wiz.transient(self.root)
        wiz.grab_set()

        state = {
            'step': 0,
            'contact': {'name':'','email':'','phone':'','address':'',},
            'summary': '',
            'experience_bullets': [],
            'skills': [],
            'technologies': [],
            'achievements': []
        }

        container = ttk.Frame(wiz, padding=12)
        container.pack(fill=tk.BOTH, expand=True)
        title_label = ttk.Label(container, text="Profile Onboarding", style='Subheading.TLabel')
        title_label.pack(anchor=tk.W)
        subtitle = ttk.Label(container, text="Guided setup – metrics matter: quantify impact (%, $, time saved, count).", style='Muted.TLabel')
        subtitle.pack(anchor=tk.W, pady=(0,8))

        progress = ttk.Progressbar(container, maximum=4, value=0)
        progress.pack(fill=tk.X, pady=(0,12))

        step_frame = ttk.Frame(container)
        step_frame.pack(fill=tk.BOTH, expand=True)

        nav_frame = ttk.Frame(container)
        nav_frame.pack(fill=tk.X)
        back_btn = ttk.Button(nav_frame, text="← Back")
        next_btn = ttk.Button(nav_frame, text="Next →")
        finish_btn = ttk.Button(nav_frame, text="Finish & Save", style='Accent.TButton')
        back_btn.pack(side=tk.LEFT)
        next_btn.pack(side=tk.RIGHT)

        def render_step():
            for child in step_frame.winfo_children():
                child.destroy()
            s = state['step']
            progress.configure(value=s)
            back_btn.configure(state='normal' if s > 0 else 'disabled')
            next_btn.configure(state='normal')
            finish_btn.pack_forget()
            if s == 0:
                ttk.Label(step_frame, text="Step 1: Contact", style='Subheading.TLabel').pack(anchor=tk.W)
                for key,label in [('name','Full Name'),('email','Email'),('phone','Phone'),('address','Location / Address')]:
                    frm = ttk.Frame(step_frame); frm.pack(fill=tk.X, pady=4)
                    ttk.Label(frm, text=label+':').pack(side=tk.LEFT)
                    ent = ttk.Entry(frm, width=40)
                    ent.pack(side=tk.LEFT, padx=6)
                    ent.insert(0, state['contact'].get(key,''))
                    ent.bind('<KeyRelease>', lambda e,k=key,en=ent: state['contact'].__setitem__(k,en.get().strip()))
            elif s == 1:
                ttk.Label(step_frame, text="Step 2: Professional Summary", style='Subheading.TLabel').pack(anchor=tk.W)
                ttk.Label(step_frame, text="Tip: Start with role focus + differentiation + quantified value.", style='Muted.TLabel').pack(anchor=tk.W)
                txt = tk.Text(step_frame, height=8, font=('Consolas',9))
                txt.pack(fill=tk.BOTH, expand=True, pady=6)
                txt.insert('1.0', state['summary'])
                def _upd(ev=None): state['summary'] = txt.get('1.0', tk.END).strip()
                txt.bind('<KeyRelease>', _upd)
            elif s == 2:
                ttk.Label(step_frame, text="Step 3: Experience Bullets", style='Subheading.TLabel').pack(anchor=tk.W)
                ttk.Label(step_frame, text="Add impactful bullets (one per line). Use metrics (%, $, time).", style='Muted.TLabel').pack(anchor=tk.W)
                txt = tk.Text(step_frame, height=10, font=('Consolas',9))
                txt.pack(fill=tk.BOTH, expand=True, pady=6)
                if state['experience_bullets']:
                    txt.insert('1.0', '\n'.join(state['experience_bullets']))
                def _upd(ev=None):
                    lines = [l.strip() for l in txt.get('1.0', tk.END).splitlines() if l.strip()]
                    state['experience_bullets'] = lines
                txt.bind('<KeyRelease>', _upd)
            elif s == 3:
                ttk.Label(step_frame, text="Step 4: Skills & Technologies", style='Subheading.TLabel').pack(anchor=tk.W)
                ttk.Label(step_frame, text="Comma separated lists. Group related items; keep concise.", style='Muted.TLabel').pack(anchor=tk.W)
                frm_sk = ttk.Frame(step_frame); frm_sk.pack(fill=tk.X, pady=4)
                ttk.Label(frm_sk, text='Skills:').pack(side=tk.LEFT)
                skills_ent = tk.Text(frm_sk, height=4, width=50, font=('Consolas',9))
                skills_ent.pack(side=tk.LEFT, padx=6)
                skills_ent.insert('1.0', ', '.join(state['skills']))
                frm_tech = ttk.Frame(step_frame); frm_tech.pack(fill=tk.X, pady=4)
                ttk.Label(frm_tech, text='Technologies:').pack(side=tk.LEFT)
                tech_ent = tk.Text(frm_tech, height=4, width=50, font=('Consolas',9))
                tech_ent.pack(side=tk.LEFT, padx=6)
                tech_ent.insert('1.0', ', '.join(state['technologies']))
                frm_ach = ttk.Frame(step_frame); frm_ach.pack(fill=tk.X, pady=4)
                ttk.Label(frm_ach, text='Achievements (one per line):').pack(side=tk.LEFT)
                ach_ent = tk.Text(frm_ach, height=4, width=50, font=('Consolas',9))
                ach_ent.pack(side=tk.LEFT, padx=6)
                ach_ent.insert('1.0', '\n'.join(state['achievements']))
                def _upd_all(ev=None):
                    state['skills'] = [s.strip() for s in skills_ent.get('1.0', tk.END).split(',') if s.strip()]
                    state['technologies'] = [t.strip() for t in tech_ent.get('1.0', tk.END).split(',') if t.strip()]
                    state['achievements'] = [a.strip() for a in ach_ent.get('1.0', tk.END).splitlines() if a.strip()]
                for w in (skills_ent, tech_ent, ach_ent):
                    w.bind('<KeyRelease>', _upd_all)
            elif s == 4:
                ttk.Label(step_frame, text="Review", style='Subheading.TLabel').pack(anchor=tk.W)
                data_preview = {
                    'contact': state['contact'],
                    'candidate': {
                        'summary': state['summary'],
                        'skills': state['skills'],
                        'technologies': state['technologies'],
                        'achievements': state['achievements']
                    },
                    'experience': [
                        {
                            'company': 'Experience Placeholder',
                            'title': 'Role Title',
                            'bullets': state['experience_bullets']
                        }
                    ] if state['experience_bullets'] else [],
                    'education': []
                }
                txt = tk.Text(step_frame, height=18, font=('Consolas',9))
                txt.pack(fill=tk.BOTH, expand=True)
                txt.insert('1.0', json.dumps(data_preview, indent=2))
                txt.config(state='disabled')
                next_btn.configure(state='disabled')
                finish_btn.pack(side=tk.RIGHT)
            else:
                return

        def go_next():
            if state['step'] < 4:
                state['step'] += 1
                render_step()
        def go_back():
            if state['step'] > 0:
                state['step'] -= 1
                render_step()
        def finish():
            # Persist to profile files
            data_dir = Path(self.root_dir, 'data')
            data_dir.mkdir(exist_ok=True)
            # Contact
            contact = {k:v for k,v in state['contact'].items() if v}
            contact['location_statement'] = contact.get('address','')
            (data_dir / 'profile_contact.json').write_text(json.dumps(contact, indent=2), encoding='utf-8')
            candidate = {
                'summary': state['summary'],
                'skills': state['skills'],
                'technologies': state['technologies'],
                'achievements': state['achievements']
            }
            (data_dir / 'profile_candidate.json').write_text(json.dumps(candidate, indent=2), encoding='utf-8')
            experience = []
            if state['experience_bullets']:
                experience.append({'company':'Experience Placeholder','title':'Role Title','bullets':state['experience_bullets']})
            (data_dir / 'profile_experience.json').write_text(json.dumps(experience, indent=2), encoding='utf-8')
            # Education left empty intentionally
            (data_dir / 'profile_education.json').write_text(json.dumps([], indent=2), encoding='utf-8')
            self.update_status('Onboarding complete')
            messagebox.showinfo('Onboarding', 'Profile saved. You can refine further in Profile tab.')
            try:
                wiz.destroy()
            except Exception:
                pass
            # Refresh current form fields (populate newly saved data without overwriting user custom edits)
            self.auto_fill_from_files(silent=True)

        next_btn.configure(command=go_next)
        back_btn.configure(command=go_back)
        finish_btn.configure(command=finish)
        render_step()

    def load_profile_into_editor(self):
        from pathlib import Path
        import json
        data_dir = Path(self.root_dir, 'data')
        files = {
            'contact': data_dir / 'profile_contact.json',
            'candidate': data_dir / 'profile_candidate.json',
            'experience': data_dir / 'profile_experience.json',
            'education': data_dir / 'profile_education.json'
        }
        merged = {}
        for key, fp in files.items():
            try:
                if fp.exists():
                    with open(fp, 'r', encoding='utf-8') as f:
                        merged[key] = json.load(f)
                else:
                    merged[key] = {} if key != 'experience' else []
            except Exception as e:
                merged[key] = {'error': str(e)}
        self.profile_text.delete('1.0', tk.END)
        self.profile_text.insert('1.0', json.dumps(merged, indent=2))
        self.update_status("Profile loaded for editing")

        # Populate form fields
        contact = merged.get('contact', {})
        candidate = merged.get('candidate', {})
        for key, ent in self.profile_entries.items():
            ent.delete(0, tk.END)
            # Use address field for location display
            if key == 'address':
                value = contact.get('address') or contact.get('location_statement', '')
            else:
                value = contact.get(key) or candidate.get(key, '')
            ent.insert(0, value)
        self.summary_text.delete('1.0', tk.END)
        self.summary_text.insert('1.0', candidate.get('summary',''))
        
        # Load skills - filter out corrupt/fragment entries
        raw_skills = candidate.get('skills', [])
        valid_skills = []
        for skill in raw_skills:
            skill_clean = skill.strip()
            # Skip fragments and corrupted entries
            if skill_clean.startswith('and ') or skill_clean.startswith('with '):
                continue
            if len(skill_clean) > 100:  # Skip overly long entries (likely corrupt)
                continue
            if '.' in skill_clean and len(skill_clean) > 50:  # Skip sentence fragments
                continue
            valid_skills.append(skill_clean)
        
        self.skills_entry.delete('1.0', tk.END)
        self.skills_entry.insert('1.0', ', '.join(valid_skills))
        
        self.tech_entry.delete('1.0', tk.END)
        self.tech_entry.insert('1.0', ', '.join(candidate.get('technologies', [])))
        self.achievements_text.delete('1.0', tk.END)
        self.achievements_text.insert('1.0', '\n'.join(candidate.get('achievements', [])))

    def save_profile_from_editor(self):
        import json, os
        from pathlib import Path
        
        data_dir = Path(self.root_dir) / 'data'
        data_dir.mkdir(parents=True, exist_ok=True)
        
        # Check if JSON editor has content
        raw = self.profile_text.get('1.0', tk.END).strip()
        
        if raw:
            # Save from JSON editor (advanced mode)
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as e:
                messagebox.showerror("Invalid JSON", f"Cannot parse JSON: {e}")
                return
            expected = ['contact','candidate','experience','education']
            for k in expected:
                if k not in data:
                    messagebox.showerror("Missing Key", f"Key '{k}' not found in root JSON object.")
                    return
        else:
            # Build data from form fields (simple mode)
            # Load existing data first
            data = {}
            for key in ['contact', 'candidate', 'experience', 'education']:
                file_path = data_dir / f'profile_{key}.json'
                if file_path.exists():
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data[key] = json.load(f)
                    except:
                        data[key] = {} if key != 'experience' else []
                else:
                    data[key] = {} if key != 'experience' else []
        
        # Save all files with form updates
        try:
            # Contact file
            with open(data_dir / 'profile_contact.json', 'w', encoding='utf-8') as f:
                address_value = self.profile_entries['address'].get().strip()
                contact_updates = {
                    'name': self.profile_entries['name'].get().strip(),
                    'email': self.profile_entries['email'].get().strip(),
                    'phone': self.profile_entries['phone'].get().strip(),
                    'address': address_value,
                    'location_statement': address_value,
                }
                if 'contact' not in data:
                    data['contact'] = {}
                data['contact'].update({k:v for k,v in contact_updates.items() if v})
                json.dump(data['contact'], f, indent=2)
            
            # Candidate file
            with open(data_dir / 'profile_candidate.json', 'w', encoding='utf-8') as f:
                candidate_updates = {
                    'summary': self.summary_text.get('1.0', tk.END).strip(),
                    'skills': [s.strip() for s in self.skills_entry.get('1.0', tk.END).strip().split(',') if s.strip()],
                    'technologies': [t.strip() for t in self.tech_entry.get('1.0', tk.END).strip().split(',') if t.strip()],
                    'achievements': [a.strip() for a in self.achievements_text.get('1.0', tk.END).splitlines() if a.strip()],
                }
                if 'candidate' not in data:
                    data['candidate'] = {}
                data['candidate'].update(candidate_updates)
                json.dump(data['candidate'], f, indent=2)
            
            # Experience file
            with open(data_dir / 'profile_experience.json', 'w', encoding='utf-8') as f:
                if 'experience' not in data:
                    data['experience'] = []
                json.dump(data['experience'], f, indent=2)
            
            # Education file
            with open(data_dir / 'profile_education.json', 'w', encoding='utf-8') as f:
                if 'education' not in data:
                    data['education'] = {}
                json.dump(data['education'], f, indent=2)
                
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save: {e}")
            return
        
        self.update_status("Profile saved successfully")
        messagebox.showinfo("Saved", "Profile saved successfully!\n\nContact, candidate info, experience, and education files updated.")

    def generate_standard_resume(self):
        from pathlib import Path
        out_dir = Path(self.default_output_dir, 'baseline')
        out_dir.mkdir(parents=True, exist_ok=True)
        # If we have reference resumes, build combined reference resume instead of simple baseline
        if getattr(self, 'reference_resumes', []):
            out_file = out_dir / 'combined_reference_resume.docx'
            try:
                self._combine_reference_resumes(str(out_file))
                # Open after creation (Windows)
                try:
                    import os
                    os.startfile(str(out_file))  # type: ignore
                except Exception:
                    import webbrowser
                    webbrowser.open(str(out_file))
            except Exception as e:
                messagebox.showerror("Generation Error", f"Failed combining reference resumes: {e}")
        else:
            out_file = out_dir / 'standard_resume.docx'
            try:
                self._generate_standard_resume_inline(str(out_file))
                try:
                    import os
                    os.startfile(str(out_file))  # type: ignore
                except Exception:
                    import webbrowser
                    webbrowser.open(str(out_file))
            except Exception as e:
                messagebox.showerror("Generation Error", str(e))

    def _combine_reference_resumes(self, out_path: str):
        """Combine all reference resumes into one organized DOCX.
        Sections: Header, Professional Summary (merged), Core Competencies, Achievements, Experience, Education.
        Parsing is heuristic; duplicates removed; metrics prioritized."""
        import json, os, re
        from pathlib import Path
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        ref_file = Path(self.root_dir) / 'data' / 'reference_resumes.json'
        if not ref_file.exists():
            raise RuntimeError("No reference resume list found.")
        try:
            with open(ref_file, 'r', encoding='utf-8') as f:
                refs = json.load(f)
        except Exception as e:
            raise RuntimeError(f"Cannot load reference resume list: {e}")
        if not refs:
            raise RuntimeError("Reference resume list is empty.")

        # Aggregation containers
        summaries = []
        skills = set()
        achievements = []
        experience_entries = []  # each: {company,title,dates,location,bullets:[]}
        education_entries = []
        contact_info = {}

        # Simple func to add metrics earlier
        def prioritize_metric(items):
            metric = []
            non = []
            for it in items:
                if re.search(r"\d+%|\d+x|\b\d{2,}\b|\b(\$\d)" , it):
                    metric.append(it)
                else:
                    non.append(it)
            return metric + non

        def dedupe_preserve_order(seq):
            seen = set(); out = []
            for s in seq:
                key = s.strip().lower()
                if key in seen or len(key) < 3:
                    continue
                seen.add(key); out.append(s.strip())
            return out

        # Basic parsing for each reference resume
        for ref in refs:
            p = Path(ref)
            if not p.exists():
                continue
            text_blocks = []
            try:
                if p.suffix.lower() == '.docx':
                    d = Document(str(p))
                    for para in d.paragraphs:
                        t = para.text.strip()
                        if t:
                            text_blocks.append(t)
                elif p.suffix.lower() in ('.txt',):
                    with open(p, 'r', encoding='utf-8', errors='ignore') as f:
                        for line in f:
                            t = line.strip()
                            if t:
                                text_blocks.append(t)
                else:
                    # Skip unsupported types (pdf etc.)
                    continue
            except Exception:
                continue

            # Heuristic extraction
            for line in text_blocks:
                low = line.lower()
                # Contact info
                if 'linkedin' in low or '@' in low or re.search(r"\+?\d[\d\s\-]{7,}\d", line):
                    if '@' in low and 'email' not in contact_info:
                        contact_info['email'] = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line).group(0) if re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line) else contact_info.get('email')
                    phone_match = re.search(r"(\+?\d{1,2}[\s-])?(\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4})", line)
                    if phone_match and 'phone' not in contact_info:
                        contact_info['phone'] = phone_match.group(0)
                    if 'linkedin' in low and 'linkedin' not in contact_info:
                        contact_info['linkedin'] = line
                # Summary indicator
                if any(h in low for h in ['professional summary','summary','profile']) and len(line) < 40:
                    continue  # heading
                elif len(line) > 60 and low.count(' ') > 8 and not line.startswith(('•','-')):
                    # treat long paragraph early as summary candidate
                    summaries.append(line)
                # Skills list
                if any(k in low for k in ['skills', 'competencies']) and ',' in line:
                    parts = [p.strip() for p in line.split(',') if p.strip()]
                    for pt in parts:
                        if len(pt) < 3: continue
                        skills.add(pt.title())
                # Bullet / achievement
                if line.startswith(('•','-','–','*')):
                    cleaned = line.lstrip('•-*– ').strip()
                    if re.search(r"\d+%|\b\d{2,}\b", cleaned):
                        achievements.append(cleaned)
                    # Experience bullet capture (will assign later)
                    experience_entries.append({'bullet_only': cleaned})
                # Education
                if any(word in low for word in ['bachelor','master','university','degree']) and len(line) < 120:
                    education_entries.append(line)

        # Organize experience bullets into a single entry (simplified)
        exp_bullets = [e['bullet_only'] for e in experience_entries if 'bullet_only' in e]
        exp_bullets = dedupe_preserve_order(exp_bullets)[:30]

        # Prepare document
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)

        # === Deduplication / compression phase ===
        def _dedupe(seq):
            seen = set(); out_local = []
            for s in seq:
                key = re.sub(r"\s+", " ", s.strip().lower())
                if key in seen:
                    continue
                seen.add(key); out_local.append(s.strip())
            return out_local
        def _compress_phrases(lines):
            trimmed = []
            for ln in lines:
                ln_clean = re.sub(r"^(Led |Managed |Responsible for |Accountable for |Oversaw |Directed )", "", ln, flags=re.I).strip()
                trimmed.append(ln_clean if ln_clean else ln)
            return trimmed
        summaries = _dedupe(summaries)
        achievements = _dedupe(achievements)
        exp_bullets = _dedupe(exp_bullets)
        exp_bullets = _compress_phrases(exp_bullets)
        achievements = _compress_phrases(achievements)

        # Header
        name = contact_info.get('name', 'Ariel Karagodskiy')
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name); r.bold = True; r.font.size = Pt(18)
        info_line = ' • '.join(filter(None,[contact_info.get('phone'), contact_info.get('email'), contact_info.get('linkedin')]))
        if info_line:
            ip = doc.add_paragraph(info_line); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ip.runs: run.font.size = Pt(9)

        # Summary section
        merged_summary = ' '.join(dedupe_preserve_order(summaries))
        # Remove sentence-level duplicates in summary
        if merged_summary:
            sent_parts = re.split(r"(?<=[.!?])\s+", merged_summary)
            sent_clean = []
            seen_sent = set()
            for s in sent_parts:
                k = s.strip().lower()
                if len(k) < 3 or k in seen_sent:
                    continue
                seen_sent.add(k); sent_clean.append(s.strip())
            merged_summary = ' '.join(sent_clean)
        if merged_summary:
            sh = doc.add_paragraph(); sr = sh.add_run('PROFESSIONAL SUMMARY'); sr.bold = True; sr.font.size = Pt(11)
            sp = doc.add_paragraph(merged_summary[:1200])
            for run in sp.runs: run.font.size = Pt(10)

        # Skills section
        if skills:
            sk = doc.add_paragraph(); skr = sk.add_run('CORE COMPETENCIES'); skr.bold=True; skr.font.size=Pt(11)
            skill_list = dedupe_preserve_order(prioritize_metric(list(skills)))[:40]
            sp = doc.add_paragraph(' • '.join(skill_list))
            for run in sp.runs: run.font.size = Pt(10)

        # Achievements section
        ach_list = dedupe_preserve_order(prioritize_metric(achievements))[:10]
        if ach_list:
            ah = doc.add_paragraph(); ar = ah.add_run('KEY ACHIEVEMENTS'); ar.bold=True; ar.font.size=Pt(11)
            for a in ach_list:
                ap = doc.add_paragraph(style='List Bullet')
                ap.add_run(a).font.size = Pt(10)

        # Experience (flattened bullets only)
        if exp_bullets:
            eh = doc.add_paragraph(); er = eh.add_run('SELECTED EXPERIENCE IMPACT'); er.bold=True; er.font.size=Pt(11)
            for b in exp_bullets:
                bp = doc.add_paragraph(style='List Bullet')
                bp.add_run(b).font.size = Pt(10)

        # Education
        edu_clean = dedupe_preserve_order(education_entries)
        if edu_clean:
            edh = doc.add_paragraph(); edr = edh.add_run('EDUCATION'); edr.bold=True; edr.font.size=Pt(11)
            for e in edu_clean[:5]:
                ep = doc.add_paragraph(e)
                for run in ep.runs: run.font.size = Pt(10)

        try:
            doc.save(out_path)
            self.update_status('Combined reference resume generated')
            messagebox.showinfo('Resume Generated', f'Created combined reference resume:\n{out_path}')
        except Exception as e:
            raise RuntimeError(f"Cannot save combined resume: {e}")

    def _generate_standard_resume_inline(self, out_path):
        import json, os
        from docx import Document
        from docx.shared import Pt, Inches
        data_dir = os.path.join(self.root_dir, 'data')
        def lj(name, default):
            try:
                pth = os.path.join(data_dir, name)
                if os.path.exists(pth):
                    with open(pth,'r',encoding='utf-8') as f:
                        return json.load(f)
            except Exception:
                pass
            return default
        contact = lj('profile_contact.json', {})
        candidate = lj('profile_candidate.json', {})
        experience = lj('profile_experience.json', [])
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(0.6)
        name = contact.get('name','Candidate Name')
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(18)
        p.alignment = 1
        info_line = ' • '.join(filter(None,[contact.get('location_statement',''),contact.get('phone',''),contact.get('email','')]))
        ip = doc.add_paragraph(info_line)
        for run in ip.runs: run.font.size = Pt(9)
        sum_p = doc.add_paragraph(candidate.get('summary','Professional bringing comprehensive experience across operations, improvement, and technical execution.'))
        for run in sum_p.runs: run.font.size = Pt(10)
        skills = list(dict.fromkeys(candidate.get('skills',[]) + candidate.get('technologies',[])))
        # Deduplicate similar skill tokens (case-insensitive, strip variants)
        norm_map = {}
        compressed_skills = []
        for sk in skills:
            base = sk.strip()
            key = base.lower()
            if key in norm_map:
                continue
            norm_map[key] = base
            compressed_skills.append(base)
        # Remove overly long or repeated phrase-like entries
        filtered_skills = [s for s in compressed_skills if len(s) <= 40 and 'summary:' not in s.lower()]
        skills = filtered_skills
        if skills:
            sp = doc.add_paragraph('CORE SKILLS: ' + ' • '.join(skills))
            for run in sp.runs: run.font.size = Pt(10)
        for entry in experience:
            co = entry.get('company','')
            ti = entry.get('title','')
            exp_head = doc.add_paragraph()
            rh = exp_head.add_run(f"{co} | {ti}")
            rh.bold = True; rh.font.size = Pt(10)
            # Deduplicate bullets within each experience
            seen_b = set()
            for b in entry.get('bullets',[]):
                b_clean = b.strip()
                low = b_clean.lower()
                # Skip if excessively similar to previous (simple prefix match)
                if any(low.startswith(prev[:35].lower()) for prev in seen_b):
                    continue
                if low in seen_b or len(low) < 5:
                    continue
                seen_b.add(low)
                bp = doc.add_paragraph(style='List Bullet')
                br = bp.add_run(b_clean)
                br.font.size = Pt(10)
        try:
            doc.save(out_path)
            self.update_status("Baseline resume generated")
            messagebox.showinfo("Baseline Resume", f"Generated: {out_path}")
        except Exception as e:
            messagebox.showerror("Save Error", str(e))
    
    def import_resume_into_profile(self):
        """Import an existing resume (DOCX/TXT/PDF) with preview and confirmation.
        Shows parsed data before importing to prevent incorrect overwrites."""
        file_path = filedialog.askopenfilename(title="Select Resume", filetypes=[("Document","*.docx *.txt *.pdf")])
        if not file_path:
            return
        
        # Parse with error handling
        try:
            parsed = self._parse_resume_file(file_path)
        except Exception as e:
            messagebox.showerror("Parse Error", f"Failed to parse resume:\n\n{str(e)}\n\nPlease ensure the file is a valid DOCX or TXT format.")
            return
        
        if not parsed:
            messagebox.showwarning("Import Failed", "Could not extract any recognizable fields from this resume.\n\nThe parser looks for:\n• Email addresses\n• Phone numbers\n• Location (City, State)\n• Section headings (Summary, Skills, Achievements)")
            return
        
        # Build preview message
        preview = "Parsed the following data:\n\n"
        for key, val in parsed.items():
            if val:
                display_val = str(val)
                if isinstance(val, list):
                    display_val = '\n  • '.join([''] + val[:3])  # Show first 3 items
                    if len(val) > 3:
                        display_val += f"\n  ... and {len(val)-3} more"
                preview += f"• {key.replace('_', ' ').title()}: {display_val}\n"
        
        preview += "\n⚠️ IMPORTANT:\n"
        preview += "• Only EMPTY fields will be filled\n"
        preview += "• Existing data will NOT be overwritten\n"
        preview += "• Review carefully - location/address may be incorrect\n\n"
        preview += "Proceed with import?"
        
        # Confirmation dialog
        if not messagebox.askyesno("Confirm Import", preview):
            self.update_status("Import cancelled")
            return
        
        # Track what gets filled and what gets skipped
        filled = []
        skipped = []
        conflicts = []
        
        # Entries (simple Entry widgets)
        entry_map = {
            'name': 'name',
            'email': 'email',
            'phone': 'phone',
            'address': 'address'
        }
        
        for parsed_key, entry_key in entry_map.items():
            if parsed_key in parsed and parsed[parsed_key]:
                ent = self.profile_entries.get(entry_key)
                if ent:
                    current = ent.get().strip()
                    if not current:
                        ent.delete(0, tk.END)
                        ent.insert(0, parsed[parsed_key])
                        filled.append(f"{entry_key}: {parsed[parsed_key]}")
                    else:
                        # Check if values differ
                        if current != parsed[parsed_key]:
                            conflicts.append(f"{entry_key}: kept '{current}' (ignored '{parsed[parsed_key]}')")
                        else:
                            skipped.append(entry_key)
        
        # Summary
        if parsed.get('summary'):
            current_summary = self.summary_text.get('1.0', tk.END).strip()
            if not current_summary:
                self.summary_text.delete('1.0', tk.END)
                self.summary_text.insert('1.0', parsed['summary'])
                filled.append(f"summary: {parsed['summary'][:50]}...")
            elif current_summary != parsed['summary']:
                conflicts.append(f"summary: kept existing (ignored imported)")
            else:
                skipped.append('summary')
        
        # Skills
        if parsed.get('skills'):
            current_skills = self.skills_entry.get().strip()
            if not current_skills:
                self.skills_entry.delete(0, tk.END)
                self.skills_entry.insert(0, parsed['skills'])
                filled.append(f"skills: {parsed['skills'][:50]}...")
            elif current_skills != parsed['skills']:
                conflicts.append(f"skills: kept existing")
            else:
                skipped.append('skills')
        
        # Technologies
        if parsed.get('technologies'):
            current_tech = self.tech_entry.get().strip()
            if not current_tech:
                self.tech_entry.delete(0, tk.END)
                self.tech_entry.insert(0, parsed['technologies'])
                filled.append(f"technologies: {parsed['technologies'][:50]}...")
            elif current_tech != parsed['technologies']:
                conflicts.append(f"technologies: kept existing")
            else:
                skipped.append('technologies')
        
        # Achievements
        if parsed.get('achievements'):
            current_ach = self.achievements_text.get('1.0', tk.END).strip()
            if not current_ach:
                self.achievements_text.delete('1.0', tk.END)
                self.achievements_text.insert('1.0', '\n'.join(parsed['achievements']))
                filled.append(f"achievements: {len(parsed['achievements'])} items")
            elif current_ach != '\n'.join(parsed['achievements']):
                conflicts.append(f"achievements: kept existing")
            else:
                skipped.append('achievements')
        
        # Build result message
        result_msg = ""
        if filled:
            result_msg += f"✅ Imported {len(filled)} field(s):\n"
            for item in filled[:5]:  # Show first 5
                result_msg += f"  • {item}\n"
            if len(filled) > 5:
                result_msg += f"  ... and {len(filled)-5} more\n"
        
        if conflicts:
            result_msg += f"\n⚠️ Kept {len(conflicts)} existing value(s):\n"
            for item in conflicts[:3]:
                result_msg += f"  • {item}\n"
            if len(conflicts) > 3:
                result_msg += f"  ... and {len(conflicts)-3} more\n"
        
        if not filled and not conflicts:
            result_msg = "No changes made - all fields already populated with identical data."
        
        result_msg += "\n💾 Remember to click 'Save Profile' to persist changes!"
        
        messagebox.showinfo("Import Complete", result_msg)
        self.update_status(f"Imported {len(filled)} field(s) from resume")

    def _parse_resume_file(self, path):
        """Heuristic parser for resume (DOCX/TXT/PDF). Returns dict of extracted fields."""
        import re
        data = {}
        text = ""
        if path.lower().endswith('.docx'):
            try:
                from docx import Document
                doc = Document(path)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                text = "\n".join(paragraphs)
            except Exception as e:
                raise e
        elif path.lower().endswith('.pdf'):
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(path)
                parts = []
                for page in reader.pages:
                    try:
                        parts.append(page.extract_text() or '')
                    except Exception:
                        pass
                text = '\n'.join(parts)
            except Exception as e:
                raise e
        else:
            try:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception as e:
                raise e
        if not text:
            return data
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        # Email
        email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
        if email_match:
            data['email'] = email_match.group(0)
        # Phone
        phone_match = re.search(r"(\+?\d{1,2}[\s-])?(\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4})", text)
        if phone_match:
            data['phone'] = phone_match.group(0)
        # Name heuristic: first line with 2-4 capitalized words
        for line in lines[:6]:
            if re.match(r"^[A-Z][a-zA-Z]+(\s+[A-Z][a-zA-Z]+){1,3}$", line):
                data['name'] = line
                break
        # Address (City, ST)
        loc_match = re.search(r"\b([A-Z][a-zA-Z]+),\s*(AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|WA|WV|WI|WY)\b", text)
        if loc_match:
            data['address'] = loc_match.group(0)
        # Summary section detection
        summary_index = None
        for i, line in enumerate(lines):
            if re.match(r"^(SUMMARY|PROFESSIONAL SUMMARY)$", line.upper()):
                summary_index = i
                break
        if summary_index is not None:
            summary_lines = []
            for j in range(summary_index+1, len(lines)):
                nxt = lines[j]
                if not nxt or (nxt.isupper() and len(nxt.split()) <= 5):
                    break
                summary_lines.append(nxt)
            if summary_lines:
                data['summary'] = ' '.join(summary_lines)
        # Skills
        skills_tokens = []
        for i, line in enumerate(lines):
            if 'skill' in line.lower():
                # Collect same line tokens
                parts = re.split(r"[,;|]\s*", line)
                for p in parts:
                    if p and 'skill' not in p.lower():
                        skills_tokens.append(p.strip())
                # Next couple lines if they look like lists
                for k in range(i+1, min(i+4, len(lines))):
                    if re.search(r"[,;|]", lines[k]):
                        for p in re.split(r"[,;|]\s*", lines[k]):
                            if p:
                                skills_tokens.append(p.strip())
                    else:
                        break
                break
        if skills_tokens:
            uniq = sorted(set(skills_tokens), key=str.lower)
            data['skills'] = ', '.join(uniq)
        # Technologies (if separate heading)
        tech_tokens = []
        for i, line in enumerate(lines):
            if 'technolog' in line.lower():
                parts = re.split(r"[,;|]\s*", line)
                for p in parts:
                    if p and 'technolog' not in p.lower():
                        tech_tokens.append(p.strip())
                break
        if tech_tokens:
            data['technologies'] = ', '.join(sorted(set(tech_tokens), key=str.lower))
        # Achievements (collect bullet-like lines after heading)
        ach_index = None
        for i, line in enumerate(lines):
            if 'achievement' in line.lower() or 'accomplishment' in line.lower():
                ach_index = i
                break
        achievements = []
        if ach_index is not None:
            for j in range(ach_index+1, len(lines)):
                nxt = lines[j]
                if not nxt or (nxt.isupper() and len(nxt.split()) <= 5):
                    break
                if re.match(r"^[\-\u2022\*]", nxt) or len(nxt.split()) > 5:
                    cleaned = re.sub(r"^[\-\u2022\*\s]+", "", nxt)
                    achievements.append(cleaned)
        if achievements:
            data['achievements'] = achievements[:20]
        return data
    
    def load_jd_file(self):
        """Load job description from a text file"""
        file_path = filedialog.askopenfilename(
            title="Select Job Description File",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    jd_content = f.read()
                self.jd_text.delete('1.0', tk.END)
                self.jd_text.insert('1.0', jd_content)
                self.update_status(f"Loaded JD from {Path(file_path).name}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load file: {str(e)}")
    
    def clear_form(self):
        """Clear all form fields"""
        self.company_entry.delete(0, tk.END)
        self.position_entry.delete(0, tk.END)
        self.jd_text.delete('1.0', tk.END)
        self.update_status("Form cleared")
    
    def ai_auto_extract(self):
        """Use AI to automatically extract company and position from JD"""
        jd_text = self.jd_text.get('1.0', tk.END).strip()
        
        if not jd_text or len(jd_text) < 50:
            messagebox.showinfo("AI Extract", "Please paste a job description first (at least 50 characters).")
            return
        
        self.update_status("🤖 AI extracting job information...")
        
        try:
            # Use AI to extract information
            info = self.ai_assistant.extract_job_info(jd_text)
            
            # Update company field if found
            if info.get('company'):
                self.company_entry.delete(0, tk.END)
                self.company_entry.insert(0, info['company'])
            
            # Update position field if found
            if info.get('position'):
                self.position_entry.delete(0, tk.END)
                self.position_entry.insert(0, info['position'])
            
            # Build results message
            found = []
            if info.get('company'):
                found.append(f"Company: {info['company']}")
            if info.get('position'):
                found.append(f"Position: {info['position']}")
            if info.get('location'):
                found.append(f"Location: {info['location']}")
            if info.get('job_type'):
                found.append(f"Type: {info['job_type']}")
            if info.get('remote_status'):
                found.append(f"Remote: {info['remote_status']}")
            
            if found:
                result_msg = "AI found:\n\n" + "\n".join(found)
                messagebox.showinfo("AI Extract Complete", result_msg)
                self.update_status("✅ AI extraction complete")
            else:
                messagebox.showwarning("AI Extract", "Could not automatically extract company/position.\nPlease enter them manually.")
                self.update_status("⚠️ AI extraction found no clear matches")
        
        except Exception as e:
            messagebox.showerror("Error", f"AI extraction failed: {str(e)}")
            self.update_status("❌ AI extraction error")

    # === Automatic JD field inference (non-blocking) ===
    def _on_jd_key_activity(self, event=None):
        """Debounce key activity in JD box and schedule passive extraction.
        Only triggers if company/position still empty and JD sufficiently long."""
        if self._jd_auto_timer:
            self.root.after_cancel(self._jd_auto_timer)
            self._jd_auto_timer = None
        # Conditions to attempt inference
        jd_text = self.jd_text.get('1.0', tk.END).strip()
        if len(jd_text) < 80:
            return
        if self.company_entry.get().strip() and self.position_entry.get().strip():
            return  # already filled
        # Schedule inference after idle delay
        self._jd_auto_timer = self.root.after(900, self._passive_jd_infer)

    def _on_jd_focus_out(self, event=None):
        """Run inference when focus leaves JD box if fields still blank."""
        if not (self.company_entry.get().strip() and self.position_entry.get().strip()):
            self._passive_jd_infer()

    def _passive_jd_infer(self):
        """Attempt lightweight heuristic extraction; fall back to AI if still missing."""
        self._jd_auto_timer = None
        jd_text = self.jd_text.get('1.0', tk.END).strip()
        if not jd_text:
            return
        # Heuristic extraction
        inferred = self._infer_basic_jd_fields(jd_text)
        changed = False
        if inferred.get('company') and not self.company_entry.get().strip():
            self.company_entry.delete(0, tk.END)
            self.company_entry.insert(0, inferred['company'][:120])
            changed = True
        if inferred.get('position') and not self.position_entry.get().strip():
            self.position_entry.delete(0, tk.END)
            self.position_entry.insert(0, inferred['position'][:120])
            changed = True
        if changed:
            self.update_status('🔎 Auto-filled fields from pasted JD')
            return
        # If heuristics failed and text large, attempt AI (silent)
        if len(jd_text) > 300 and hasattr(self.ai_assistant, 'extract_job_info'):
            try:
                info = self.ai_assistant.extract_job_info(jd_text)
                ai_changed = False
                if info.get('company') and not self.company_entry.get().strip():
                    self.company_entry.delete(0, tk.END)
                    self.company_entry.insert(0, info['company'][:120])
                    ai_changed = True
                if info.get('position') and not self.position_entry.get().strip():
                    self.position_entry.delete(0, tk.END)
                    self.position_entry.insert(0, info['position'][:120])
                    ai_changed = True
                if ai_changed:
                    self.update_status('🤖 AI auto-filled fields from JD')
            except Exception:
                pass  # silent fail

    def _infer_basic_jd_fields(self, jd_text: str) -> dict:
        """Heuristic field inference without AI calls.
        Returns dict with possible 'company' and 'position'."""
        import re
        lines = [ln.strip() for ln in jd_text.splitlines() if ln.strip()]
        out = {}
        # Look for explicit labels first
        for ln in lines[:25]:
            low = ln.lower()
            if 'company:' in low:
                out['company'] = ln.split(':',1)[1].strip()
            if any(tag in low for tag in ['position:', 'title:','role:']):
                out['position'] = ln.split(':',1)[1].strip()
            if out.get('company') and out.get('position'):
                return out
        # Infer company from URL if present
        url_match = re.search(r'https?://([^/\s]+)', jd_text)
        if url_match and 'company' not in out:
            domain = url_match.group(1)
            parts = domain.split('.')
            if len(parts) >= 2:
                candidate = parts[-2] if parts[-1] in ('com','org','net','io','co') else parts[0]
                if len(candidate) > 2:
                    out['company'] = candidate.replace('-', ' ').title()
        # Position inference: first line containing common job nouns
        job_keywords = ['engineer','manager','analyst','specialist','director','consultant','developer','scientist','technician','coordinator','lead','architect','administrator']
        if 'position' not in out:
            for ln in lines[:40]:
                low = ln.lower()
                if any(k in low for k in job_keywords) and len(ln.split()) <= 15:
                    out['position'] = ln.strip(':- ')[:150]
                    break
        # Fallback: capitalized phrase at top
        if 'position' not in out and lines:
            first = lines[0]
            if re.match(r'[A-Z][A-Za-z/&\- ]{3,}', first) and len(first.split()) <= 8:
                out['position'] = first.strip()
        return out
    
    
    def submit_application(self):
        """Submit the application for processing"""
        company = self.company_entry.get().strip()
        position = self.position_entry.get().strip()
        jd_text = self.jd_text.get('1.0', tk.END).strip()
        
        # Validation
        if not company:
            messagebox.showwarning("Validation Error", "Please enter a company name.")
            return
        if not position:
            messagebox.showwarning("Validation Error", "Please enter a position title.")
            return
        if not jd_text or len(jd_text) < 50:
            messagebox.showwarning("Validation Error", "Please enter a valid job description (at least 50 characters).")
            return
        
        # Disable submit button
        self.update_status("Processing application...")
        
        # Run in separate thread to avoid freezing UI
        thread = threading.Thread(
            target=self.process_application,
            args=(company, position, jd_text)
        )
        thread.daemon = True
        thread.start()
    
    def process_application(self, company, position, jd_text):
        """Process the application in background thread"""
        try:
            # Call the automation system
            result = self.automation.apply_to_job(company, position, jd_text)
            
            # Update UI in main thread
            self.root.after(0, lambda: self.on_application_complete(result, company, position))
        except Exception as e:
            self.root.after(0, lambda: self.on_application_error(str(e)))
    
    def on_application_complete(self, result, company, position):
        """Handle successful application generation"""
        self.update_status("Application generated successfully!")
        
        message = f"✅ Application materials generated for {position} at {company}!\n\n"
        message += f"📁 Output folder: {result.get('output_folder', 'N/A')}\n\n"
        message += "Files created:\n"
        message += "  • Resume (DOCX)\n"
        message += "  • Cover Letter (DOCX)\n"
        message += "  • Job Description (TXT)\n"
        message += "  • Summary (JSON)\n\n"
        message += "Would you like to open the output folder?"
        
        if messagebox.askyesno("Success", message):
            output_folder = result.get('output_folder')
            if output_folder and Path(output_folder).exists():
                webbrowser.open(str(output_folder))
        
        # Clear form and refresh applications list
        self.clear_form()
        self.refresh_applications()
    
    def on_application_error(self, error_msg):
        """Handle application generation error"""
        self.update_status("Error occurred")
        messagebox.showerror("Error", f"Failed to generate application:\n\n{error_msg}")
    
    def calculate_match_score(self, jd_text, profile_skills):
        """
        Calculate estimated match score between job description and candidate profile.
        NOTE: This is an automated estimate only. Many factors beyond this algorithm 
        determine hiring decisions. Always apply if you're interested in the role!
        
        Returns: tuple (score 0-100, category 'high'/'medium'/'low')
        """
        if not jd_text:
            return 0, 'low'
        
        try:
            import json
            from pathlib import Path
            import re
            
            # Load profile data
            profile_skills_data = []
            profile_exp_data = []
            
            skills_path = ROOT / 'data' / 'profile_skills.json'
            exp_path = ROOT / 'data' / 'profile_experience.json'
            candidate_path = ROOT / 'data' / 'profile_candidate.json'
            
            if skills_path.exists():
                with open(skills_path, 'r', encoding='utf-8') as f:
                    skills_json = json.load(f)
                    profile_skills_data = skills_json.get('skills', []) + skills_json.get('technologies', [])
            
            if exp_path.exists():
                with open(exp_path, 'r', encoding='utf-8') as f:
                    profile_exp_data = json.load(f)
            
            candidate_years = 0
            if candidate_path.exists():
                with open(candidate_path, 'r', encoding='utf-8') as f:
                    cand = json.load(f)
                    candidate_years = cand.get('years_experience', 0)
            
            # Normalize text for comparison
            jd_lower = jd_text.lower()
            
            # Extract JD requirements more accurately
            # Look for required vs preferred skills
            required_section = re.search(r'required.*?(?=preferred|qualifications|responsibilities|\Z)', jd_lower, re.DOTALL)
            preferred_section = re.search(r'preferred.*?(?=responsibilities|duties|about|\Z)', jd_lower, re.DOTALL)
            
            required_text = required_section.group(0) if required_section else jd_lower
            preferred_text = preferred_section.group(0) if preferred_section else ""
            
            # Skills matching (35% weight) - more nuanced
            skill_score = 0
            if profile_skills_data:
                required_skill_matches = 0
                preferred_skill_matches = 0
                total_profile_skills = len(profile_skills_data)
                
                for skill in profile_skills_data:
                    skill_lower = skill.lower()
                    # Exact phrase matching and partial word matching
                    if skill_lower in required_text or any(word in required_text for word in skill_lower.split()):
                        required_skill_matches += 2  # Required skills worth more
                    elif skill_lower in preferred_text or any(word in preferred_text for word in skill_lower.split()):
                        preferred_skill_matches += 1
                
                # Calculate based on matches found
                max_possible = total_profile_skills * 2
                actual_matches = required_skill_matches + preferred_skill_matches
                skill_score = (actual_matches / max_possible) * 35 if max_possible > 0 else 0
            
            # Experience level matching (25% weight)
            exp_score = 0
            years_patterns = [
                r'(\d+)(?:-\d+)?\+?\s*years?\s+(?:of\s+)?experience',
                r'(\d+)(?:-\d+)?\s*(?:\+|or more)\s*years?',
                r'minimum\s+of\s+(\d+)\s+years?'
            ]
            
            required_years = None
            for pattern in years_patterns:
                match = re.search(pattern, required_text)
                if match:
                    required_years = int(match.group(1))
                    break
            
            if required_years:
                if candidate_years >= required_years:
                    exp_score = 25
                elif candidate_years >= required_years * 0.8:
                    exp_score = 20
                elif candidate_years >= required_years * 0.6:
                    exp_score = 15
                else:
                    exp_score = 10
            else:
                exp_score = 20  # No specific requirement, assume moderate fit
            
            # Keyword relevance in experience bullets (25% weight)
            keyword_score = 0
            if profile_exp_data:
                # Extract domain-specific keywords (technical terms, tools, methodologies)
                jd_words = set(re.findall(r'\b[A-Za-z][A-Za-z0-9+#\.]{2,}\b', jd_text))  # Keep case for acronyms
                common_words = {
                    'with', 'from', 'that', 'this', 'have', 'will', 'experience', 'work', 'team', 'using',
                    'able', 'strong', 'excellent', 'good', 'ability', 'skills', 'knowledge', 'understanding',
                    'years', 'including', 'such', 'other', 'must', 'should', 'required', 'preferred'
                }
                jd_keywords = {w.lower() for w in jd_words if w.lower() not in common_words and len(w) > 3}
                
                # Weighted keyword matching in bullets
                total_bullets = 0
                weighted_matches = 0
                
                for job in profile_exp_data:
                    for bullet in job.get('bullets', []):
                        total_bullets += 1
                        bullet_lower = bullet.lower()
                        
                        # Count how many keywords appear
                        matches_in_bullet = sum(1 for kw in jd_keywords if kw in bullet_lower)
                        
                        # Weighted: 1-2 keywords = partial, 3+ = strong match
                        if matches_in_bullet >= 3:
                            weighted_matches += 1.0
                        elif matches_in_bullet >= 1:
                            weighted_matches += 0.5
                
                if total_bullets > 0:
                    keyword_score = (weighted_matches / total_bullets) * 25
            
            # Industry/domain matching (15% weight)
            domain_score = 0
            industries = ['manufacturing', 'software', 'engineering', 'healthcare', 'finance', 
                         'education', 'retail', 'technology', 'construction', 'consulting']
            
            for industry in industries:
                if industry in jd_lower:
                    # Check if this industry appears in experience
                    for job in profile_exp_data:
                        company_lower = job.get('company', '').lower()
                        title_lower = job.get('title', '').lower()
                        if industry in company_lower or industry in title_lower:
                            domain_score = 15
                            break
                if domain_score > 0:
                    break
            
            if domain_score == 0:
                domain_score = 8  # Partial credit for transferable skills
            
            # Total score
            total_score = min(100, int(skill_score + exp_score + keyword_score + domain_score))

            # Adaptive categorization thresholds
            thresh = self.feedback_config.get('row_color_thresholds', {'high':65,'medium':40})
            if total_score >= thresh.get('high',65):
                category = 'high_match'
            elif total_score >= thresh.get('medium',40):
                category = 'medium_match'
            else:
                category = 'low_match'
            
            return total_score, category
            
        except Exception as e:
            print(f"Error calculating match score: {e}")
            return 0, 'low'
    
    def calculate_job_fit_score(self, jd_text, company, position):
        """
        Calculate how well the job matches the candidate's personal preferences.
        Considers: employment type, location, salary, work arrangement, responsibility level, etc.
        
        Returns:
            tuple: (fit_score_percentage, category) where category is 'high_match', 'medium_match', or 'low_match'
        """
        try:
            # Load preferences
            preferences_file = ROOT / 'data' / 'profile_preferences.json'
            if not preferences_file.exists():
                return 0, 'low_match'
            
            with open(preferences_file, 'r', encoding='utf-8') as f:
                prefs = json.load(f).get('job_preferences', {})
            
            jd_lower = jd_text.lower()
            total_score = 0
            max_score = 0
            
            # 1. Employment Type (20% weight)
            emp_type_score = 0
            emp_type_prefs = prefs.get('employment_type', {})
            if 'full-time' in jd_lower or 'full time' in jd_lower:
                emp_type_score = emp_type_prefs.get('full_time', 100)
            elif 'part-time' in jd_lower or 'part time' in jd_lower:
                emp_type_score = emp_type_prefs.get('part_time', 30)
            elif 'contract' in jd_lower:
                emp_type_score = emp_type_prefs.get('contract', 50)
            else:
                emp_type_score = emp_type_prefs.get('full_time', 100)  # Assume full-time if not specified
            total_score += emp_type_score * 0.20
            max_score += 100 * 0.20
            
            # 2. Work Arrangement (15% weight)
            work_arr_score = 0
            work_arr_prefs = prefs.get('work_arrangement', {})
            if 'remote' in jd_lower and 'hybrid' not in jd_lower:
                work_arr_score = work_arr_prefs.get('remote', 100)
            elif 'hybrid' in jd_lower:
                work_arr_score = work_arr_prefs.get('hybrid', 80)
            elif 'on-site' in jd_lower or 'onsite' in jd_lower or 'on site' in jd_lower:
                work_arr_score = work_arr_prefs.get('onsite', 60)
            else:
                work_arr_score = work_arr_prefs.get('hybrid', 80)  # Assume hybrid if not specified
            total_score += work_arr_score * 0.15
            max_score += 100 * 0.15
            
            # 3. Location Match (15% weight)
            location_score = 50  # Default neutral score
            loc_prefs = prefs.get('location_preferences', {})
            preferred_states = loc_prefs.get('preferred_states', [])
            
            # Check if any preferred state is in the JD
            for state in preferred_states:
                if state.lower() in jd_lower:
                    location_score = 100
                    break
            
            # Check relocation
            if 'relocation' in jd_lower:
                if loc_prefs.get('willing_to_relocate', True):
                    if loc_prefs.get('relocation_assistance_required', False):
                        if 'relocation assistance' in jd_lower or 'relocation package' in jd_lower:
                            location_score = max(location_score, 90)
                    else:
                        location_score = max(location_score, 85)
            
            total_score += location_score * 0.15
            max_score += 100 * 0.15
            
            # 4. Salary Match (25% weight)
            salary_score = 50  # Default neutral
            comp_prefs = prefs.get('compensation', {})
            min_salary = comp_prefs.get('minimum_salary', 0)
            target_salary = comp_prefs.get('target_salary', 0)
            ideal_salary = comp_prefs.get('ideal_salary', 0)
            
            # Extract salary from JD using regex
            import re
            salary_patterns = [
                r'\\$(\\d{1,3}(?:,\\d{3})*(?:\\.\\d{2})?)[\\s-]*(k|thousand)?',
                r'(\\d{1,3})k',
                r'salary.*?(\\d{1,3}(?:,\\d{3})*)',
            ]
            
            found_salary = None
            for pattern in salary_patterns:
                matches = re.findall(pattern, jd_lower)
                if matches:
                    try:
                        salary_str = matches[0] if isinstance(matches[0], str) else matches[0][0]
                        salary_str = salary_str.replace(',', '').replace('$', '').replace('k', '000')
                        found_salary = float(salary_str)
                        if found_salary < 1000:  # It's in thousands (e.g., "85k")
                            found_salary *= 1000
                        break
                    except:
                        pass
            
            if found_salary:
                if found_salary >= ideal_salary:
                    salary_score = 100
                elif found_salary >= target_salary:
                    salary_score = 80
                elif found_salary >= min_salary:
                    salary_score = 60
                else:
                    salary_score = 30
            
            total_score += salary_score * 0.25
            max_score += 100 * 0.25
            
            # 5. Responsibility Level (15% weight)
            resp_score = 50  # Default
            resp_prefs = prefs.get('responsibility_level', {})
            
            if any(word in jd_lower for word in ['director', 'vp', 'vice president']):
                resp_score = resp_prefs.get('director', 70)
            elif any(word in jd_lower for word in ['senior manager', 'sr. manager', 'sr manager']):
                resp_score = resp_prefs.get('senior_manager', 90)
            elif 'manager' in jd_lower or 'lead' in jd_lower:
                resp_score = resp_prefs.get('manager', 100)
            elif any(word in jd_lower for word in ['team lead', 'tech lead', 'technical lead']):
                resp_score = resp_prefs.get('team_lead', 80)
            else:
                resp_score = resp_prefs.get('individual_contributor', 50)
            
            total_score += resp_score * 0.15
            max_score += 100 * 0.15
            
            # 6. Travel Requirements (10% weight)
            travel_score = 100  # Default: assume acceptable
            work_style = prefs.get('work_style', {})
            travel_tolerance = work_style.get('travel_tolerance_percent', 25)
            
            # Extract travel percentage
            travel_matches = re.findall(r'(\\d{1,3})%?\\s*travel', jd_lower)
            if travel_matches:
                try:
                    travel_required = int(travel_matches[0])
                    if travel_required <= travel_tolerance:
                        travel_score = 100
                    elif travel_required <= travel_tolerance + 20:
                        travel_score = 70
                    else:
                        travel_score = 40
                except:
                    pass
            
            total_score += travel_score * 0.10
            max_score += 100 * 0.10
            
            # Calculate final percentage
            if max_score > 0:
                final_score = int((total_score / max_score) * 100)
            else:
                final_score = 0
            
            # Determine category using adaptive thresholds
            fit_thresh = self.feedback_config.get('job_fit_thresholds', {'high':70,'medium':50})
            if final_score >= fit_thresh.get('high',70):
                category = 'high_match'
            elif final_score >= fit_thresh.get('medium',50):
                category = 'medium_match'
            else:
                category = 'low_match'
            
            return final_score, category
            
        except Exception as e:
            print(f"Error calculating job fit score: {e}")
            return 0, 'low_match'
    
    def refresh_applications(self):
        """Refresh the applications list from tracker"""
        # Clear existing items
        for item in self.app_tree.get_children():
            self.app_tree.delete(item)
        
        # Load from tracker
        try:
            import openpyxl
            from pathlib import Path
            
            tracker_path = self.tracker_path
            
            if not tracker_path.exists():
                return
            
            wb = openpyxl.load_workbook(tracker_path, data_only=True)
            ws = wb['Job Applications']

            # Ensure Source / Job URL headers exist (AG=33, AH=34)
            try:
                if ws.cell(row=1, column=33).value is None:
                    ws.cell(row=1, column=33).value = 'Source'
                if ws.cell(row=1, column=34).value is None:
                    ws.cell(row=1, column=34).value = 'Job URL'
                wb.save(tracker_path)
            except Exception:
                pass
            
            # Read data (skip header)
            visual_index = 0  # For zebra striping
            for row in range(2, min(ws.max_row + 1, 200)):  # show more
                date_applied = ws.cell(row=row, column=1).value
                company = ws.cell(row=row, column=2).value
                position = ws.cell(row=row, column=3).value
                status = ws.cell(row=row, column=32).value  # AF column
                source_site = None
                try:
                    source_site = ws.cell(row=row, column=33).value
                except Exception:
                    source_site = None
                if not source_site:
                    source_site = "Other"
                
                if company and position:
                    # Format date - handle Excel serial dates (like 45981)
                    date_str = ""
                    if date_applied:
                        if isinstance(date_applied, datetime):
                            date_str = date_applied.strftime("%b %d, %Y")
                        elif isinstance(date_applied, (int, float)):
                            # Excel serial date (days since 1900-01-01)
                            try:
                                from datetime import timedelta
                                excel_epoch = datetime(1899, 12, 30)  # Excel's epoch
                                parsed = excel_epoch + timedelta(days=int(date_applied))
                                date_str = parsed.strftime("%b %d, %Y")
                            except:
                                date_str = str(date_applied)
                        else:
                            raw = str(date_applied)
                            parsed = None
                            for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y", "%m-%d-%Y"):
                                try:
                                    parsed = datetime.strptime(raw, fmt)
                                    break
                                except Exception:
                                    pass
                            date_str = parsed.strftime("%b %d, %Y") if parsed else raw
                    
                    # Try to load JD and calculate both match scores
                    match_score = 0
                    match_category = 'low_match'
                    fit_score = 0
                    fit_category = 'low_match'
                    jd_text = None
                    
                    # Look for JD file in outputs
                    output_dir = ROOT / 'outputs' / company
                    if output_dir.exists():
                        # Find the most recent JD file for this position
                        jd_files = list(output_dir.glob(f"**/*{position.replace(' ', '_')}*/job_description.txt"))
                        if not jd_files:
                            jd_files = list(output_dir.glob("**/job_description.txt"))
                        
                        if jd_files:
                            # Use the most recent
                            jd_file = max(jd_files, key=lambda p: p.stat().st_mtime)
                            try:
                                with open(jd_file, 'r', encoding='utf-8') as f:
                                    jd_text = f.read()
                                    # Calculate skills match
                                    match_score, match_category = self.calculate_match_score(jd_text, None)
                                    # Calculate job fit
                                    fit_score, fit_category = self.calculate_job_fit_score(jd_text, company, position)
                            except:
                                pass
                    
                    # Format score displays
                    match_display = f"{match_score}%" if match_score > 0 else "N/A"
                    fit_display = f"{fit_score}%" if fit_score > 0 else "N/A"
                    
                    # Adaptive thresholds for row coloring
                    row_thresh = self.feedback_config.get('row_color_thresholds', {'high':65,'medium':40})
                    combined_score = (match_score + fit_score) / 2 if (match_score > 0 and fit_score > 0) else max(match_score, fit_score)
                    if combined_score >= row_thresh.get('high',65):
                        row_category = 'high_match'
                    elif combined_score >= row_thresh.get('medium',40):
                        row_category = 'medium_match'
                    else:
                        row_category = 'low_match'
                    
                    zebra_tag = 'zebra_even' if visual_index % 2 == 0 else 'zebra_odd'
                    self.app_tree.insert(
                        "",
                        tk.END,
                        iid=str(row),  # store workbook row index
                        values=(company, position, date_str, source_site, status or "Applied", match_display, fit_display),
                        tags=(row_category, zebra_tag)
                    )
                    visual_index += 1
            # Auto-resize after populating rows so all columns are visible
            try:
                self.root.update_idletasks()
                tree_width = self.app_tree.winfo_reqwidth()
                padding = 140
                target_width = max(self.root.winfo_width(), tree_width + padding)
                if target_width > self.root.winfo_width():
                    self.root.geometry(f"{target_width}x{self.root.winfo_height()}")
                    self.root.minsize(target_width, 600)
            except Exception:
                pass
            
            wb.close()
        except Exception as e:
            print(f"Error loading tracker: {e}")

    # === Feedback Loop Functions ===
    def _load_feedback_config(self):
        """Load adaptive threshold configuration from JSON or create defaults."""
        import json
        if hasattr(self, 'feedback_config_path') and self.feedback_config_path.exists():
            try:
                with open(self.feedback_config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                pass
        return {
            'row_color_thresholds': {'high': 65, 'medium': 40},
            'job_fit_thresholds': {'high': 70, 'medium': 50},
            'adaptation': {'enabled': True}
        }

    def submit_feedback(self):
        """Persist user ratings for selected application and schedule adaptation."""
        try:
            sel = self.app_tree.selection()
            if not sel:
                messagebox.showinfo("Feedback", "Select an application first.")
                return
            iid = sel[0]
            values = self.app_tree.item(iid, 'values')
            if len(values) < 7:
                messagebox.showerror("Feedback", "Unexpected row format; cannot log feedback.")
                return
            company, position, date_applied, source_site, status, match_display, fit_display = values
            def _parse_pct(s):
                try:
                    return int(s.replace('%', ''))
                except Exception:
                    return 0
            match_score = _parse_pct(match_display)
            fit_score = _parse_pct(fit_display)
            user_match = getattr(self, 'match_rating_var', None).get() if hasattr(self, 'match_rating_var') else 0
            user_fit = getattr(self, 'fit_rating_var', None).get() if hasattr(self, 'fit_rating_var') else 0
            if user_match < 1 or user_fit < 1:
                messagebox.showinfo("Feedback", "Please set both ratings (1-5) before submitting.")
                return
            import json
            entry = {
                'ts': datetime.utcnow().isoformat(),
                'company': company,
                'position': position,
                'match_score': match_score,
                'fit_score': fit_score,
                'user_match_rating': user_match,
                'user_fit_rating': user_fit,
                'row_tags': self.app_tree.item(iid, 'tags')
            }
            log = []
            if self.feedback_log_path.exists():
                try:
                    with open(self.feedback_log_path, 'r', encoding='utf-8') as f:
                        log = json.load(f)
                except Exception:
                    log = []
            log.append(entry)
            self.feedback_log_path.parent.mkdir(parents=True, exist_ok=True)
            with open(self.feedback_log_path, 'w', encoding='utf-8') as f:
                json.dump(log, f, indent=2)
            self.update_status("✅ Feedback saved")
            # Reset sliders
            if hasattr(self, 'match_rating_var'):
                self.match_rating_var.set(0)
            if hasattr(self, 'fit_rating_var'):
                self.fit_rating_var.set(0)
            # Schedule adaptation check
            self.root.after(750, self._maybe_adapt_thresholds)
        except Exception as e:
            self.update_status("❌ Feedback error")
            print("Feedback error", e)

    def _maybe_adapt_thresholds(self):
        """Adjust thresholds if user ratings systematically disagree with current categories."""
        if not self.feedback_config.get('adaptation', {}).get('enabled', True):
            return
        import json
        if not self.feedback_log_path.exists():
            return
        try:
            with open(self.feedback_log_path, 'r', encoding='utf-8') as f:
                entries = json.load(f)
        except Exception:
            return
        if len(entries) < 5:
            # Not enough data yet; recheck later
            self.root.after(60000, self._maybe_adapt_thresholds)
            return
        recent = entries[-25:]
        high_under = 0  # user rated high (>=4) but system low/medium
        low_over = 0    # user rated low (<=2) but system high
        for e in recent:
            um = e.get('user_match_rating', 0)
            tags = e.get('row_tags', [])
            cat = None
            for t in tags:
                if t in ('high_match', 'medium_match', 'low_match'):
                    cat = t
                    break
            if um >= 4 and cat in ('low_match', 'medium_match'):
                high_under += 1
            if um <= 2 and cat == 'high_match':
                low_over += 1
        row_thresh = self.feedback_config.get('row_color_thresholds', {'high': 65, 'medium': 40})
        changed = False
        if high_under >= 5:
            row_thresh['high'] = max(50, row_thresh['high'] - 3)
            row_thresh['medium'] = max(25, row_thresh['medium'] - 2)
            changed = True
        if low_over >= 5:
            row_thresh['high'] = min(90, row_thresh['high'] + 3)
            row_thresh['medium'] = min(row_thresh['high'] - 10, row_thresh['medium'] + 2)
            changed = True
        if changed:
            self.feedback_config['row_color_thresholds'] = row_thresh
            try:
                with open(self.feedback_config_path, 'w', encoding='utf-8') as f:
                    json.dump(self.feedback_config, f, indent=2)
            except Exception:
                pass
            self.update_status(f"🔁 Adapted thresholds: high>={row_thresh['high']} medium>={row_thresh['medium']}")
            self.refresh_applications()
        # Schedule next periodic check
        self.root.after(60000, self._maybe_adapt_thresholds)

    # === View Applications Editing Logic ===
    def _on_status_changed(self, event=None):
        """Auto-save when status is changed from dropdown"""
        selected = self.app_tree.selection()
        if selected:
            self.update_selected_application()

    def on_application_select(self, event):
        selected = self.app_tree.selection()
        if not selected:
            return
        iid = selected[0]
        values = self.app_tree.item(iid, 'values')
        if not values:
            return
        company, position, date_applied, source_site, status, match_score, fit_score = values
        
        # Populate basic fields
        self.view_company_var.set(company)
        self.view_position_var.set(position)
        self.view_date_var.set(date_applied)
        self.view_status_var.set(status)
        
        # Populate new detail fields
        self.view_source_var.set(source_site if source_site else "Unknown")
        
        # Format and display match scores
        try:
            match_val = float(match_score) if match_score else 0
            self.view_match_var.set(f"{match_val:.0f}%")
        except:
            self.view_match_var.set("N/A")
        
        try:
            fit_val = float(fit_score) if fit_score else 0
            self.view_fit_var.set(f"{fit_val:.0f}%")
        except:
            self.view_fit_var.set("N/A")
        
        # Load additional details from workbook
        try:
            import openpyxl
            tracker_path = self.tracker_path
            if tracker_path.exists():
                wb = openpyxl.load_workbook(tracker_path)
                ws = wb['Job Applications']
                row_index = int(iid)
                
                # Load URL (stored in column 34 if present)
                url_val = ws.cell(row=row_index, column=34).value
                self.view_url_text.delete('1.0', tk.END)
                if url_val:
                    self.view_url_text.insert(tk.END, str(url_val))
                
                # Load notes (column 49)
                notes_val = ws.cell(row=row_index, column=49).value
                self.view_notes_text.delete('1.0', tk.END)
                if notes_val:
                    self.view_notes_text.insert(tk.END, str(notes_val))
                
                # Load lessons learned (column 50)
                lessons_val = ws.cell(row=row_index, column=50).value
                self.view_lessons_text.delete('1.0', tk.END)
                if lessons_val:
                    self.view_lessons_text.insert(tk.END, str(lessons_val))
                
                wb.close()
        except Exception as e:
            print(f"Error loading details: {e}")

        # Recompute match / fit scores live using JD file if available
        try:
            output_dir = ROOT / 'outputs' / company
            jd_file = None
            if output_dir.exists():
                # Find latest JD file referencing position
                pattern = f"**/*{position.replace(' ', '_')}*/job_description.txt"
                jd_files = list(output_dir.glob(pattern)) or list(output_dir.glob('**/job_description.txt'))
                if jd_files:
                    jd_file = max(jd_files, key=lambda p: p.stat().st_mtime)
            if jd_file and jd_file.exists():
                with open(jd_file, 'r', encoding='utf-8') as f:
                    jd_text = f.read()
                new_match, m_cat = self.calculate_match_score(jd_text, None)
                new_fit, f_cat = self.calculate_job_fit_score(jd_text, company, position)
                self.view_match_var.set(f"{new_match}%")
                self.view_fit_var.set(f"{new_fit}%")
        except Exception:
            pass

    def update_selected_application(self):
        selected = self.app_tree.selection()
        if not selected:
            messagebox.showinfo("No Selection", "Select an application first.")
            return
        iid = selected[0]
        try:
            import openpyxl
            tracker_path = self.tracker_path
            if not tracker_path.exists():
                messagebox.showerror("Tracker Missing", "Tracker file not found.")
                return
            wb = openpyxl.load_workbook(tracker_path)
            ws = wb['Job Applications']
            row_index = int(iid)

            # Normalize fields to standard
            company = self._normalize_ws(self.view_company_var.get())
            position = self._normalize_ws(self.view_position_var.get())
            date_text = self._normalize_ws(self.view_date_var.get())
            status = self.view_status_var.get().strip()
            url_val = self.view_url_text.get('1.0', tk.END).strip()
            notes_val = self._normalize_notes(self.view_notes_text.get('1.0', tk.END))
            lessons_val = self._normalize_notes(self.view_lessons_text.get('1.0', tk.END))

            # Validate and format date
            parsed_date = None
            if date_text:
                parsed_date = self._parse_date_flexible(date_text)
                if not parsed_date:
                    messagebox.showerror("Invalid Date", "Please enter date as MM/DD/YYYY (e.g., 11/20/2025).")
                    return

            # Reflect normalized in UI
            self.view_company_var.set(company)
            self.view_position_var.set(position)
            if parsed_date:
                self.view_date_var.set(parsed_date.strftime('%m/%d/%Y'))
            self.view_notes_text.delete('1.0', tk.END)
            if notes_val:
                self.view_notes_text.insert(tk.END, notes_val)
            self.view_lessons_text.delete('1.0', tk.END)
            if lessons_val:
                self.view_lessons_text.insert(tk.END, lessons_val)

            # Update workbook (avoid overwriting formulas or other data)
            ws.cell(row=row_index, column=2).value = company  # Company
            ws.cell(row=row_index, column=3).value = position  # Position correct column
            # Ensure headers for Source / Job URL
            if ws.cell(row=1, column=33).value is None:
                ws.cell(row=1, column=33).value = 'Source'
            if ws.cell(row=1, column=34).value is None:
                ws.cell(row=1, column=34).value = 'Job URL'
            # Persist source & URL
            source_val = self.view_source_var.get().strip() or None
            ws.cell(row=row_index, column=33).value = source_val
            ws.cell(row=row_index, column=34).value = url_val if url_val else None
            if parsed_date:
                ws.cell(row=row_index, column=1).value = parsed_date
                ws.cell(row=row_index, column=1).number_format = 'MM/DD/YYYY'
            ws.cell(row=row_index, column=32).value = status
            ws.cell(row=row_index, column=49).value = notes_val if notes_val else None
            ws.cell(row=row_index, column=50).value = lessons_val if lessons_val else None

            wb.save(tracker_path)
            wb.close()
            
            # Visual feedback - highlight the saved row briefly
            self.app_tree.item(iid, tags=('saved',))
            self.app_tree.tag_configure('saved', background='#d4f4dd')  # Light green
            
            # Refresh and restore selection
            self.refresh_applications()
            self.app_tree.selection_set(iid)
            
            # Update status bar
            self.update_status(f"✓ Saved changes to {company} - {position}")
            
            # Remove highlight after 2 seconds
            self.root.after(2000, lambda: self.app_tree.item(iid, tags=()))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to update application: {e}")
            import traceback
            traceback.print_exc()

    def delete_selected_application(self):
        selected = self.app_tree.selection()
        if not selected:
            messagebox.showinfo("No Selection", "Select an application first.")
            return
        iid = selected[0]
        if not messagebox.askyesno("Confirm Delete", "Delete selected application row from tracker? This cannot be undone."):
            return
        try:
            import openpyxl
            from pathlib import Path
            tracker_path = self.tracker_path
            if not tracker_path.exists():
                messagebox.showerror("Tracker Missing", "Tracker file not found.")
                return
            wb = openpyxl.load_workbook(tracker_path)
            ws = wb['Job Applications']
            row_index = int(iid)
            ws.delete_rows(row_index, 1)
            wb.save(tracker_path)
            wb.close()
            self.refresh_applications()
            self.clear_view_selection()
            messagebox.showinfo("Deleted", "Application deleted.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete: {e}")

    def export_selected_application(self):
        selected = self.app_tree.selection()
        if not selected:
            messagebox.showinfo("No Selection", "Select an application first.")
            return
        iid = selected[0]
        values = self.app_tree.item(iid, 'values')
        if not values:
            return
        company, position, date_applied, source_site, status, match_score, fit_score = values
        notes = self.view_notes_text.get('1.0', tk.END).strip()
        lessons = self.view_lessons_text.get('1.0', tk.END).strip()
        # Simple CSV export (append to single file)
        import csv
        export_path = ROOT / 'exports'
        export_path.mkdir(exist_ok=True)
        file_path = export_path / 'selected_applications.csv'
        file_exists = file_path.exists()
        try:
            with open(file_path, 'a', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                if not file_exists:
                    writer.writerow(['Company','Position','Date Applied','Status','Notes','Lessons Learned'])
                writer.writerow([company, position, date_applied, status, notes, lessons])
            messagebox.showinfo("Exported", f"Row exported to {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")

    def clear_view_selection(self):
        self.app_tree.selection_remove(self.app_tree.selection())
        self.view_company_var.set('')
        self.view_position_var.set('')
        self.view_date_var.set('')
        self.view_status_var.set('')
        self.view_source_var.set('')
        self.view_match_var.set('N/A')
        self.view_fit_var.set('N/A')
        self.view_url_text.delete('1.0', tk.END)
        self.view_notes_text.delete('1.0', tk.END)
        self.view_lessons_text.delete('1.0', tk.END)

    # === Formatting Helpers ===
    def _normalize_ws(self, s: str) -> str:
        import re
        if not s:
            return ''
        return re.sub(r"\s+", " ", s).strip()

    def _normalize_notes(self, text: str) -> str:
        import re
        if not text:
            return ''
        # Normalize line endings and trim trailing spaces per line
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        lines = [re.sub(r"\s+$", "", ln) for ln in text.split('\n')]
        # Collapse consecutive blank lines to single
        out_lines = []
        blank = False
        for ln in lines:
            if ln.strip() == '':
                if not blank:
                    out_lines.append('')
                blank = True
            else:
                out_lines.append(ln)
                blank = False
        return '\n'.join(out_lines).strip()

    def _parse_date_flexible(self, s: str):
        from datetime import datetime
        s = s.strip()
        patterns = [
            '%m/%d/%Y', '%m-%d-%Y', '%Y-%m-%d', '%b %d, %Y', '%B %d, %Y'
        ]
        for p in patterns:
            try:
                return datetime.strptime(s, p)
            except ValueError:
                continue
        return None

    def _format_display_date(self, value, epoch=None) -> str:
        """Return MM/DD/YYYY for various Excel date representations."""
        from datetime import datetime
        if not value:
            return ""
        if isinstance(value, datetime):
            return value.strftime('%m/%d/%Y')
        # Excel serial number
        try:
            if isinstance(value, (int, float)):
                try:
                    from openpyxl.utils.datetime import from_excel
                    if epoch is None:
                        # Default to Windows 1900 calendar
                        from openpyxl.utils.datetime import CALENDAR_WINDOWS_1900
                        dt = from_excel(value, CALENDAR_WINDOWS_1900)
                    else:
                        dt = from_excel(value, epoch)
                    return dt.strftime('%m/%d/%Y')
                except Exception:
                    pass
            # Try parsing strings
            if isinstance(value, str):
                dt = self._parse_date_flexible(value)
                if dt:
                    return dt.strftime('%m/%d/%Y')
        except Exception:
            pass
        return str(value)
    
    def open_tracker(self):
        """Open the Excel tracker"""
        tracker_path = self.tracker_path
        if tracker_path.exists():
            webbrowser.open(str(tracker_path))
        else:
            messagebox.showinfo("Not Found", "Tracker file not found. Apply to a job first to create it.")
    
    def open_output_folder(self):
        """Open the outputs folder"""
        output_path = ROOT / 'outputs'
        if output_path.exists():
            webbrowser.open(str(output_path))
        else:
            messagebox.showinfo("Not Found", "Output folder not found. Apply to a job first to create it.")
    
    def upload_tracker(self):
        """Upload and import existing Excel tracker"""
        file_path = filedialog.askopenfilename(
            title="Select Excel Tracker to Upload",
            filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            import openpyxl
            import csv
            from pathlib import Path
            
            self.update_status("📤 Importing tracker data...")
            
            # Load the Excel file
            wb = openpyxl.load_workbook(file_path, data_only=True)
            ws = wb.active
            
            # Read headers from first row
            headers = []
            for cell in ws[1]:
                headers.append(cell.value if cell.value else "")
            
            # Find the CSV tracker file
            csv_tracker = ROOT / 'job_applications_tracker.csv'
            
            # Read existing data if it exists
            existing_data = {}
            if csv_tracker.exists():
                with open(csv_tracker, 'r', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        # Use company + position as unique key
                        key = f"{row.get('Company', '')}_{row.get('Position', '')}"
                        existing_data[key] = row
            
            # Import new data from Excel
            imported_count = 0
            updated_count = 0
            
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not row or not any(row):  # Skip empty rows
                    continue
                
                # Create dict from row data
                row_data = {}
                for idx, value in enumerate(row):
                    if idx < len(headers):
                        header = headers[idx]
                        row_data[header] = value if value is not None else ""
                
                # Check if this application exists
                company = row_data.get('Company', '').strip()
                position = row_data.get('Position', '').strip()
                
                if not company or not position:
                    continue  # Skip rows without company/position
                
                key = f"{company}_{position}"
                
                if key in existing_data:
                    # Update existing entry
                    existing_data[key].update(row_data)
                    updated_count += 1
                else:
                    # Add new entry
                    existing_data[key] = row_data
                    imported_count += 1
            
            # Write combined data back to CSV
            if existing_data:
                with open(csv_tracker, 'w', newline='', encoding='utf-8') as f:
                    # Use headers from Excel file
                    writer = csv.DictWriter(f, fieldnames=headers)
                    writer.writeheader()
                    for row_data in existing_data.values():
                        writer.writerow(row_data)
                
                # Refresh the view
                self.refresh_applications()
                
                message = f"✅ Import Complete!\n\n"
                message += f"• New applications: {imported_count}\n"
                message += f"• Updated applications: {updated_count}\n"
                message += f"• Total applications: {len(existing_data)}"
                
                messagebox.showinfo("Upload Successful", message)
                self.update_status(f"✅ Imported {imported_count} new, updated {updated_count} applications")
            else:
                messagebox.showwarning("No Data", "No valid application data found in the uploaded tracker.")
                self.update_status("⚠️ No data imported")
            
            wb.close()
            
        except Exception as e:
            messagebox.showerror("Upload Failed", f"Failed to import tracker:\n\n{str(e)}")
            self.update_status("❌ Upload failed")

    
    def add_reference_resumes(self):
        """Add multiple reference resume documents"""
        file_paths = filedialog.askopenfilenames(
            title="Select Reference Resumes (can select multiple)",
            filetypes=[("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if file_paths:
            added_count = 0
            for file_path in file_paths:
                if file_path not in self.reference_resumes:
                    self.reference_resumes.append(file_path)
                    # Display just the filename
                    from pathlib import Path
                    filename = Path(file_path).name
                    self.ref_listbox.insert(tk.END, filename)
                    added_count += 1
            
            if added_count > 0:
                self.update_status(f"Added {added_count} reference resume(s)")
                self._save_reference_resumes()
            else:
                messagebox.showinfo("No New Files", "All selected files were already in the reference list.")
    
    def remove_reference_resume(self):
        """Remove selected reference resume from list"""
        selection = self.ref_listbox.curselection()
        if selection:
            index = selection[0]
            removed_file = self.reference_resumes.pop(index)
            self.ref_listbox.delete(index)
            from pathlib import Path
            self.update_status(f"Removed {Path(removed_file).name}")
            self._save_reference_resumes()
        else:
            messagebox.showwarning("No Selection", "Please select a reference resume to remove.")
    
    def clear_reference_resumes(self):
        """Clear all reference resumes"""
        if self.reference_resumes:
            count = len(self.reference_resumes)
            if messagebox.askyesno("Confirm Clear", f"Remove all {count} reference resume(s)?"):
                self.reference_resumes.clear()
                self.ref_listbox.delete(0, tk.END)
                self.update_status("Cleared all reference resumes")
                self._save_reference_resumes()
        else:
            messagebox.showinfo("Empty", "No reference resumes to clear.")
    
    def _save_reference_resumes(self):
        """Save reference resume list to JSON file"""
        import json
        from pathlib import Path
        ref_file = Path(self.root_dir) / 'data' / 'reference_resumes.json'
        ref_file.parent.mkdir(parents=True, exist_ok=True)
        try:
            with open(ref_file, 'w', encoding='utf-8') as f:
                json.dump(self.reference_resumes, f, indent=2)
        except Exception as e:
            print(f"Warning: Could not save reference resumes: {e}")
    
    def _load_reference_resumes(self):
        """Load reference resume list from JSON file"""
        import json
        from pathlib import Path
        ref_file = Path(self.root_dir) / 'data' / 'reference_resumes.json'
        if ref_file.exists():
            try:
                with open(ref_file, 'r', encoding='utf-8') as f:
                    self.reference_resumes = json.load(f)
                    # Populate listbox
                    for file_path in self.reference_resumes:
                        filename = Path(file_path).name
                        self.ref_listbox.insert(tk.END, filename)
            except Exception as e:
                print(f"Warning: Could not load reference resumes: {e}")
    
    def browse_profile(self):
        """Browse for profile file"""
        file_path = filedialog.askopenfilename(
            title="Select Profile File",
            filetypes=[("YAML Files", "*.yaml"), ("All Files", "*.*")]
        )
        if file_path:
            self.profile_path_var.set(file_path)
    
    def browse_output(self):
        """Browse for output folder"""
        folder_path = filedialog.askdirectory(title="Select Output Folder")
        if folder_path:
            self.output_path_var.set(folder_path)
    
    def update_status(self, message):
        """Update the status bar"""
        self.status_bar.config(text=f"  {message}")
        self.root.update_idletasks()


def main():
    root = tk.Tk()
    app = JobApplicationGUI(root)
    root.mainloop()


if __name__ == '__main__':
    main()
