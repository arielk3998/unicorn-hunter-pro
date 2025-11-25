"""
Resume scoring engine to evaluate resume quality against industry best practices.
Scores on a 1-100 scale across multiple dimensions.
"""
import os
import re
from docx import Document

ROOT = os.path.dirname(os.path.dirname(__file__))  # Project root
OUTPUT_DIR = os.path.join(ROOT, 'outputs')
RESUME_PATH = os.path.join(OUTPUT_DIR, '3M_Supply_Chain_Engineer_Resume_Enhanced.docx')

def score_resume(doc_path):
    """Comprehensive resume scoring across 10 dimensions."""
    
    doc = Document(doc_path)
    text = '\n'.join([p.text for p in doc.paragraphs])
    
    scores = {}
    
    # 1. Quantifiable Achievements (15 points)
    metrics_pattern = r'\d+%|\d+x|\d+\+|\$\d+[KMB]?|\d+,\d+'
    metrics = re.findall(metrics_pattern, text)
    scores['quantifiable_achievements'] = min(15, len(metrics) * 2)
    
    # 2. Action Verbs & Strong Language (10 points)
    action_verbs = ['Led', 'Managed', 'Delivered', 'Implemented', 'Enhanced', 'Reduced', 
                    'Increased', 'Developed', 'Designed', 'Optimized', 'Automated', 
                    'Collaborated', 'Spearheaded', 'Conducted', 'Established', 'Achieved']
    verb_count = sum(text.count(verb) for verb in action_verbs)
    scores['action_verbs'] = min(10, verb_count)
    
    # 3. ATS Optimization (15 points)
    ats_score = 0
    # Check for standard section headers
    required_sections = ['Professional Summary', 'Experience', 'Education', 'Skills']
    ats_score += sum(3 for section in required_sections if section.upper() in text.upper())
    # No tables/graphics warning (assume OK if docx loads)
    ats_score += 3
    scores['ats_optimization'] = min(15, ats_score)
    
    # 4. Keyword Density for Target Role (15 points)
    # 3M Supply Chain Engineer keywords
    target_keywords = ['NPI', 'Scale-up', 'CAPEX', 'Lean', 'Six Sigma', 'Process Optimization',
                       'Manufacturing', 'Quality', 'Cross-functional', 'Supply Chain', 
                       'Automation', 'ERP', 'Documentation', 'Compliance', 'Root Cause']
    keyword_matches = sum(1 for kw in target_keywords if kw.lower() in text.lower())
    scores['keyword_density'] = min(15, keyword_matches)
    
    # 5. Conciseness & Clarity (10 points)
    word_count = len(text.split())
    # Optimal 2-page resume: 500-800 words
    if 500 <= word_count <= 800:
        scores['conciseness'] = 10
    elif 400 <= word_count < 500 or 800 < word_count <= 900:
        scores['conciseness'] = 7
    else:
        scores['conciseness'] = 5
    
    # 6. Professional Formatting (10 points)
    format_score = 0
    # Check for consistent structure (paragraphs > 10 = well-structured)
    if len(doc.paragraphs) > 15:
        format_score += 5
    # Check for bullet points
    if '•' in text or any('List' in p.style.name for p in doc.paragraphs if p.style):
        format_score += 5
    scores['formatting'] = format_score
    
    # 7. Relevance to Target Role (15 points)
    role_alignment = 0
    # Check for supply chain/manufacturing focus
    domain_terms = ['supply chain', 'manufacturing', 'production', 'assembly', 
                    'engineering', 'process', 'quality', 'operations']
    role_alignment += sum(2 for term in domain_terms if term in text.lower())
    scores['role_relevance'] = min(15, role_alignment)
    
    # 8. Skills Organization (5 points)
    # Check for dedicated skills section
    if 'competencies' in text.lower() or 'skills' in text.lower():
        scores['skills_organization'] = 5
    else:
        scores['skills_organization'] = 2
    
    # 9. Contact Info Completeness (3 points)
    contact_score = 0
    if re.search(r'\(\d{3}\)\s?\d{3}-\d{4}', text):  # Phone
        contact_score += 1
    if '@' in text:  # Email
        contact_score += 1
    if 'linkedin.com' in text.lower():  # LinkedIn
        contact_score += 1
    scores['contact_info'] = contact_score
    
    # 10. Error-Free Writing (2 points)
    # Basic check for common issues
    error_score = 2
    if text.count('  ') > 5:  # Double spaces
        error_score -= 0.5
    if re.search(r'\b(i|Im|dont|cant)\b', text):  # Missing apostrophes/caps
        error_score -= 0.5
    scores['error_free'] = max(0, error_score)
    
    # Calculate total
    total_score = sum(scores.values())
    
    # Generate detailed feedback
    feedback = {
        'total_score': round(total_score),
        'grade': get_grade(total_score),
        'dimension_scores': scores,
        'strengths': [],
        'improvements': []
    }
    
    # Identify strengths
    if scores['quantifiable_achievements'] >= 12:
        feedback['strengths'].append("Excellent use of quantified metrics and achievements")
    if scores['keyword_density'] >= 12:
        feedback['strengths'].append("Strong alignment with target role keywords")
    if scores['ats_optimization'] >= 12:
        feedback['strengths'].append("Well-optimized for ATS parsing")
    if scores['role_relevance'] >= 12:
        feedback['strengths'].append("Highly relevant experience for target position")
    
    # Identify improvements
    if scores['quantifiable_achievements'] < 10:
        feedback['improvements'].append("Add more quantified achievements with specific metrics")
    if scores['keyword_density'] < 10:
        feedback['improvements'].append("Incorporate more industry-specific keywords (NPI, CAPEX, Lean Six Sigma)")
    if scores['action_verbs'] < 7:
        feedback['improvements'].append("Use stronger action verbs to lead bullet points")
    if scores['conciseness'] < 8:
        feedback['improvements'].append("Adjust length to optimal 500-800 word range for 2-page format")
    if scores['formatting'] < 8:
        feedback['improvements'].append("Enhance visual structure with consistent formatting and bullet points")
    
    return feedback

def get_grade(score):
    """Convert numeric score to letter grade."""
    if score >= 90:
        return 'A+ (Exceptional)'
    elif score >= 85:
        return 'A (Excellent)'
    elif score >= 80:
        return 'B+ (Very Good)'
    elif score >= 75:
        return 'B (Good)'
    elif score >= 70:
        return 'C+ (Above Average)'
    elif score >= 65:
        return 'C (Average)'
    else:
        return 'D (Needs Improvement)'

def print_score_report(feedback):
    """Print formatted score report."""
    print("\n" + "="*60)
    print("RESUME QUALITY SCORE REPORT")
    print("="*60)
    print(f"\nOVERALL SCORE: {feedback['total_score']}/100")
    print(f"GRADE: {feedback['grade']}")
    print("\n" + "-"*60)
    print("DIMENSION BREAKDOWN:")
    print("-"*60)
    
    dimensions = {
        'quantifiable_achievements': 'Quantifiable Achievements (15pts)',
        'ats_optimization': 'ATS Optimization (15pts)',
        'keyword_density': 'Keyword Density (15pts)',
        'role_relevance': 'Role Relevance (15pts)',
        'action_verbs': 'Action Verbs & Language (10pts)',
        'conciseness': 'Conciseness & Clarity (10pts)',
        'formatting': 'Professional Formatting (10pts)',
        'skills_organization': 'Skills Organization (5pts)',
        'contact_info': 'Contact Info Completeness (3pts)',
        'error_free': 'Error-Free Writing (2pts)'
    }
    
    for key, label in dimensions.items():
        score = feedback['dimension_scores'][key]
        max_pts = int(re.search(r'\((\d+)pts\)', label).group(1))
        bar_length = int((score / max_pts) * 20)
        bar = '█' * bar_length + '░' * (20 - bar_length)
        print(f"{label:45} {bar} {score:.1f}/{max_pts}")
    
    print("\n" + "-"*60)
    if feedback['strengths']:
        print("STRENGTHS:")
        for strength in feedback['strengths']:
            print(f"  ✓ {strength}")
    
    if feedback['improvements']:
        print("\nSUGGESTED IMPROVEMENTS:")
        for improvement in feedback['improvements']:
            print(f"  → {improvement}")
    
    print("\n" + "="*60)

if __name__ == '__main__':
    feedback = score_resume(RESUME_PATH)
    print_score_report(feedback)
