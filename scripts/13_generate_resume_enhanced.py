"""
Enhanced 2-page ATS-friendly resume generator with JD-aware skill mapping.
Emphasizes transferable skills and quantified achievements for maximum impact.
"""
import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(__file__))  # Project root (parent of scripts/)
DATA_DIR = os.path.join(ROOT, 'data')
OUTPUT_DIR = os.path.join(ROOT, 'outputs')

CONTACT_FILE = os.path.join(DATA_DIR, 'profile_contact.json')
EXP_JSON = os.path.join(DATA_DIR, 'profile_experience.json')
PROFILE_JSON = os.path.join(DATA_DIR, 'profile_candidate.json')
STYLE_SPEC = os.path.join(DATA_DIR, 'config_style_spec.json')
OUT_ENHANCED = os.path.join(OUTPUT_DIR, '3M_Supply_Chain_Engineer_Resume_Enhanced.docx')

# ATS-friendly color scheme (subtle but memorable)
ACCENT = RGBColor(0x00, 0x51, 0x99)  # Professional blue
FONT_PRIMARY = 'Arial'
FONT_HEADINGS = 'Arial'

# JD-aligned skill mapping for 3M Supply Chain Engineer role
SKILL_CATEGORIES = {
    'Manufacturing Excellence': [
        'NPI Integration & Scale-Up',
        'Process Optimization',
        'Production Throughput Enhancement',
        'Manufacturing Process Design',
        'Equipment Validation & Testing'
    ],
    'Quality & Compliance': [
        'Customer Complaint Resolution',
        'Cost of Poor Quality Reduction',
        'Quality Assurance Protocols',
        'Regulatory Compliance (Aerospace)',
        'Technical Documentation Standards'
    ],
    'Process Engineering': [
        'Lean Six Sigma Methodologies',
        'Value Stream Mapping',
        'Root Cause Analysis',
        'Process Standardization (SOPs)',
        'Automation & Workflow Optimization'
    ],
    'Technical Leadership': [
        'Cross-Functional Team Collaboration',
        'Subject Matter Expertise (Assembly/Testing)',
        'Technical Training & Mentoring',
        'Stakeholder Communication',
        'Project Management'
    ],
    'Systems & Technology': [
        'ERP Systems (Aurora, Dynamics 365)',
        'CAPEX Planning & Equipment Design Support',
        'Microsoft Office Suite (Advanced)',
        'Data Analysis & Reporting',
        'New Technology Adoption'
    ]
}

def load_json(path):
    """Load JSON file with error handling."""
    if not os.path.exists(path):
        return {}
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (OSError, json.JSONDecodeError):
        return {}

def load_style_spec():
    return load_json(STYLE_SPEC)

def load_contact():
    return load_json(CONTACT_FILE)

def load_experience():
    data = load_json(EXP_JSON)
    return data if isinstance(data, list) else []

def load_profile():
    return load_json(PROFILE_JSON)

def apply_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7):
    """Set narrow margins to maximize space while maintaining DANS compliance."""
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    # DANS: Ensure standard page size
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)

def add_dans_metadata(doc, contact):
    """Add DANS-compliant document metadata for digital application navigation."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - Professional Resume"
    core_props.subject = "Resume - ATS & DANS Optimized"
    core_props.keywords = "resume, ATS, DANS, professional, supply chain, manufacturing, engineering"
    core_props.category = "Resume"
    core_props.comments = "Digital Application Navigation System (DANS) compliant with ATS optimization"

def add_horizontal_line(paragraph):
    """Add subtle horizontal line below paragraph for visual separation."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '005199')
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_header(doc, contact):
    """Build professional DANS-compliant header with name and contact info."""
    # Name - DANS: Use Heading 1 for document structure
    name_p = doc.add_paragraph()
    name_p.style = 'Heading 1'
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(contact.get('name', 'Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = FONT_HEADINGS
    r.font.color.rgb = ACCENT
    name_p.paragraph_format.space_after = Pt(2)
    
    # Contact line
    contact_parts = [
        contact.get('location_statement', 'Tucson, AZ'),
        contact.get('phone', '(520) 591-9667'),
        contact.get('email', 'Ariel.pk@outlook.com'),
        contact.get('linkedin', 'linkedin.com/in/ariel-karagodskiy-98b912194')
    ]
    contact_line = ' • '.join([p for p in contact_parts if p])
    
    p = doc.add_paragraph(contact_line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.name = FONT_PRIMARY
    p.paragraph_format.space_after = Pt(8)
    
    add_horizontal_line(p)

def add_section_title(doc, title, space_before=6):
    """Add DANS-compliant section heading with consistent formatting."""
    p = doc.add_paragraph()
    # DANS: Use Heading 2 for proper document hierarchy
    p.style = 'Heading 2'
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper())  # DANS: uppercase for clarity
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = FONT_HEADINGS
    r.font.color.rgb = ACCENT
    return p

def write_professional_summary(doc, summary_text=None):
    """Write compelling summary targeting specific role."""
    add_section_title(doc, 'Professional Summary')
    
    if summary_text is None:
        summary_text = (
            "Results-driven Manufacturing & Supply Chain Engineer with 10+ years optimizing production systems, "
            "leading NPI scale-up initiatives, and implementing Lean Six Sigma process improvements. Proven track "
            "record: 300% production throughput increase, 50% testing efficiency gains, 20% error reduction, and "
            "15% downtime reduction. Subject matter expert in cross-functional collaboration, CAPEX planning, and "
            "translating business requirements into scalable manufacturing solutions while maintaining regulatory compliance."
        )
    
    p = doc.add_paragraph(summary_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.name = FONT_PRIMARY
    p.paragraph_format.space_after = Pt(6)

def write_key_achievements(doc, achievements=None):
    """Highlight quantified achievements aligned with JD priorities.

    Accepts optional list of pre-deduplicated achievements.
    """
    add_section_title(doc, 'Key Achievements & Impact')

    if achievements is None:
        achievements = [
            "Delivered 300% production throughput increase and 50% testing efficiency gains for human-rated aerospace components through systematic process reengineering",
            "Spearheaded ERP migration (Aurora → Dynamics 365) reducing documentation errors 20% and enabling cross-departmental standardization",
            "Led cross-functional technical writing team achieving 95% on-time delivery while elevating quality standards organization-wide",
            "Decreased system downtime 15% and improved safety compliance 20% through structured protocols and targeted training initiatives",
            "Automated technical workflows achieving zero-defect accuracy and eliminating manual calculation errors"
        ]

    for achievement in achievements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(achievement)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.name = FONT_PRIMARY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

def write_core_competencies(doc, custom_categories=None):
    """Create organized skills matrix showcasing transferable capabilities.

    Allows override via custom_categories dict.
    """
    add_section_title(doc, 'Core Competencies & Technical Skills')

    categories = custom_categories if isinstance(custom_categories, dict) else SKILL_CATEGORIES

    for category, skills in categories.items():
        p = doc.add_paragraph()
        cat_run = p.add_run(f"{category}: ")
        cat_run.bold = True
        cat_run.font.size = Pt(10)
        cat_run.font.name = FONT_PRIMARY
        skills_run = p.add_run(' • '.join(skills))
        skills_run.font.size = Pt(10)
        skills_run.font.name = FONT_PRIMARY
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)

def write_professional_experience(doc, entries):
    """Write detailed experience section with JD-aligned bullet points."""
    add_section_title(doc, 'Professional Experience')
    
    # Enhanced bullet points mapping to 3M JD requirements
    enhanced_bullets = {
        'FLSmidth': [
            "Led ERP migration (Aurora → Dynamics 365) as technical SME for cross-functional team, reducing process errors 20% and enabling seamless integration across IT, Operations, and Quality",
            "Managed team of 4 technical writers delivering SOPs, training modules, and user manuals with 95% on-time rate while establishing department-wide documentation standards",
            "Conducted gap analysis and process mapping optimizing workflows and increasing productivity through systematic tool integration and compliance frameworks",
            "Partnered with R&D and engineering teams to translate complex specifications into actionable, compliance-ready documentation supporting cost of poor quality reduction"
        ],
        'Jansen Aircraft System Control': [
            "Served as process optimization SME delivering 300% throughput improvement and 50% testing efficiency gains for human-rated aerospace valves through systematic protocol development",
            "Automated technical workflows using algorithmic solutions, achieving zero-defect accuracy and eliminating manual calculation errors",
            "Collaborated with R&D and Quality teams on root cause analysis and corrective actions for unexpected issues, ensuring regulatory compliance",
            "Trained and mentored technicians on specialized tooling and testing SOPs, elevating adherence to aerospace quality standards",
            "Strengthened supply chain integration through transparent cross-functional communication and production support protocols"
        ],
        'Honeywell': [
            "Led process improvement initiatives enhancing manufacturing efficiency and reducing defects through Lean principles and specialized tooling design",
            "Conducted rigorous quality assurance inspections supporting zero-defect objectives and contributing to continuous improvement programs",
            "Developed and delivered SOP training for new hires while contributing to documentation standardization projects",
            "Implemented safety protocols increasing workplace compliance 20% through targeted training and awareness initiatives",
            "Supported NPI activities via assembly process validation and new product testing protocols"
        ]
    }
    
    for entry in entries:
        company = entry.get('company', '').strip()
        title = entry.get('title', '').strip()
        location = entry.get('location', '').strip()
        dates = entry.get('dates', '').strip()
        
        # Company header with two-column layout
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell, right_cell = table.rows[0].cells
        
        # Left: Company | Title
        left_p = left_cell.paragraphs[0]
        company_run = left_p.add_run(company)
        company_run.bold = True
        company_run.font.size = Pt(11)
        company_run.font.name = FONT_PRIMARY
        
        left_p.add_run(' | ')
        
        title_run = left_p.add_run(title)
        title_run.font.size = Pt(10)
        title_run.font.name = FONT_PRIMARY
        
        # Right: Location | Dates
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        loc_run = right_p.add_run(f"{location} | {dates}")
        loc_run.font.size = Pt(10)
        loc_run.font.name = FONT_PRIMARY
        
        # Enhanced bullets for this company
        bullets = enhanced_bullets.get(company, entry.get('bullets', []))
        
        for bullet in bullets[:8]:  # Allow more bullets per role for 2-page format
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(bullet)
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = FONT_PRIMARY
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Small space between roles
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)

def write_education(doc):
    """Write education section."""
    add_section_title(doc, 'Education & Professional Development')
    
    # Degree
    p = doc.add_paragraph()
    degree_run = p.add_run('Bachelor of Science (B.S.) in Computer Information Systems')
    degree_run.bold = True
    degree_run.font.size = Pt(10)
    degree_run.font.name = FONT_PRIMARY
    p.paragraph_format.space_after = Pt(1)
    
    p2 = doc.add_paragraph('Post University, Waterbury, CT | 2020 – 2022')
    for r in p2.runs:
        r.font.size = Pt(10)
        r.font.name = FONT_PRIMARY
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.left_indent = Inches(0.25)
    
    # Additional
    p3 = doc.add_paragraph()
    honors_run = p3.add_run('Honors: ')
    honors_run.bold = True
    honors_run.font.size = Pt(10)
    honors_run.font.name = FONT_PRIMARY
    
    honors_text = p3.add_run('President\'s List (GPA 3.5+), National Society of Leadership and Success')
    honors_text.font.size = Pt(10)
    honors_text.font.name = FONT_PRIMARY
    p3.paragraph_format.left_indent = Inches(0.25)

def write_additional_info(doc):
    """Add certifications, clearances, or other relevant info."""
    add_section_title(doc, 'Additional Information')
    
    info = [
        "Military Service: U.S. Army National Guard (2012-2015) – Apache Repairer (15R)",
        "Clearance: Eligible for security clearance (aerospace industry background)",
        "Relocation: Open to relocation to Maplewood, MN or other 3M locations",
        "Travel: Available for up to 15% travel as required",
        "Technical Proficiencies: Microsoft Office Suite (Expert), ERP Systems (Aurora, Dynamics 365), G Suite, Technical Documentation Tools"
    ]
    
    for item in info:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.name = FONT_PRIMARY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

def build_enhanced_resume(summary_override=None, output_path=None, achievements_override=None, skill_categories_override=None):
    """Generate comprehensive 2-page ATS-friendly resume.
    
    Args:
        summary_override: Optional custom summary text to replace default
        output_path: Optional custom output path (uses OUT_ENHANCED if None)
    """
    contact = load_contact()
    experience = load_experience()
    
    doc = Document()
    apply_margins(doc)
    
    build_header(doc, contact)
    write_professional_summary(doc, summary_text=summary_override)
    write_key_achievements(doc, achievements=achievements_override)
    write_core_competencies(doc, custom_categories=skill_categories_override)
    write_professional_experience(doc, experience)
    write_education(doc)
    write_additional_info(doc)
    
    output_file = output_path if output_path else OUT_ENHANCED
    
    try:
        doc.save(output_file)
    except PermissionError:
        alt = output_file.replace('.docx', '_updated.docx')
        doc.save(alt)
        print(f"Enhanced 2-page resume generated: {alt}")
        return alt
    
    print(f"Enhanced 2-page resume generated: {output_file}")
    return output_file

if __name__ == '__main__':
    build_enhanced_resume()
