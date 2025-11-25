"""Cover Letter Generator

Generates a tailored cover letter DOCX leveraging match analysis and
candidate profile data. Integrates with the existing automation pipeline.
"""

from __future__ import annotations
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import re


GENERIC_REMOVALS = {"job","role","position","team","teams","company"}
MIN_KEYWORD_LEN = 4

def add_dans_metadata(doc, profile, company, position):
    """Add DANS-compliant document metadata for digital application navigation."""
    core_props = doc.core_properties
    core_props.author = profile.get('name', 'Ariel Karagodskiy')
    core_props.title = f"Cover Letter - {position} at {company}"
    core_props.subject = "Cover Letter - ATS & DANS Optimized"
    core_props.keywords = f"cover letter, ATS, DANS, {position}, {company}"
    core_props.category = "Cover Letter"
    core_props.comments = "Digital Application Navigation System (DANS) compliant with ATS optimization"

def configure_dans_layout(doc):
    """Configure DANS-compliant page layout with standard dimensions and margins."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)  # DANS standard
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def _sanitize_keywords(jd_keywords: list[str] | None, limit: int = 3) -> list[str]:
    if not jd_keywords:
        return []
    cleaned = []
    seen = set()
    for kw in jd_keywords:
        kw_clean = kw.strip().lower()
        kw_clean = kw_clean.replace('-', ' ')
        if kw_clean in GENERIC_REMOVALS:
            continue
        if len(kw_clean) < MIN_KEYWORD_LEN:
            continue
        # Skip if purely numeric or already used
        if kw_clean.isdigit() or kw_clean in seen:
            continue
        seen.add(kw_clean)
        # Normalize spacing
        kw_norm = ' '.join(w for w in kw_clean.split() if w not in GENERIC_REMOVALS)
        if kw_norm and kw_norm not in cleaned:
            cleaned.append(kw_norm.title())
        if len(cleaned) >= limit:
            break
    return cleaned

def _join_keywords_readable(words: list[str]) -> str:
    if not words:
        return ''
    if len(words) == 1:
        return words[0]
    if len(words) == 2:
        return f"{words[0]} and {words[1]}"
    return f"{', '.join(words[:-1])}, and {words[-1]}"

def build_cover_letter(company: str, position: str, jd_text: str, profile: dict, output_path: str, jd_keywords: list[str] | None = None) -> str:
    """Build DANS-compliant cover letter driven by JD keywords only."""
    doc = Document()
    add_dans_metadata(doc, profile, company, position)
    configure_dans_layout(doc)
    today = datetime.now().strftime("%B %d, %Y")

    # Header
    doc.add_paragraph(profile.get("name", ""))
    contact_line = f"{profile.get('email','')} | {profile.get('phone','')} | {profile.get('address','')}"
    doc.add_paragraph(contact_line)
    doc.add_paragraph("\n")

    # Date & Company
    doc.add_paragraph(today)
    doc.add_paragraph(company)
    doc.add_paragraph("\n")

    greeting = f"Dear Hiring Team,"
    doc.add_paragraph(greeting)
    doc.add_paragraph("\n")

    # Opening (JD-keyword + metric hook)
    keyword_phrase = ''
    top_keywords = _sanitize_keywords(jd_keywords, limit=3)
    if top_keywords:
        keyword_phrase = f" with strengths spanning {_join_keywords_readable(top_keywords)}"

    years = profile.get('years_experience', 0)
    hook_metric = ''
    achievements = profile.get('achievements', [])
    metric_src = next((a for a in achievements if re.search(r'\d+%|\d{2,}', str(a))), '')
    if metric_src:
        hook_metric = f" (including {metric_src.split('.')[0]})"

    # Prevent double period after company name
    company_clean = company.rstrip('.')
    opening = (
        f"I am writing to express my interest in the {position} role at {company_clean}. "
        f"With {years}+ years leading process engineering{keyword_phrase}{hook_metric}, I offer a proven ability to convert operational objectives into measurable improvements."
    )
    op_p = doc.add_paragraph(opening)
    for r in op_p.runs:
        r.font.size = Pt(11)
    doc.add_paragraph("\n")

    # Value proposition with specific examples (expand to 2/3 page)
    value1 = (
        "I specialize in optimizing manufacturing and technical workflows by aligning data, automation, and cross-functional collaboration to reduce waste and elevate quality. "
        "This approach blends hands-on problem solving with structured continuous improvement, producing durable gains rather than temporary fixes."
    )
    v1_p = doc.add_paragraph(value1)
    for r in v1_p.runs:
        r.font.size = Pt(11)
    doc.add_paragraph("\n")
    
    # Add key achievements paragraph
    if achievements:
        metric_achievements = [a for a in achievements if re.search(r'\d+%|\d{2,}', str(a))]
        top_ach = metric_achievements[:3] if metric_achievements else achievements[:2]
        if top_ach:
            ach_preamble = "Selected achievements demonstrating ROI and operational impact:" if metric_achievements else "Selected achievements:" 
            ach_para = doc.add_paragraph()
            run_head = ach_para.add_run(ach_preamble + " ")
            run_head.bold = True
            run_head.font.size = Pt(11)
            for idx, ach in enumerate(top_ach, 1):
                # Bold metrics inside achievements
                segments = re.split(r'(\d+%|\$\d+[\d,]*|\d{2,})', str(ach))
                bullet_run = ach_para.add_run(f"({idx}) ")
                bullet_run.bold = True
                for seg in segments:
                    if not seg:
                        continue
                    rseg = ach_para.add_run(seg)
                    if re.match(r'(\d+%|\$\d+[\d,]*|\d{2,})', seg):
                        rseg.bold = True
                    rseg.font.size = Pt(11)
                ach_para.add_run("  ")
            doc.add_paragraph("\n")
    
    # Alignment with company
    value2 = (
        f"I am drawn to {company_clean}'s focus on disciplined execution and innovation. "
        f"I would leverage my background in data-driven process optimization, stakeholder alignment, and continuous improvement to advance strategic objectives while increasing throughput and quality." 
    )
    v2_p = doc.add_paragraph(value2)
    for r in v2_p.runs:
        r.font.size = Pt(11)
    doc.add_paragraph("\n")

    closing = (
        f"I welcome the opportunity to discuss how my experience can advance {company_clean}'s near-term initiatives and long-term roadmap. "
        f"Thank you for your consideration; I look forward to the possibility of speaking with you."
    )
    c_p = doc.add_paragraph(closing)
    for r in c_p.runs:
        r.font.size = Pt(11)
    doc.add_paragraph("\n")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(profile.get("name", ""))

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)


if __name__ == "__main__":
    # Simple manual test placeholder
    sample_profile = {"name": "Sample Name", "email": "email@example.com", "phone": "123", "address": "City, ST", "years_experience": 10}
    sample_match = {"breakdown": {"tech_percent": 70, "process_percent": 65, "leadership_percent": 55, "npi_percent": 40}, "gaps": ["Regulatory", "CAD"]}
    build_cover_letter("Acme Corp", "Senior Engineer", "JD text here", sample_match, sample_profile, "cover_letter_test.docx")
    print("Generated cover_letter_test.docx")
