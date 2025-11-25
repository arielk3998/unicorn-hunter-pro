import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(__file__)
CONTACT_FILE = os.path.join(ROOT, 'contact_info.json')
RESUME_TXT = os.path.join(ROOT, 'tailored_resume_3M_short.txt')
COVER_TXT = os.path.join(ROOT, 'cover_letter_3M_short.txt')
RESUME_OUT = os.path.join(ROOT, 'tailored_resume_3M_short.docx')
COVER_OUT = os.path.join(ROOT, 'cover_letter_3M_short.docx')

ACCENT_COLOR = RGBColor(0x18, 0x3C, 0x5A)
BODY_FONT = 'Calibri'

def load_contact():
    if os.path.exists(CONTACT_FILE):
        with open(CONTACT_FILE,'r',encoding='utf-8') as f:
            return json.load(f)
    return {}

def read_lines(path):
    with open(path,'r',encoding='utf-8') as f:
        return [l.rstrip() for l in f]

def is_bullet(line):
    return line.strip().startswith(('•','-','*'))

def clean_bullet(text):
    return text.lstrip('•-* ').strip()

def add_dans_metadata(doc, contact, doc_type="Resume"):
    """Add DANS-compliant document metadata."""
    core_props = doc.core_properties
    core_props.author = contact.get('name', 'Ariel Karagodskiy')
    core_props.title = f"{contact.get('name', 'Candidate')} - Professional {doc_type}"
    core_props.subject = f"{doc_type} - ATS & DANS Optimized"
    core_props.keywords = f"{doc_type.lower()}, ATS, DANS, professional"
    core_props.category = doc_type
    core_props.comments = "Digital Application Navigation System (DANS) compliant"

def configure_dans_layout(doc):
    """Configure DANS-compliant page layout."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def build_resume(contact, lines):
    doc = Document()
    add_dans_metadata(doc, contact, "Resume")
    configure_dans_layout(doc)
    # Header - DANS: Use Heading 1
    header = doc.add_paragraph()
    header.style = 'Heading 1'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run(contact.get('name','Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT_COLOR
    r.font.name = BODY_FONT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(' • '.join([p for p in [contact.get('address',''), contact.get('location_statement','')] if p]))
    sub_run.font.size = Pt(9)
    sub_run.font.name = BODY_FONT

    contact_line = ' • '.join([p for p in [contact.get('email',''), contact.get('phone',''), contact.get('linkedin',''), contact.get('github','')] if p])
    cpara = doc.add_paragraph(contact_line)
    cpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rr in cpara.runs:
        rr.font.size = Pt(9)
        rr.font.name = BODY_FONT

    current = None
    buf = []

    def flush(sec, buf):
        if not sec or not buf: return
        p = doc.add_paragraph()
        rt = p.add_run(sec)
        rt.bold = True
        rt.font.size = Pt(11)
        rt.font.color.rgb = ACCENT_COLOR
        rt.font.name = BODY_FONT
        for b in buf:
            if is_bullet(b):
                pb = doc.add_paragraph(style='List Bullet')
                pb.add_run(clean_bullet(b))
            else:
                pb = doc.add_paragraph(b)
            for rrun in pb.runs:
                rrun.font.size = Pt(10)
                rrun.font.name = BODY_FONT

    for line in lines:
        if not line.strip():
            continue
        if line.isupper() and len(line.split()) < 8 and not is_bullet(line) and not line.startswith('Ariel'):
            flush(current, buf)
            current = line.title()
            buf = []
        else:
            buf.append(line)
    flush(current, buf)
    doc.save(RESUME_OUT)
    return RESUME_OUT

def build_cover_letter(contact, lines):
    doc = Document()
    add_dans_metadata(doc, contact, "Cover Letter")
    configure_dans_layout(doc)
    # Header - DANS: Use Heading 1
    header = doc.add_paragraph()
    header.style = 'Heading 1'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run(contact.get('name','Ariel Karagodskiy'))
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT_COLOR
    r.font.name = BODY_FONT

    info = ' • '.join([p for p in [contact.get('email',''), contact.get('phone','')] if p])
    ip = doc.add_paragraph(info)
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ip.runs:
        run.font.size = Pt(9)
        run.font.name = BODY_FONT

    for line in lines:
        if not line.strip():
            continue
        p = doc.add_paragraph(line)
        for rrun in p.runs:
            rrun.font.size = Pt(11)
            rrun.font.name = BODY_FONT
    try:
        doc.save(COVER_OUT)
    except PermissionError:
        alt = COVER_OUT.replace('.docx','_updated.docx')
        doc.save(alt)
    return COVER_OUT

def main():
    contact = load_contact()
    resume_lines = read_lines(RESUME_TXT)
    cover_lines = read_lines(COVER_TXT)
    r = build_resume(contact, resume_lines)
    c = build_cover_letter(contact, cover_lines)
    print('Generated condensed DOCX:', r, 'and', c)

if __name__ == '__main__':
    main()